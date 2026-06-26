"""
KRATOS Core — esqueleto FastAPI reutilizável (template para novos apps).

Inclui (genérico, sem domínio):
- Servir páginas estáticas (1 arquivo por página) + mount /frontend.
- Acesso por convite (token) com kill-switch ACCESS_CONTROL e chave-mestra ADMIN_TOKEN.
- Geração de relatório PDF (reportlab) e DOCX (python-docx) com compartilhamento.
- Health checks. Pronto para Passenger/a2wsgi sob subcaminho (root_path).

O QUE FAZER NO NOVO APP (procure por "TODO-DOMINIO"):
- Adicionar a fonte de dados do domínio e as rotas/funcionalidades específicas.
- Preencher a página principal (frontend/index.html) com o conteúdo do app.
"""
import os
import json
import time
import uuid
import secrets
import threading
from pathlib import Path

from fastapi import FastAPI, Request
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

APP_ROOT = Path(__file__).resolve().parent
FRONTEND_DIR = APP_ROOT / "frontend"
DATA_DIR = APP_ROOT / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)

APP_NAME = os.getenv("APP_NAME") or "KRATOS Core"
ADMIN_TOKEN = (os.getenv("ADMIN_TOKEN") or "").strip()
ACCESS_CONTROL_ON = (os.getenv("ACCESS_CONTROL") or "off").strip().lower() in ("1", "on", "true", "yes", "sim")
ACCESS_COOKIE = "kratos_access"
ACCESS_MAX_AGE = 60 * 60 * 24 * 30
_lock = threading.Lock()

app = FastAPI(title=APP_NAME)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True,
                   allow_methods=["*"], allow_headers=["*"])


# ===================== Acesso por convite (token) =====================
def _invites_file() -> Path:
    return DATA_DIR / "access_invites.json"

def _load_invites() -> list:
    try:
        p = _invites_file()
        if p.exists():
            d = json.loads(p.read_text(encoding="utf-8"))
            return list(d.get("invites") or []) if isinstance(d, dict) else (d if isinstance(d, list) else [])
    except (OSError, json.JSONDecodeError):
        pass
    return []

def _save_invites(items): _invites_file().write_text(json.dumps({"invites": items}, ensure_ascii=False, indent=2), encoding="utf-8")

def _invite_status(inv, now=None):
    now = now or time.time()
    if inv.get("revoked"): return "revogado"
    exp = inv.get("expiresAt")
    try:
        if exp and now > float(exp): return "expirado"
    except (TypeError, ValueError): pass
    return "ativo"

def create_invite(label, days=None):
    now = time.time()
    try: days = int(days) if days else None
    except (TypeError, ValueError): days = None
    inv = {"token": secrets.token_urlsafe(24), "label": (str(label or "").strip() or "Convidado")[:80],
           "createdAt": now, "expiresAt": (now + days * 86400) if days and days > 0 else None,
           "revoked": False, "lastAccessAt": None, "accessCount": 0}
    with _lock:
        items = _load_invites(); items.append(inv); _save_invites(items)
    return inv

def revoke_invite(token):
    with _lock:
        items = _load_invites(); found = False
        for inv in items:
            if inv.get("token") == token and not inv.get("revoked"): inv["revoked"] = True; found = True
        if found: _save_invites(items)
    return found

def token_valid(token):
    if not token: return False
    now = time.time()
    return any(inv.get("token") == token and _invite_status(inv, now) == "ativo" for inv in _load_invites())

def _public_invite(inv, now=None):
    return {"token": inv.get("token"), "label": inv.get("label"), "createdAt": inv.get("createdAt"),
            "expiresAt": inv.get("expiresAt"), "revoked": bool(inv.get("revoked")),
            "lastAccessAt": inv.get("lastAccessAt"), "accessCount": int(inv.get("accessCount") or 0),
            "status": _invite_status(inv, now)}

EXEMPT = ("/entrar", "/api/access/", "/admin", "/api/admin/", "/versao", "/healthz", "/favicon", "/robots.txt", "/frontend/")

def _app_path(request: Request):
    path = request.url.path
    root = (request.scope.get("root_path") or "").rstrip("/")
    if root and path.startswith(root): path = path[len(root):] or "/"
    return path, root

def _admin_ok(request: Request) -> bool:
    if not ADMIN_TOKEN: return False
    given = (request.headers.get("X-Admin-Token") or request.query_params.get("token") or "").strip()
    return given == ADMIN_TOKEN

@app.middleware("http")
async def access_gate(request: Request, call_next):
    if not ACCESS_CONTROL_ON:
        return await call_next(request)
    path, root = _app_path(request)
    if any(path == p or path.startswith(p) for p in EXEMPT):
        return await call_next(request)
    qt = (request.query_params.get("access") or "").strip()
    token = (qt or request.cookies.get(ACCESS_COOKIE) or request.headers.get("X-Access-Token") or "").strip()
    master = bool(ADMIN_TOKEN) and token == ADMIN_TOKEN
    if master or token_valid(token):
        resp = await call_next(request)
        if qt:
            secure = request.headers.get("x-forwarded-proto", request.url.scheme) == "https"
            resp.set_cookie(ACCESS_COOKIE, token, max_age=ACCESS_MAX_AGE, httponly=True, samesite="lax", secure=secure)
        return resp
    if path.startswith("/api/") or "application/json" in request.headers.get("accept", ""):
        return JSONResponse({"ok": False, "error": "acesso restrito — convite necessário"}, status_code=401)
    return RedirectResponse(url=f"{root}/entrar")

async def _json(request: Request):
    try:
        d = await request.json(); return d if isinstance(d, dict) else {}
    except Exception:
        return {}

@app.get("/api/access/status")
async def access_status(request: Request):
    token = (request.cookies.get(ACCESS_COOKIE) or request.query_params.get("access") or "").strip()
    ok = (not ACCESS_CONTROL_ON) or (bool(ADMIN_TOKEN) and token == ADMIN_TOKEN) or token_valid(token)
    return {"ok": True, "accessControl": ACCESS_CONTROL_ON, "hasValidAccess": bool(ok)}

@app.post("/api/access/validate")
async def access_validate(request: Request):
    token = str((await _json(request)).get("token") or "").strip()
    if not ((bool(ADMIN_TOKEN) and token == ADMIN_TOKEN) or token_valid(token)):
        return JSONResponse({"ok": True, "valid": False})
    resp = JSONResponse({"ok": True, "valid": True})
    secure = request.headers.get("x-forwarded-proto", request.url.scheme) == "https"
    resp.set_cookie(ACCESS_COOKIE, token, max_age=ACCESS_MAX_AGE, httponly=True, samesite="lax", secure=secure)
    return resp

@app.post("/api/access/logout")
async def access_logout():
    resp = JSONResponse({"ok": True}); resp.delete_cookie(ACCESS_COOKIE); return resp

@app.get("/api/admin/invites")
async def admin_list(request: Request):
    if not _admin_ok(request): return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    now = time.time()
    rows = sorted((_public_invite(i, now) for i in _load_invites()), key=lambda x: x.get("createdAt") or 0, reverse=True)
    return {"ok": True, "accessControl": ACCESS_CONTROL_ON, "invites": rows}

@app.post("/api/admin/invites")
async def admin_create(request: Request):
    if not _admin_ok(request): return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    d = await _json(request); return {"ok": True, "invite": _public_invite(create_invite(d.get("label"), d.get("expiresInDays")))}

@app.post("/api/admin/invites/revoke")
async def admin_revoke(request: Request):
    if not _admin_ok(request): return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    return {"ok": bool(revoke_invite(str((await _json(request)).get("token") or "").strip()))}


# ===================== Relatório PDF / DOCX (genérico) =====================
def _report_sections():
    """TODO-DOMINIO: devolva as seções reais do seu app.
    Formato: [(titulo, [linhas...]), ...]"""
    from datetime import datetime
    return [("Resumo", [f"{APP_NAME} — relatório gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}.",
                        "Substitua _report_sections() pelos dados reais do domínio."])]

def _render_pdf(sections, title):
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=A4, title=title, author="Jossian Brito")
    st = getSampleStyleSheet(); E = [Paragraph(title, st["Title"]), Spacer(1, 8)]
    for h, lines in sections:
        E.append(Paragraph(h, st["Heading2"]))
        for ln in lines: E.append(Paragraph(str(ln), st["BodyText"]))
        E.append(Spacer(1, 6))
    doc.build(E); return buf.getvalue()

def _render_docx(sections, title):
    import io
    from docx import Document
    doc = Document(); doc.add_heading(title, 0)
    for h, lines in sections:
        doc.add_heading(h, level=2)
        for ln in lines: doc.add_paragraph(str(ln))
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

@app.post("/api/report-file")
async def report_file(request: Request):
    import asyncio
    from datetime import datetime
    fmt = str((await _json(request)).get("format") or "pdf").lower().strip()
    if fmt not in ("pdf", "docx"):
        return JSONResponse({"ok": False, "error": "formato inválido"}, status_code=400)
    title = f"{APP_NAME} — Relatório"
    sections = _report_sections()
    try:
        if fmt == "pdf":
            blob = await asyncio.to_thread(_render_pdf, sections, title); media = "application/pdf"
        else:
            blob = await asyncio.to_thread(_render_docx, sections, title)
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except ImportError as exc:
        lib = "python-docx" if fmt == "docx" else "reportlab"
        return JSONResponse({"ok": False, "error": f"biblioteca {lib} não instalada ({exc})."}, status_code=400)
    fname = "Relatorio_" + datetime.now().strftime("%Y%m%d_%H%M") + "." + fmt
    return Response(content=blob, media_type=media, headers={"Content-Disposition": f'attachment; filename="{fname}"'})


# ===================== Páginas e health =====================
@app.get("/healthz")
@app.get("/api/health")
def health():
    return {"ok": True, "app": APP_NAME}

@app.get("/")
def root():
    return FileResponse(FRONTEND_DIR / "index.html", media_type="text/html")

@app.get("/entrar")
@app.get("/entrar/")
def entrar():
    return FileResponse(FRONTEND_DIR / "entrar.html", media_type="text/html")

@app.get("/admin")
@app.get("/admin/")
def admin():
    return FileResponse(FRONTEND_DIR / "admin.html", media_type="text/html")

@app.get("/versao")
@app.get("/versao/")
def versao():
    return FileResponse(FRONTEND_DIR / "versao.html", media_type="text/html")

# TODO-DOMINIO: adicionar aqui as rotas/funcionalidades específicas do seu app.

app.mount("/frontend", StaticFiles(directory=str(FRONTEND_DIR)), name="frontend")

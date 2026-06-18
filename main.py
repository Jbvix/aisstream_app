import os
import math
import unicodedata
from datetime import datetime
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse, Response
from dotenv import load_dotenv
import asyncio
import json
import websockets
import time
import threading
import uuid
import secrets
import urllib.request
import urllib.error
from pathlib import Path
from collections import deque
from websockets.exceptions import ConnectionClosed

import praticagem_saa
import obsidian_supabase
import obsidian_notes

load_dotenv()

APP_ROOT = Path(__file__).resolve().parent
FRONTEND_DIR = APP_ROOT / "frontend"

PORT = int(os.getenv("PORT", 8080))
AIS_MODE = os.getenv("AIS_MODE", "mock").lower()
AISSTREAM_API_KEY = os.getenv("AISSTREAM_API_KEY", "")
DEFAULT_AREA = os.getenv("DEFAULT_AREA", "rio").lower()
AISSTREAM_URL = "wss://stream.aisstream.io/v0/stream"
SAAM_BGRA_FLEET_NAME = "SAAM-BGRA"
SAAM_BGRA_NAMES = {
    "710012550": "SAAM PATAXO",
    "710001249": "SAAM PARECI",
    "710021750": "SAAM CHILE",
    "710001593": "SAAM HOLANDA",
    "710016030": "SAAM LANCELOT",
    "710015310": "SAAM ARTHUR",
}
SAAM_BGRA_MMSI_SET = set(SAAM_BGRA_NAMES.keys())
# BBox operacional da Baia de Guanabara (lon/lat aproximados)
GUANABARA_GEOFENCE = {
    "name": "Baia de Guanabara",
    "minLat": -23.08,
    "maxLat": -22.75,
    "minLon": -43.35,
    "maxLon": -43.05,
}
BG_INTERNO_GEOFENCE_ID = "bg-interno-persistente"
BG_INTERNO_GEOFENCE_NAME = "Baia de Guanabara Interno"
LEGACY_GEOFENCE_STORAGE_PATH = APP_ROOT / "data" / "geofences.json"
DASHBOARD_USER_ID = (os.getenv("DASHBOARD_USER_ID") or "default").strip() or "default"
SAAM_MMSI_ABBR = {
    "710012550": "PX",
    "710001249": "PA",
    "710021750": "CH",
    "710001593": "HL",
    "710016030": "LT",
    "710015310": "AT",
}
# Na programação da Praticagem (campo EMP.RB), o código "SAA" representa a SAAM,
# que é a nossa empresa. WIL e CAM são concorrentes.
OWN_COMPANY_EMP_RB = "SAA"
COMPETITOR_TUGS = {
    "WIL": [
        {"mmsi": "710005290", "name": "LYRA"},
        {"mmsi": "710000009", "name": "HERCULES"},
        {"mmsi": "710000391", "name": "WEZEN"},
        {"mmsi": "710018340", "name": "PEGASUS"},
    ],
    "CAM": [
        {"mmsi": "710009550", "name": "C NEBLINA"},
        {"mmsi": "710008322", "name": "C HARPIA"},
        {"mmsi": "710010250", "name": "C ARRAIAL / C SALVADOR"},
    ],
}
COMPETITOR_COMPANY_BY_MMSI = {
    t["mmsi"]: company for company, tugs in COMPETITOR_TUGS.items() for t in tugs
}
GROK_API_KEY = (os.getenv("XAI_API_KEY") or "").strip()
GROK_MODEL = (os.getenv("XAI_MODEL") or "grok-3-mini").strip()
ASSISTANT_PROFILE = (os.getenv("ASSISTANT_PROFILE") or "hibrido").strip().lower()
# Estatísticas persistidas (tug_geofence_stats.json): só berço e polígono contam manobra + tempo.
# Saída da base rebocador não soma manobra nem horas nesse ficheiro.
SAAM_MANEUVER_STATS_GEOFENCE_TYPES = frozenset({"berco", "polygon"})

# Controle de horas de operação do rebocador (prevenção de fadiga — reunião 12/06).
# Operação = qualquer movimento (manobra OU deslocamento); repouso = atracado/parado.
# Após TUG_OP_LIMIT_HOURS de operação, o rebocador deve entrar em recuperação;
# TUG_RECOVERY_HOURS de repouso contínuo zeram a fadiga (revezamento).
TUG_OP_LIMIT_HOURS = float(os.getenv("TUG_OP_LIMIT_HOURS") or 8.0)
TUG_OP_WARN_HOURS = float(os.getenv("TUG_OP_WARN_HOURS") or 7.0)
TUG_RECOVERY_HOURS = float(os.getenv("TUG_RECOVERY_HOURS") or 8.0)
TUG_OP_MOVING_SOG = 0.5          # nós: acima disso o rebocador está em movimento (operação)
TUG_OP_MAX_GAP_SEC = 600         # lacuna > 10 min = sinal perdido; não contabiliza
_last_op_save_ts = 0.0

# Navio comercial (não SAAM-BGRA) → indicador «manobra SAA» no dashboard (berço/polígono/base)
SAA_MANEUVER_SHIP_CATEGORIES = frozenset(
    {"carga", "petroleiro", "passageiros", "lazer", "pesca", "outros"}
)


def _vessel_matches_saa_maneuver_lamp(vessel):
    if not isinstance(vessel, dict):
        return False
    if bool(vessel.get("isSaamBgra")):
        return False
    cat = vessel.get("shipCategory") or "outros"
    return cat in SAA_MANEUVER_SHIP_CATEGORIES


def _vessel_display_name_for_tooltip(vessel):
    if not isinstance(vessel, dict):
        return "—"
    name = (vessel.get("shipName") or "").strip()
    if name:
        return name
    mmsi = vessel.get("mmsi")
    return f"MMSI {mmsi}" if mmsi is not None else "—"


def _distance_m_approx(v1, v2):
    lat1, lon1 = v1.get("latitude"), v1.get("longitude")
    lat2, lon2 = v2.get("latitude"), v2.get("longitude")
    if not all(isinstance(x, (int, float)) for x in (lat1, lon1, lat2, lon2)):
        return float("inf")
    dy = (float(lat1) - float(lat2)) * 111_320.0
    dx = (float(lon1) - float(lon2)) * 111_320.0
    return (dx * dx + dy * dy) ** 0.5


def _berco_em_manobra_text(inside_geo):
    tugs = [v for v in inside_geo if bool(v.get("isSaamBgra"))]
    ships = [v for v in inside_geo if not bool(v.get("isSaamBgra"))]
    if not tugs:
        return "—"
    if not ships:
        tug_names = ", ".join(_vessel_display_name_for_tooltip(v) for v in tugs[:4])
        return f"{tug_names} (sem navio)"
    pairs = []
    for tug in tugs[:4]:
        nearest_ship = min(ships, key=lambda ship: _distance_m_approx(tug, ship))
        pairs.append(
            f"{_vessel_display_name_for_tooltip(tug)} + {_vessel_display_name_for_tooltip(nearest_ship)}"
        )
    return " | ".join(pairs)

AREAS = {
    "suape": {
        "name": "Suape",
        "center": [-8.393, -34.968],
        "zoom": 11,
        "boundingBoxes": [[[-8.25, -35.15], [-8.55, -34.75]]]
    },
    "santos": {
        "name": "Santos",
        "center": [-23.975, -46.33],
        "zoom": 11,
        "boundingBoxes": [[[-23.82, -46.5], [-24.15, -46.15]]]
    },
    "rio": {
        "name": "Rio de Janeiro",
        "center": [-22.895, -43.165],
        "zoom": 11,
        "boundingBoxes": [[[-22.75, -43.4], [-23.05, -42.95]]]
    },
    "paranagua": {
        "name": "Paranaguá",
        "center": [-25.509, -48.505],
        "zoom": 11,
        "boundingBoxes": [[[-25.35, -48.75], [-25.75, -48.3]]]
    },
    "bahia": {
        "name": "Baía de Todos-os-Santos",
        "center": [-12.900, -38.516],
        "zoom": 11,
        "boundingBoxes": [[[-12.7, -38.8], [-13.1, -38.2]]]
    },
    "mucuripe_pecem": {
        "name": "Mucuripe / Pecem",
        "center": [-3.66, -38.65],
        "zoom": 10,
        "boundingBoxes": [[[-3.40, -38.95], [-3.95, -38.35]]]
    },
    "itaguai": {
        "name": "Itaguai",
        "center": [-22.93, -43.85],
        "zoom": 10,
        "boundingBoxes": [[[-22.70, -44.20], [-23.20, -43.55]]]
    },
    "vitoria": {
        "name": "Vitoria",
        "center": [-20.31, -40.29],
        "zoom": 10,
        "boundingBoxes": [[[-20.10, -40.55], [-20.55, -40.05]]]
    },
    "rio_grande": {
        "name": "Rio Grande",
        "center": [-32.12, -52.10],
        "zoom": 10,
        "boundingBoxes": [[[-31.90, -52.40], [-32.35, -51.80]]]
    },
    "itajai": {
        "name": "Itajai",
        "center": [-26.91, -48.67],
        "zoom": 10,
        "boundingBoxes": [[[-26.70, -48.90], [-27.15, -48.45]]]
    },
    "sao_francisco_do_sul": {
        "name": "Sao Francisco do Sul",
        "center": [-26.24, -48.64],
        "zoom": 10,
        "boundingBoxes": [[[-26.05, -48.85], [-26.45, -48.40]]]
    },
    "rotterdam": {
        "name": "Rotterdam",
        "center": [51.94, 4.14],
        "zoom": 10,
        "boundingBoxes": [[[51.75, 3.80], [52.20, 4.60]]]
    },
    "antwerp_bruges": {
        "name": "Antwerp-Bruges",
        "center": [51.29, 4.32],
        "zoom": 10,
        "boundingBoxes": [[[51.10, 3.95], [51.55, 4.75]]]
    },
    "hamburg": {
        "name": "Hamburg",
        "center": [53.54, 9.96],
        "zoom": 10,
        "boundingBoxes": [[[53.35, 9.55], [53.75, 10.35]]]
    },
    "algeciras": {
        "name": "Algeciras",
        "center": [36.13, -5.45],
        "zoom": 10,
        "boundingBoxes": [[[35.95, -5.85], [36.35, -5.10]]]
    },
    "tangier_med": {
        "name": "Tangier Med",
        "center": [35.88, -5.50],
        "zoom": 10,
        "boundingBoxes": [[[35.70, -5.90], [36.10, -5.10]]]
    },
    "suez": {
        "name": "Suez Canal",
        "center": [30.60, 32.33],
        "zoom": 9,
        "boundingBoxes": [[[29.95, 31.85], [31.10, 32.95]]]
    },
    "jebel_ali": {
        "name": "Jebel Ali",
        "center": [25.02, 55.06],
        "zoom": 10,
        "boundingBoxes": [[[24.80, 54.75], [25.30, 55.35]]]
    },
    "singapore": {
        "name": "Singapore",
        "center": [1.23, 103.84],
        "zoom": 10,
        "boundingBoxes": [[[1.05, 103.55], [1.45, 104.15]]]
    },
    "shanghai": {
        "name": "Shanghai",
        "center": [31.33, 121.75],
        "zoom": 9,
        "boundingBoxes": [[[30.95, 121.20], [31.80, 122.40]]]
    },
    "ningbo_zhoushan": {
        "name": "Ningbo-Zhoushan",
        "center": [29.93, 122.24],
        "zoom": 9,
        "boundingBoxes": [[[29.40, 121.70], [30.45, 122.85]]]
    },
    "shenzhen": {
        "name": "Shenzhen",
        "center": [22.55, 114.20],
        "zoom": 10,
        "boundingBoxes": [[[22.30, 113.80], [22.85, 114.55]]]
    },
    "hong_kong": {
        "name": "Hong Kong",
        "center": [22.30, 114.17],
        "zoom": 10,
        "boundingBoxes": [[[22.15, 113.90], [22.55, 114.45]]]
    },
    "busan": {
        "name": "Busan",
        "center": [35.10, 129.04],
        "zoom": 10,
        "boundingBoxes": [[[34.90, 128.70], [35.35, 129.40]]]
    },
    "colombo": {
        "name": "Colombo",
        "center": [6.95, 79.85],
        "zoom": 10,
        "boundingBoxes": [[[6.75, 79.60], [7.20, 80.10]]]
    },
    "los_angeles_long_beach": {
        "name": "Los Angeles / Long Beach",
        "center": [33.74, -118.24],
        "zoom": 10,
        "boundingBoxes": [[[33.55, -118.60], [34.00, -117.95]]]
    },
    "new_york_new_jersey": {
        "name": "New York / New Jersey",
        "center": [40.64, -74.07],
        "zoom": 10,
        "boundingBoxes": [[[40.40, -74.35], [40.95, -73.70]]]
    },
    "houston": {
        "name": "Houston",
        "center": [29.72, -95.18],
        "zoom": 10,
        "boundingBoxes": [[[29.45, -95.55], [30.05, -94.85]]]
    },
    "panama_canal": {
        "name": "Panama Canal",
        "center": [9.11, -79.66],
        "zoom": 10,
        "boundingBoxes": [[[8.85, -79.95], [9.45, -79.35]]]
    },
    "valparaiso_san_antonio": {
        "name": "Valparaiso / San Antonio",
        "center": [-33.35, -71.64],
        "zoom": 10,
        "boundingBoxes": [[[-33.70, -72.10], [-32.95, -71.30]]]
    },
    "durban": {
        "name": "Durban",
        "center": [-29.87, 31.04],
        "zoom": 10,
        "boundingBoxes": [[[-30.10, 30.70], [-29.55, 31.35]]]
    },
    "sydney_botany": {
        "name": "Sydney (Botany)",
        "center": [-33.95, 151.22],
        "zoom": 10,
        "boundingBoxes": [[[-34.15, 150.95], [-33.65, 151.50]]]
    },
    "brasil_sudeste": {
        "name": "Brasil Sudeste",
        "center": [-23.8, -43.5],
        "zoom": 6,
        "boundingBoxes": [[[-23.3, -42.5], [-24.2, -44.1]]]
    },
    "miami_teste": {
        "name": "Miami (Teste Live)",
        "center": [25.72, -80.04],
        "zoom": 10,
        "boundingBoxes": [[[25.835302, -80.207729], [25.602700, -79.879297]]]
    },
    "mundo_teste": {
        "name": "Mundo (Diagnóstico)",
        "center": [0.0, 0.0],
        "zoom": 2,
        "boundingBoxes": [[[-90, -180], [90, 180]]]
    }
}


from fastapi.staticfiles import StaticFiles

app = FastAPI()


@app.get("/")
def root():
    return FileResponse(FRONTEND_DIR / "index.html", media_type="text/html")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def access_gate_middleware(request: Request, call_next):
    """Trava o app (mapa + painel) atras de token de convite quando ACCESS_CONTROL_ON.
    Admin (/admin, /api/admin/*) e a propria pagina de entrada ficam de fora; o
    ADMIN_TOKEN funciona como chave-mestra."""
    if not ACCESS_CONTROL_ON:
        return await call_next(request)
    path, root = _app_relative_path(request)
    if _is_access_exempt(path):
        return await call_next(request)
    qtoken = (request.query_params.get("access") or "").strip()
    token = (
        qtoken
        or request.cookies.get(ACCESS_COOKIE_NAME)
        or request.headers.get("X-Access-Token")
        or ""
    ).strip()
    is_master = bool(ADMIN_TOKEN) and token == ADMIN_TOKEN
    if is_master or access_token_is_valid(token):
        response = await call_next(request)
        if qtoken:  # chegou pelo link de convite -> fixa o cookie de sessao
            secure = request.headers.get("x-forwarded-proto", request.url.scheme) == "https"
            response.set_cookie(
                ACCESS_COOKIE_NAME, token, max_age=ACCESS_COOKIE_MAX_AGE,
                httponly=True, samesite="lax", secure=secure,
            )
        if not is_master:
            try:
                _touch_access_invite(token)
            except Exception:
                pass
        return response
    accept = request.headers.get("accept", "")
    if path.startswith("/api/") or "application/json" in accept:
        return JSONResponse({"ok": False, "error": "acesso restrito — convite necessário"}, status_code=401)
    return RedirectResponse(url=f"{root}/entrar")

current_area_key = DEFAULT_AREA if DEFAULT_AREA in AREAS else "rio"
current_mode = "live"
live_connected = False
last_error = None
last_ais_message_at = None
total_messages = 0
live_subscription_update_event = asyncio.Event()
last_subscription_update_monotonic = 0.0
vessel_state_by_mmsi = {}
latest_vessel_by_mmsi = {}
recent_vessels = deque(maxlen=4000)
last_vessel_seq = 0

# Persistência do snapshot de embarcações: sobrevive a restart do Passenger
# (deploy/cPanel), evitando o mapa abrir vazio até o AIS repovoar.
vessels_snapshot_lock = threading.Lock()
_last_vessels_persist_monotonic = 0.0
VESSELS_PERSIST_MIN_INTERVAL_SEC = 20.0
# Posições mais antigas que isto não são restauradas no boot (evita "fantasmas").
VESSELS_SNAPSHOT_MAX_AGE_SEC = 6 * 60 * 60
live_worker_task = None
_praticagem_auto_sync_task: asyncio.Task | None = None
_obsidian_auto_sync_task: asyncio.Task | None = None
_obsidian_last_export_ts = 0.0  # debounce: timestamp do último export bem-sucedido
_obsidian_export_running = False
live_worker_thread = None
live_worker_lock = threading.Lock()
geofences = []
geofence_lock = threading.Lock()
saam_geofence_stats_lock = threading.Lock()
last_saam_inside_geofences = {}
saam_geofence_sessions = {}
saam_last_position_for_nm = {}
_nm_save_budget = {"n": 0, "last_mono": 0.0}
tug_stats_state = {"byMmsi": {}, "updatedAt": None}
saa_maneuvers_lock = threading.RLock()
saa_maneuvers_list = []


def user_data_dir(user_id: str) -> Path:
    p = APP_ROOT / "data" / "users" / user_id
    p.mkdir(parents=True, exist_ok=True)
    return p


def geofences_file_for_user(user_id: str) -> Path:
    return user_data_dir(user_id) / "geofences.json"


def ensure_data_dir():
    (APP_ROOT / "data").mkdir(parents=True, exist_ok=True)


def vessels_snapshot_file() -> Path:
    return user_data_dir(DASHBOARD_USER_ID) / "vessels_snapshot.json"


def save_vessels_snapshot(force: bool = False):
    """Grava o snapshot de embarcações em disco (com throttle)."""
    global _last_vessels_persist_monotonic
    now = time.monotonic()
    if not force and (now - _last_vessels_persist_monotonic) < VESSELS_PERSIST_MIN_INTERVAL_SEC:
        return
    _last_vessels_persist_monotonic = now
    try:
        items = []
        for v in list(latest_vessel_by_mmsi.values()):
            if not isinstance(v, dict):
                continue
            items.append({k: val for k, val in v.items() if k != "raw"})
        payload = {"savedAt": int(time.time()), "vessels": items}
        path = vessels_snapshot_file()
        tmp = path.with_suffix(".json.tmp")
        with vessels_snapshot_lock:
            tmp.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            tmp.replace(path)
    except OSError:
        pass


def load_vessels_snapshot():
    """Restaura o último snapshot conhecido no boot (descarta posições antigas)."""
    path = vessels_snapshot_file()
    try:
        if not path.exists():
            return
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return
    saved_at = payload.get("savedAt")
    if isinstance(saved_at, (int, float)) and (time.time() - saved_at) > VESSELS_SNAPSHOT_MAX_AGE_SEC:
        return
    for v in payload.get("vessels", []):
        mmsi = str(v.get("mmsi") or "").strip()
        if not mmsi:
            continue
        v["restoredFromSnapshot"] = True
        latest_vessel_by_mmsi[mmsi] = v


# ===== Frota dinâmica (SAAM e concorrentes editáveis em runtime) =====

def fleet_config_file() -> Path:
    return user_data_dir(DASHBOARD_USER_ID) / "fleet_config.json"


def _current_fleet_config() -> dict:
    return {
        "saam": [
            {"mmsi": m, "name": SAAM_BGRA_NAMES.get(m, ""), "abbr": SAAM_MMSI_ABBR.get(m, m[-2:])}
            for m in sorted(SAAM_BGRA_MMSI_SET)
        ],
        "competitors": {
            company: [dict(t) for t in tugs] for company, tugs in COMPETITOR_TUGS.items()
        },
    }


def _rebuild_fleet_structures(saam: list, competitors: dict):
    """Atualiza in-place as estruturas globais de frota (todas as referências
    do app passam a enxergar a frota nova) e corrige o flag isSaamBgra do buffer."""
    SAAM_BGRA_NAMES.clear()
    SAAM_MMSI_ABBR.clear()
    for t in saam:
        m = str(t.get("mmsi") or "").strip()
        if not m:
            continue
        SAAM_BGRA_NAMES[m] = str(t.get("name") or f"SAAM {m[-4:]}").strip()
        SAAM_MMSI_ABBR[m] = (str(t.get("abbr") or "").strip() or m[-2:]).upper()[:3]
    SAAM_BGRA_MMSI_SET.clear()
    SAAM_BGRA_MMSI_SET.update(SAAM_BGRA_NAMES.keys())
    COMPETITOR_TUGS.clear()
    for company, tugs in (competitors or {}).items():
        rows = []
        for t in tugs or []:
            m = str(t.get("mmsi") or "").strip()
            if m:
                rows.append({"mmsi": m, "name": str(t.get("name") or m).strip()})
        COMPETITOR_TUGS[str(company).strip().upper()] = rows
    COMPETITOR_COMPANY_BY_MMSI.clear()
    COMPETITOR_COMPANY_BY_MMSI.update(
        {t["mmsi"]: company for company, tugs in COMPETITOR_TUGS.items() for t in tugs}
    )
    # Reclassifica embarcações já no buffer (o flag é gravado na ingestão).
    for m, v in list(latest_vessel_by_mmsi.items()):
        if isinstance(v, dict):
            v["isSaamBgra"] = m in SAAM_BGRA_MMSI_SET


def save_fleet_config():
    try:
        fleet_config_file().write_text(
            json.dumps(_current_fleet_config(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except OSError:
        pass


_fleet_config_loaded = False


def ensure_fleet_config_loaded():
    """Carrega a frota persistida (se existir) sobrepondo os padrões do código."""
    global _fleet_config_loaded
    if _fleet_config_loaded:
        return
    _fleet_config_loaded = True
    try:
        p = fleet_config_file()
        if not p.exists():
            return
        cfg = json.loads(p.read_text(encoding="utf-8"))
        if isinstance(cfg, dict) and (cfg.get("saam") or cfg.get("competitors")):
            _rebuild_fleet_structures(cfg.get("saam") or [], cfg.get("competitors") or {})
    except (OSError, json.JSONDecodeError):
        pass


def migrate_legacy_geofences_if_needed(user_id: str):
    dest = geofences_file_for_user(user_id)
    if dest.exists():
        return
    if LEGACY_GEOFENCE_STORAGE_PATH.exists():
        try:
            dest.parent.mkdir(parents=True, exist_ok=True)
            dest.write_text(LEGACY_GEOFENCE_STORAGE_PATH.read_text(encoding="utf-8"), encoding="utf-8")
        except Exception:
            pass


def _normalize_text(value) -> str:
    return (
        unicodedata.normalize("NFD", str(value or ""))
        .encode("ascii", "ignore")
        .decode("ascii")
        .strip()
        .lower()
    )


def _is_bg_interno_geofence(geofence: dict) -> bool:
    if not isinstance(geofence, dict):
        return False
    if str(geofence.get("id") or "").strip() == BG_INTERNO_GEOFENCE_ID:
        return True
    return _normalize_text(geofence.get("name")) == _normalize_text(BG_INTERNO_GEOFENCE_NAME)


def _default_bg_interno_geometry():
    min_lat = GUANABARA_GEOFENCE["minLat"]
    max_lat = GUANABARA_GEOFENCE["maxLat"]
    min_lon = GUANABARA_GEOFENCE["minLon"]
    max_lon = GUANABARA_GEOFENCE["maxLon"]
    return {
        "coordinates": [
            [max_lat, min_lon],
            [max_lat, max_lon],
            [min_lat, max_lon],
            [min_lat, min_lon],
        ]
    }


def ensure_bg_interno_geofence():
    changed = False
    target = None
    for g in geofences:
        if _is_bg_interno_geofence(g):
            target = g
            break
    if target is None:
        now = get_now_iso()
        geofences.append(
            {
                "id": BG_INTERNO_GEOFENCE_ID,
                "name": BG_INTERNO_GEOFENCE_NAME,
                "type": "polygon",
                "geometry": _default_bg_interno_geometry(),
                "fleetScope": "all",
                "isActive": True,
                "color": "#4aa8ff",
                "persistent": True,
                "createdAt": now,
                "updatedAt": now,
            }
        )
        changed = True
    else:
        if target.get("id") != BG_INTERNO_GEOFENCE_ID:
            target["id"] = BG_INTERNO_GEOFENCE_ID
            changed = True
        if target.get("name") != BG_INTERNO_GEOFENCE_NAME:
            target["name"] = BG_INTERNO_GEOFENCE_NAME
            changed = True
        if target.get("type") != "polygon":
            target["type"] = "polygon"
            changed = True
        if not isinstance(target.get("geometry"), dict) or not target.get("geometry", {}).get("coordinates"):
            target["geometry"] = _default_bg_interno_geometry()
            changed = True
        if target.get("fleetScope") != "all":
            target["fleetScope"] = "all"
            changed = True
        if target.get("isActive") is not True:
            target["isActive"] = True
            changed = True
        if target.get("persistent") is not True:
            target["persistent"] = True
            changed = True
        if changed:
            target["updatedAt"] = get_now_iso()
    return changed


def load_geofences():
    global geofences
    ensure_data_dir()
    uid = DASHBOARD_USER_ID
    migrate_legacy_geofences_if_needed(uid)
    path = geofences_file_for_user(uid)
    if not path.exists():
        geofences = []
        return
    try:
        content = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(content, list):
            geofences = content
        else:
            geofences = []
    except Exception:
        geofences = []
    if ensure_bg_interno_geofence():
        save_geofences()


def ensure_geofences_loaded():
    """Lazy-load defensivo para ambientes Passenger onde startup pode não hidratar estado."""
    if geofences:
        return
    with geofence_lock:
        if geofences:
            return
        load_geofences()


def save_geofences():
    uid = DASHBOARD_USER_ID
    path = geofences_file_for_user(uid)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(geofences, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def tug_stats_file_for_user(user_id: str) -> Path:
    return user_data_dir(user_id) / "tug_geofence_stats.json"


def load_tug_stats():
    global tug_stats_state
    path = tug_stats_file_for_user(DASHBOARD_USER_ID)
    if not path.exists():
        tug_stats_state = {"byMmsi": {}, "updatedAt": get_now_iso()}
        return
    try:
        tug_stats_state = json.loads(path.read_text(encoding="utf-8"))
        if "byMmsi" not in tug_stats_state:
            tug_stats_state["byMmsi"] = {}
    except Exception:
        tug_stats_state = {"byMmsi": {}, "updatedAt": get_now_iso()}


def save_tug_stats():
    path = tug_stats_file_for_user(DASHBOARD_USER_ID)
    with saam_geofence_stats_lock:
        tug_stats_state["updatedAt"] = get_now_iso()
        path.write_text(
            json.dumps(tug_stats_state, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


def saa_maneuvers_file_for_user(user_id: str) -> Path:
    return user_data_dir(user_id) / "saa_maneuvers.json"


def strategy_memory_file_for_user(user_id: str) -> Path:
    return user_data_dir(user_id) / "strategy_memory.json"


def schedule_monitor_file_for_user(user_id: str) -> Path:
    return user_data_dir(user_id) / "saa_schedule_monitor.json"


def load_saa_maneuvers():
    global saa_maneuvers_list
    path = saa_maneuvers_file_for_user(DASHBOARD_USER_ID)
    if not path.exists():
        saa_maneuvers_list = []
        return
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
        saa_maneuvers_list = raw if isinstance(raw, list) else raw.get("items", [])
    except Exception:
        saa_maneuvers_list = []


def ensure_saa_maneuvers_loaded():
    """Se a lista em memória estiver vazia, relê saa_maneuvers.json (startup parcial ou processo sem on_event)."""
    path = saa_maneuvers_file_for_user(DASHBOARD_USER_ID)
    if not path.exists():
        return
    with saa_maneuvers_lock:
        if saa_maneuvers_list:
            return
        try:
            raw = json.loads(path.read_text(encoding="utf-8"))
            data = raw if isinstance(raw, list) else raw.get("items", [])
            if data:
                saa_maneuvers_list[:] = data
        except Exception:
            pass


def save_saa_maneuvers():
    path = saa_maneuvers_file_for_user(DASHBOARD_USER_ID)
    with saa_maneuvers_lock:
        path.write_text(json.dumps(saa_maneuvers_list, ensure_ascii=False, indent=2), encoding="utf-8")


def _saa_maneuver_dedupe_key(item: dict):
    def norm(v):
        return str(v or "").strip().upper()

    return (
        norm(item.get("source")),
        norm(item.get("pob")),
        norm(item.get("vesselName")),
        norm(item.get("berthName")),
        norm(item.get("empRb")),
        norm(item.get("status")),
        norm(item.get("cal")),
        norm(item.get("loa")),
        norm(item.get("boca")),
        norm(item.get("dwt")),
        norm(item.get("gt")),
        norm(item.get("m")),
    )


def _dedupe_saa_maneuvers(items):
    seen = set()
    out = []
    for item in items or []:
        key = _saa_maneuver_dedupe_key(item if isinstance(item, dict) else {})
        if key in seen:
            continue
        seen.add(key)
        out.append(item)
    return out


def load_strategy_memory():
    path = strategy_memory_file_for_user(DASHBOARD_USER_ID)
    if not path.exists():
        return []
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
        return raw if isinstance(raw, list) else []
    except Exception:
        return []


def append_strategy_memory(note: str, author: str = "user"):
    note = (note or "").strip()
    if not note:
        return
    items = load_strategy_memory()
    items.insert(0, {"at": get_now_iso(), "author": author, "note": note})
    path = strategy_memory_file_for_user(DASHBOARD_USER_ID)
    path.write_text(json.dumps(items[:200], ensure_ascii=False, indent=2), encoding="utf-8")


# ===== Telemetria / auditoria do KRATOS (página do administrador) =====
ADMIN_TOKEN = (os.getenv("ADMIN_TOKEN") or "").strip()


# ===== Controle de acesso por convite (token) =====
# Quando ACCESS_CONTROL_ON, o app (mapa + painel) passa a exigir um token de
# convite valido (cookie, ?access= ou cabecalho). O ADMIN_TOKEN funciona sempre
# como chave-mestra, evitando lockout do dono. Desligado por padrao: o deploy do
# codigo nao tranca o site ate o dono ativar ACCESS_CONTROL=on no servidor.
ACCESS_CONTROL_ON = (os.getenv("ACCESS_CONTROL") or "off").strip().lower() in ("1", "on", "true", "yes", "sim")
ACCESS_COOKIE_NAME = "kratos_access"
ACCESS_COOKIE_MAX_AGE = 60 * 60 * 24 * 30  # 30 dias
_access_lock = threading.Lock()


def access_invites_file() -> Path:
    return user_data_dir(DASHBOARD_USER_ID) / "access_invites.json"


def _load_access_invites() -> list:
    try:
        p = access_invites_file()
        if p.exists():
            data = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return list(data.get("invites") or [])
            if isinstance(data, list):
                return data
    except (OSError, json.JSONDecodeError):
        pass
    return []


def _save_access_invites(invites: list):
    access_invites_file().write_text(
        json.dumps({"invites": invites}, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _invite_status(inv: dict, now: float | None = None) -> str:
    now = now if now is not None else time.time()
    if inv.get("revoked"):
        return "revogado"
    exp = inv.get("expiresAt")
    try:
        if exp and now > float(exp):
            return "expirado"
    except (TypeError, ValueError):
        pass
    return "ativo"


def create_access_invite(label: str, expires_in_days=None) -> dict:
    now = time.time()
    try:
        days = int(expires_in_days) if expires_in_days else None
    except (TypeError, ValueError):
        days = None
    inv = {
        "token": secrets.token_urlsafe(24),
        "label": (str(label or "").strip() or "Convidado")[:80],
        "createdAt": now,
        "expiresAt": (now + days * 86400) if days and days > 0 else None,
        "revoked": False,
        "lastAccessAt": None,
        "accessCount": 0,
    }
    with _access_lock:
        invites = _load_access_invites()
        invites.append(inv)
        _save_access_invites(invites)
    return inv


def revoke_access_invite(token: str) -> bool:
    found = False
    with _access_lock:
        invites = _load_access_invites()
        for inv in invites:
            if inv.get("token") == token and not inv.get("revoked"):
                inv["revoked"] = True
                found = True
        if found:
            _save_access_invites(invites)
    return found


def access_token_is_valid(token: str) -> bool:
    if not token:
        return False
    now = time.time()
    for inv in _load_access_invites():
        if inv.get("token") == token and _invite_status(inv, now) == "ativo":
            return True
    return False


def _touch_access_invite(token: str):
    """Atualiza ultimo acesso no maximo 1x/min por token (evita escrita por request)."""
    now = time.time()
    with _access_lock:
        invites = _load_access_invites()
        changed = False
        for inv in invites:
            if inv.get("token") == token:
                last = 0.0
                try:
                    last = float(inv.get("lastAccessAt") or 0)
                except (TypeError, ValueError):
                    last = 0.0
                if now - last > 60:
                    inv["lastAccessAt"] = now
                    inv["accessCount"] = int(inv.get("accessCount") or 0) + 1
                    changed = True
        if changed:
            _save_access_invites(invites)


def _public_invite_view(inv: dict, now: float | None = None) -> dict:
    return {
        "token": inv.get("token"),
        "label": inv.get("label"),
        "createdAt": inv.get("createdAt"),
        "expiresAt": inv.get("expiresAt"),
        "revoked": bool(inv.get("revoked")),
        "lastAccessAt": inv.get("lastAccessAt"),
        "accessCount": int(inv.get("accessCount") or 0),
        "status": _invite_status(inv, now),
    }


ACCESS_EXEMPT_PREFIXES = (
    "/entrar", "/api/access/", "/admin", "/api/admin/",
    "/favicon", "/robots.txt", "/healthz",
)


def _is_access_exempt(path: str) -> bool:
    return any(path == p or path.startswith(p) for p in ACCESS_EXEMPT_PREFIXES)


def _app_relative_path(request: Request) -> tuple[str, str]:
    """(caminho relativo ao app, prefixo de montagem). Sob Passenger o app vive em
    /aisstream: request.url.path inclui o prefixo (root_path) e os redirects
    precisam preservá-lo."""
    path = request.url.path
    root = (request.scope.get("root_path") or "").rstrip("/")
    if root and path.startswith(root):
        path = path[len(root):] or "/"
    return path, root
_kratos_events_lock = threading.Lock()
KRATOS_EVENTS_MAX = 5000  # eventos mantidos em disco (rolling)


def kratos_events_file() -> Path:
    return user_data_dir(DASHBOARD_USER_ID) / "kratos_events.jsonl"


def log_kratos_event(event_type: str, data: dict | None = None):
    """Registra um evento de uso/auditoria do KRATOS em JSONL (append, rolling).

    Tipos: chat, voice_session, feedback, error, no_data, report, insights, tour.
    """
    try:
        rec = {"at": get_now_iso(), "ts": int(time.time()), "type": event_type}
        if isinstance(data, dict):
            rec.update(data)
        path = kratos_events_file()
        with _kratos_events_lock:
            with path.open("a", encoding="utf-8") as fh:
                fh.write(json.dumps(rec, ensure_ascii=False) + "\n")
            # Compactação ocasional para não crescer sem limite.
            if rec["ts"] % 50 == 0:
                _trim_kratos_events(path)
    except OSError:
        pass


def _trim_kratos_events(path: Path):
    try:
        lines = path.read_text(encoding="utf-8").splitlines()
        if len(lines) > KRATOS_EVENTS_MAX:
            path.write_text("\n".join(lines[-KRATOS_EVENTS_MAX:]) + "\n", encoding="utf-8")
    except OSError:
        pass


def load_kratos_events(limit: int = 2000) -> list:
    try:
        path = kratos_events_file()
        if not path.exists():
            return []
        lines = path.read_text(encoding="utf-8").splitlines()
        out = []
        for ln in lines[-limit:]:
            ln = ln.strip()
            if not ln:
                continue
            try:
                out.append(json.loads(ln))
            except json.JSONDecodeError:
                continue
        return out
    except OSError:
        return []


def _saa_dim_defaults() -> dict:
    """Dimensões / programação; vazios na fila «Pedra» sem tabela."""
    return {"pob": "", "cal": "", "loa": "", "dwt": "", "m": "", "boca": "", "gt": ""}


def _saa_dims_from_payload(payload: dict) -> dict:
    d = _saa_dim_defaults()
    for k in d:
        v = payload.get(k)
        if v is not None and str(v).strip():
            d[k] = str(v).strip()
    return d


def _haversine_nm(lat1, lon1, lat2, lon2):
    """Distância ortodrómica aproximada (esfera) em milhas náuticas (MN)."""
    r_nm = 3440.065  # raio médio da Terra em MN
    p1, p2 = math.radians(lat1), math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlmb = math.radians(lon2 - lon1)
    h = math.sin(dphi / 2) ** 2 + math.cos(p1) * math.cos(p2) * math.sin(dlmb / 2) ** 2
    c = 2 * math.atan2(math.sqrt(h), math.sqrt(max(0.0, 1.0 - h)))
    return r_nm * c


def _parse_vessel_timestamp_unix(vessel):
    raw = vessel.get("timestamp")
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    if s.endswith("Z"):
        s = s[:-1] + "+00:00"
    try:
        return datetime.fromisoformat(s).timestamp()
    except ValueError:
        return None


def _accumulate_saam_exit(mmsi, geofence_id, geofence_name, duration_sec, vessel):
    bym = tug_stats_state.setdefault("byMmsi", {})
    ent = bym.setdefault(
        mmsi,
        {
            "abbr": SAAM_MMSI_ABBR.get(mmsi, mmsi[-3:]),
            "name": vessel.get("shipName") or "",
            "totalSeconds": 0.0,
            "totalManeuvers": 0,
            "totalNauticalMiles": 0.0,
            "byGeofence": {},
        },
    )
    ent["name"] = vessel.get("shipName") or ent.get("name")
    gf = ent["byGeofence"].setdefault(
        geofence_id,
        {"name": geofence_name, "seconds": 0.0, "maneuvers": 0},
    )
    gf["name"] = geofence_name
    gf["seconds"] = float(gf.get("seconds", 0)) + duration_sec
    gf["maneuvers"] = int(gf.get("maneuvers", 0)) + 1
    ent["totalSeconds"] = float(ent.get("totalSeconds", 0)) + duration_sec
    ent["totalManeuvers"] = int(ent.get("totalManeuvers", 0)) + 1


def _bump_saam_nautical_miles(mmsi, vessel, delta_nm):
    if delta_nm <= 0:
        return False
    bym = tug_stats_state.setdefault("byMmsi", {})
    ent = bym.setdefault(
        mmsi,
        {
            "abbr": SAAM_MMSI_ABBR.get(mmsi, mmsi[-3:]),
            "name": vessel.get("shipName") or "",
            "totalSeconds": 0.0,
            "totalManeuvers": 0,
            "totalNauticalMiles": 0.0,
            "byGeofence": {},
        },
    )
    ent["name"] = vessel.get("shipName") or ent.get("name")
    ent["totalNauticalMiles"] = float(ent.get("totalNauticalMiles", 0)) + float(delta_nm)
    return True


def _maybe_persist_nm_stats():
    global _nm_save_budget
    _nm_save_budget["n"] = _nm_save_budget.get("n", 0) + 1
    now = time.monotonic()
    last = float(_nm_save_budget.get("last_mono", 0.0))
    if _nm_save_budget["n"] >= 35 or (now - last) >= 150.0:
        _nm_save_budget["n"] = 0
        _nm_save_budget["last_mono"] = now
        save_tug_stats()


def update_saam_nautical_miles(vessel):
    """Soma milhas náuticas percorridas entre posições AIS consecutivas (rebocadores SAAM-BGRA)."""
    if not vessel.get("isSaamBgra"):
        return
    mmsi = str(vessel.get("mmsi"))
    lat = vessel.get("latitude")
    lon = vessel.get("longitude")
    if not isinstance(lat, (int, float)) or not isinstance(lon, (int, float)):
        return
    lat = float(lat)
    lon = float(lon)
    t_cur = _parse_vessel_timestamp_unix(vessel)
    try:
        sog = float(vessel.get("sog") or 0)
    except (TypeError, ValueError):
        sog = 0.0

    did_nm = False
    mono_now = time.monotonic()
    with saam_geofence_stats_lock:
        prev = saam_last_position_for_nm.get(mmsi)
        saam_last_position_for_nm[mmsi] = {
            "lat": lat,
            "lon": lon,
            "t_ais": t_cur,
            "sog": sog,
            "mono": mono_now,
        }
        if not prev:
            return
        lat0, lon0 = float(prev["lat"]), float(prev["lon"])
        t0 = prev.get("t_ais")
        if t_cur is not None and t0 is not None:
            dt_s = t_cur - t0
        else:
            pm = prev.get("mono")
            dt_s = (mono_now - float(pm)) if pm is not None else 0.0
        if dt_s <= 0 or dt_s > 6 * 3600:
            return
        nm = _haversine_nm(lat0, lon0, lat, lon)
        sog_ref = max(sog, float(prev.get("sog") or 0), 2.0)
        max_nm = max(0.08, min(100.0, (dt_s / 3600.0) * max(28.0, sog_ref * 1.6)))
        if nm > max_nm:
            return
        if nm < 1e-6:
            return
        did_nm = _bump_saam_nautical_miles(mmsi, vessel, nm)

    if did_nm:
        _maybe_persist_nm_stats()


def update_saam_fleet_geofence_stats(vessel):
    if not vessel.get("isSaamBgra"):
        return
    mmsi = str(vessel.get("mmsi"))
    now = time.time()
    inside_ids = set()
    id_to_name = {}
    valid_maneuver_gids = set()
    with geofence_lock:
        maneuver_geofences = [
            g
            for g in geofences
            if g.get("isActive", True) and g.get("type") in SAAM_MANEUVER_STATS_GEOFENCE_TYPES
        ]
        for g in maneuver_geofences:
            gid = g.get("id")
            if gid:
                valid_maneuver_gids.add(gid)
        for g in maneuver_geofences:
            if not geofence_matches_vessel_scope(vessel, g):
                continue
            gid = g.get("id")
            if not gid:
                continue
            id_to_name[gid] = g.get("name", gid)
            if is_inside_geofence(vessel, g):
                inside_ids.add(gid)
    had_exit = False
    with saam_geofence_stats_lock:
        for key in list(saam_geofence_sessions.keys()):
            if key[0] == mmsi and key[1] not in valid_maneuver_gids:
                saam_geofence_sessions.pop(key, None)
        prev_stored = last_saam_inside_geofences.get(mmsi, set())
        prev = {gid for gid in prev_stored if gid in valid_maneuver_gids}
        for gid in prev - inside_ids:
            had_exit = True
            sess = saam_geofence_sessions.pop((mmsi, gid), None)
            if sess:
                dt = max(0.0, now - sess["t0"])
                _accumulate_saam_exit(mmsi, gid, id_to_name.get(gid, str(gid)), dt, vessel)
        for gid in inside_ids - prev:
            saam_geofence_sessions[(mmsi, gid)] = {"t0": now}
        last_saam_inside_geofences[mmsi] = set(inside_ids)
    if had_exit:
        save_tug_stats()


def update_saam_operating_hours(vessel):
    """Acumula horas de OPERAÇÃO (movimento) do rebocador SAAM para controle de
    fadiga. Usa o delta de relógio global por MMSI (robusto a múltiplos streams).
    Repouso (parado) >= TUG_RECOVERY_HOURS zera a fadiga (revezamento)."""
    if not vessel.get("isSaamBgra"):
        return
    global _last_op_save_ts
    mmsi = str(vessel.get("mmsi"))
    try:
        sog = float(vessel.get("sog") or 0)
    except (TypeError, ValueError):
        sog = 0.0
    now = time.time()
    with saam_geofence_stats_lock:
        op = tug_stats_state.setdefault("operating", {})
        ent = op.get(mmsi)
        if not ent:
            op[mmsi] = {
                "lastTs": now, "operatingSec": 0.0, "restingSec": 0.0,
                "recovered": False, "moving": sog >= TUG_OP_MOVING_SOG,
                "name": vessel.get("shipName") or mmsi, "byDay": {},
            }
            return
        dt = now - float(ent.get("lastTs") or now)
        ent["lastTs"] = now
        if dt < 0:
            dt = 0.0
        if dt > TUG_OP_MAX_GAP_SEC:
            dt = 0.0  # sinal perdido: não contabiliza a lacuna
        moving = sog >= TUG_OP_MOVING_SOG
        ent["moving"] = moving
        crossed = None
        if moving:
            prev_h = float(ent.get("operatingSec", 0)) / 3600.0
            ent["operatingSec"] = float(ent.get("operatingSec", 0)) + dt
            ent["restingSec"] = 0.0
            ent["recovered"] = False
            new_h = ent["operatingSec"] / 3600.0
            if prev_h < TUG_OP_WARN_HOURS <= new_h:
                crossed = "warn"
            if prev_h < TUG_OP_LIMIT_HOURS <= new_h:
                crossed = "limit"
            day = datetime.fromtimestamp(now).strftime("%Y-%m-%d")
            bd = ent.setdefault("byDay", {})
            bd[day] = float(bd.get(day, 0)) + dt
            if len(bd) > 40:
                for k in sorted(bd)[:-30]:
                    bd.pop(k, None)
        else:
            ent["restingSec"] = float(ent.get("restingSec", 0)) + dt
            if ent["restingSec"] >= TUG_RECOVERY_HOURS * 3600 and float(ent.get("operatingSec", 0)) > 0:
                ent["operatingSec"] = 0.0
                ent["recovered"] = True
        ent["name"] = vessel.get("shipName") or ent.get("name") or mmsi
    # Persiste ao cruzar marco ou a cada ~2 min.
    if crossed or (now - _last_op_save_ts > 120):
        _last_op_save_ts = now
        save_tug_stats()


def point_in_polygon(latitude, longitude, polygon):
    # polygon: [[lat, lon], [lat, lon], ...]
    if not polygon or len(polygon) < 3:
        return False
    inside = False
    j = len(polygon) - 1
    for i in range(len(polygon)):
        yi, xi = polygon[i][0], polygon[i][1]
        yj, xj = polygon[j][0], polygon[j][1]
        intersects = ((yi > latitude) != (yj > latitude)) and (
            longitude < (xj - xi) * (latitude - yi) / ((yj - yi) or 1e-12) + xi
        )
        if intersects:
            inside = not inside
        j = i
    return inside


def is_inside_geofence(vessel, geofence):
    geometry = geofence.get("geometry") or {}
    gtype = geofence.get("type")
    latitude = vessel.get("latitude")
    longitude = vessel.get("longitude")
    if latitude is None or longitude is None:
        return False
    if gtype in {"berco", "base_rebocador", "polygon"}:
        polygon = geometry.get("coordinates") or []
        return point_in_polygon(latitude, longitude, polygon)
    if gtype == "circle":
        center = geometry.get("center") or []
        radius_m = geometry.get("radiusMeters") or 0
        if len(center) != 2 or radius_m <= 0:
            return False
        # aprox. metros por grau
        dy = (latitude - center[0]) * 111_320
        dx = (longitude - center[1]) * 111_320
        return (dx * dx + dy * dy) ** 0.5 <= radius_m
    return False


def geofence_matches_vessel_scope(vessel, geofence):
    scope = geofence.get("fleetScope", "all")
    if scope == "all":
        return True
    if scope == SAAM_BGRA_FLEET_NAME:
        return bool(vessel.get("isSaamBgra"))
    return False


def get_vessel_geofences(vessel):
    ensure_geofences_loaded()
    names = []
    with geofence_lock:
        active_geofences = [g for g in geofences if g.get("isActive", True)]
    for geofence in active_geofences:
        if not geofence_matches_vessel_scope(vessel, geofence):
            continue
        if is_inside_geofence(vessel, geofence):
            names.append(geofence.get("name", "Sem nome"))
    return names


def vessel_in_rebocador_base(vessel):
    """Dentro de algum geofence ativo tipo base_rebocador (escopo respeitado)."""
    ensure_geofences_loaded()
    with geofence_lock:
        active_geofences = [g for g in geofences if g.get("isActive", True)]
    for geofence in active_geofences:
        if geofence.get("type") != "base_rebocador":
            continue
        if not geofence_matches_vessel_scope(vessel, geofence):
            continue
        if is_inside_geofence(vessel, geofence):
            return True
    return False


def build_geofence_occupancy():
    ensure_geofences_loaded()
    with geofence_lock:
        active_geofences = [g for g in geofences if g.get("isActive", True)]
    vessels = list(latest_vessel_by_mmsi.values())
    occupancy = []
    for geofence in active_geofences:
        inside = []
        for vessel in vessels:
            if not geofence_matches_vessel_scope(vessel, geofence):
                continue
            if is_inside_geofence(vessel, geofence):
                inside.append(
                    {
                        "mmsi": vessel.get("mmsi"),
                        "shipName": vessel.get("shipName"),
                        "shipCategory": vessel.get("shipCategory"),
                        "fleet": vessel.get("fleet"),
                        "timestamp": vessel.get("timestamp"),
                        "heading": vessel.get("heading"),
                        "course": vessel.get("course"),
                        "isSaamBgra": bool(vessel.get("isSaamBgra")),
                        "latitude": vessel.get("latitude"),
                        "longitude": vessel.get("longitude"),
                    }
                )
        occupancy.append(
            {
                "geofenceId": geofence.get("id"),
                "name": geofence.get("name"),
                "type": geofence.get("type"),
                "fleetScope": geofence.get("fleetScope", "all"),
                "vesselCount": len(inside),
                "insideVessels": inside,
            }
        )
    return occupancy


def build_live_subscription():
    return {
        "APIKey": AISSTREAM_API_KEY,
        "BoundingBoxes": AREAS[current_area_key]["boundingBoxes"],
        "FilterMessageTypes": [
            "PositionReport",
            "StandardClassBPositionReport",
            "ExtendedClassBPositionReport",
            "ShipStaticData",
            "StaticDataReport",
        ],
    }


_vessels_snapshot_loaded = False


def ensure_vessels_snapshot_loaded():
    global _vessels_snapshot_loaded
    if _vessels_snapshot_loaded:
        return
    _vessels_snapshot_loaded = True
    load_vessels_snapshot()


def ensure_live_worker_started():
    global live_worker_thread
    ensure_vessels_snapshot_loaded()
    ensure_fleet_config_loaded()
    if current_mode != "live" or not AISSTREAM_API_KEY:
        return
    with live_worker_lock:
        if live_worker_thread and live_worker_thread.is_alive():
            return
        live_worker_thread = threading.Thread(
            target=lambda: asyncio.run(live_background_worker()),
            daemon=True,
            name="ais-live-worker",
        )
        live_worker_thread.start()


def normalize_ship_type_code(metadata, message_body):
    report_b = message_body.get("ReportB") if isinstance(message_body, dict) else None
    raw_type = (
        metadata.get("ShipType")
        or metadata.get("ship_type")
        or (report_b.get("ShipType") if isinstance(report_b, dict) else None)
        or message_body.get("TypeAndCargo")
        or message_body.get("ShipType")
        or message_body.get("Type")
    )
    try:
        if raw_type is None or raw_type == "":
            return None
        return int(raw_type)
    except Exception:
        return None


def infer_ship_category(ship_type_code, ship_name):
    if ship_type_code is not None:
        if 30 <= ship_type_code <= 39:
            if ship_type_code in [36, 37]:
                return "lazer"
            return "pesca"
        if 50 <= ship_type_code <= 59:
            return "rebocador_servico"
        if 60 <= ship_type_code <= 69:
            return "passageiros"
        if 70 <= ship_type_code <= 79:
            return "carga"
        if 80 <= ship_type_code <= 89:
            return "petroleiro"

    name = (ship_name or "").upper()
    if (
        "TUG" in name
        or "REBOC" in name
        or "AHTS" in name
        or "PSV" in name
        or "OSRV" in name
        or "SUPPLY" in name
        or "SVITZER" in name
        or "SAAM" in name
        or "CBO" in name
        or "HOS " in name
    ):
        return "rebocador_servico"
    if "TANK" in name or "PETRO" in name:
        return "petroleiro"
    if "CARGO" in name or "BULK" in name or "CONTAINER" in name or "BARGE" in name:
        return "carga"
    if "FISH" in name or "PESCA" in name:
        return "pesca"
    if "PASSENGER" in name or "FERRY" in name or "CATAMARAN" in name or "CAT " in name:
        return "passageiros"
    return "outros"


def _dim_segment(value):
    try:
        if value is None:
            return None
        v = int(value)
        if v < 0 or v > 511:
            return None
        return v
    except Exception:
        return None


def extract_ship_dimensions_meters(message_type, message_body):
    """
    Retorna (length_m, beam_m) quando AIS traz dimensões em metros (A/B/C/D).
    LOA ~= A + B, Boca ~= C + D.
    """
    dim = None
    if isinstance(message_body, dict):
        dim = message_body.get("Dimension") or message_body.get("dimension")
    if not isinstance(dim, dict):
        return None, None
    a = _dim_segment(dim.get("A"))
    b = _dim_segment(dim.get("B"))
    c = _dim_segment(dim.get("C"))
    d = _dim_segment(dim.get("D"))
    if a is None or b is None:
        return None, None
    length_m = float(a + b)
    beam_m = None
    if c is not None and d is not None:
        beam_m = float(c + d)
    if length_m <= 0 or length_m > 600:
        return None, beam_m
    if beam_m is not None and (beam_m <= 0 or beam_m > 120):
        beam_m = None
    return length_m, beam_m


def extract_ship_ref_offsets(message_type, message_body):
    """
    Offsets do ponto de referência AIS (antena) ao casco, em metros: A/B/C/D.
    A = proa, B = popa, C = bombordo, D = boreste. LOA = A+B, Boca = C+D.
    Permite desenhar o casco em escala real ancorado na antena (estilo MarineTraffic).
    """
    dim = None
    if isinstance(message_body, dict):
        dim = message_body.get("Dimension") or message_body.get("dimension")
    if not isinstance(dim, dict):
        return None
    a = _dim_segment(dim.get("A"))
    b = _dim_segment(dim.get("B"))
    c = _dim_segment(dim.get("C"))
    d = _dim_segment(dim.get("D"))
    if a is None or b is None or c is None or d is None:
        return None
    loa = a + b
    beam = c + d
    if loa <= 0 or loa > 600 or beam <= 0 or beam > 120:
        return None
    return [float(a), float(b), float(c), float(d)]


def estimate_length_from_category(ship_category):
    """Fallback visual quando não há dimensão AIS (metros aproximados)."""
    if ship_category == "rebocador_servico":
        return 32.0
    if ship_category == "pesca":
        return 28.0
    if ship_category == "lazer":
        return 18.0
    if ship_category == "passageiros":
        return 140.0
    if ship_category == "carga":
        return 220.0
    if ship_category == "petroleiro":
        return 250.0
    return 90.0


def push_recent_vessel(vessel_payload):
    global last_vessel_seq
    last_vessel_seq += 1
    item = dict(vessel_payload)
    item["_seq"] = last_vessel_seq
    recent_vessels.append(item)


def classify_geofence(latitude, longitude):
    if latitude is None or longitude is None:
        return None
    if (
        GUANABARA_GEOFENCE["minLat"] <= latitude <= GUANABARA_GEOFENCE["maxLat"]
        and GUANABARA_GEOFENCE["minLon"] <= longitude <= GUANABARA_GEOFENCE["maxLon"]
    ):
        return GUANABARA_GEOFENCE["name"]
    return None


# --- REST endpoints ---
@app.get("/api/status")
def get_status(expand: str | None = None):
    ensure_live_worker_started()
    out = {
        "app": "KRATOS - Inteligência Naval Estratégica",
        "version": "1.0.0",
        "mode": current_mode,
        "area": AREAS.get(current_area_key, {}),
        "areaKey": current_area_key,
        "liveConnected": live_connected,
        "lastError": last_error,
        "lastAisMessageAt": last_ais_message_at,
        "totalMessages": total_messages,
    }
    if expand == "dashboard":
        out["dashboard"] = build_dashboard_overview_dict()
    return out


async def _sync_saa_from_praticagem_impl():
    """
    Substitui apenas entradas com source=praticagem por uma leitura atual do site público.
    Entradas manuais (sem source ou source diferente) mantêm-se.
    """
    url = (os.getenv("PRATICAGEM_RJ_URL") or praticagem_saa.DEFAULT_PRATICAGEM_URL).strip()
    try:
        rows, parse_diag = await asyncio.to_thread(
            praticagem_saa.fetch_parse_praticagem_diagnostics, url
        )
    except Exception as e:
        return JSONResponse(
            {"ok": False, "error": f"Falha ao aceder Praticagem: {e}"},
            status_code=502,
        )
    if not rows:
        return {
            "ok": True,
            "imported": 0,
            "warning": (
                "Nenhuma linha extraída. Verifique PRATICAGEM_RJ_URL, firewall ou bloqueio do site ao servidor; "
                "veja parseDiagnostics."
            ),
            "parseDiagnostics": parse_diag,
        }
    now = get_now_iso()
    new_items = []
    for r in rows:
        dims = _saa_dim_defaults()
        for k in dims:
            if r.get(k):
                dims[k] = str(r[k]).strip()
        if r.get("rowSource") == "programacao":
            new_items.append(
                {
                    "id": str(uuid.uuid4()),
                    "vesselName": r.get("vesselName") or "—",
                    "berthName": (r.get("berthName") or "").strip() or "—",
                    "empRb": (r.get("empRb") or "").strip() or "—",
                    "status": (r.get("status") or "").strip() or "—",
                    "note": "Fonte: praticagem-rj.com.br (programação pública)",
                    "source": "praticagem",
                    "recordedAt": now,
                    **dims,
                }
            )
        else:
            new_items.append(
                {
                    "id": str(uuid.uuid4()),
                    "vesselName": r.get("vesselName") or "—",
                    "berthName": f"Pedra / {r.get('zone', '—')}",
                    "empRb": (r.get("empRb") or "SAA").strip() or "SAA",
                    "status": (r.get("status") or "").strip() or "—",
                    "note": "Fonte: praticagem-rj.com.br (fila na Pedra)",
                    "source": "praticagem",
                    "recordedAt": now,
                    **dims,
                }
            )
    new_items = _dedupe_saa_maneuvers(new_items)
    with saa_maneuvers_lock:
        kept = [x for x in saa_maneuvers_list if x.get("source") != "praticagem"]
        saa_maneuvers_list[:] = new_items + kept
        save_saa_maneuvers()
    zones = sorted({(r.get("zone") or "").strip() for r in rows if (r.get("zone") or "").strip()})
    return {
        "ok": True,
        "imported": len(new_items),
        "zones": zones,
        "parseDiagnostics": parse_diag,
    }


@app.post("/api/status/sync-praticagem-saa")
async def sync_saa_maneuvers_from_praticagem_status_path():
    return await _sync_saa_from_praticagem_impl()


@app.post("/api/praticagem/saa-sync")
async def sync_saa_maneuvers_from_praticagem_short_path():
    return await _sync_saa_from_praticagem_impl()


@app.post("/api/dashboard/saa-maneuvers/sync-praticagem")
async def sync_saa_maneuvers_from_praticagem_dashboard_api():
    return await _sync_saa_from_praticagem_impl()


@app.post("/dashboard/api/saa-maneuvers/sync-praticagem")
async def sync_saa_maneuvers_from_praticagem_under_dashboard():
    return await _sync_saa_from_praticagem_impl()


# --- Integração Obsidian (Supabase Storage) — Sprint 1: Exportador Base -------

def _obsidian_auto_enabled() -> bool:
    return (os.getenv("OBSIDIAN_AUTO_SYNC", "0") or "0").strip().lower() in {
        "1", "true", "on", "yes", "sim",
    }


def _obsidian_auto_interval() -> int:
    try:
        return int(os.getenv("OBSIDIAN_AUTO_SYNC_SECONDS", "300") or "300")
    except (TypeError, ValueError):
        return 300


@app.get("/api/obsidian/status")
@app.get("/dashboard/api/obsidian/status")
def obsidian_status():
    """Diagnóstico da ponte KRATOS -> Supabase Storage (sem expor segredos)."""
    status = obsidian_supabase.config_status()
    status["autoSync"] = _obsidian_auto_enabled()
    status["autoSyncSeconds"] = _obsidian_auto_interval()
    status["lastExportTs"] = _obsidian_last_export_ts or None
    return status


@app.post("/api/obsidian/test-upload")
@app.post("/dashboard/api/obsidian/test-upload")
async def obsidian_test_upload():
    """Sobe uma nota de saúde no bucket para validar a conexão (Sprint 1)."""
    if not obsidian_supabase.is_configured():
        return JSONResponse(
            {
                "ok": False,
                "error": (
                    "Supabase não configurado. Defina SUPABASE_URL, SUPABASE_KEY "
                    "e SUPABASE_BUCKET no .env."
                ),
                "status": obsidian_supabase.config_status(),
            },
            status_code=400,
        )
    try:
        result = await obsidian_supabase.check_connection_async()
    except obsidian_supabase.ObsidianSupabaseError as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=502)
    return result


def _build_obsidian_vault_notes():
    """Reúne o estado do KRATOS e gera as notas interligadas (Sprint 2)."""
    overview = build_dashboard_overview_dict()
    vessels = list(latest_vessel_by_mmsi.values())
    metocean = _fetch_metocean_context()
    saam_fleet = [{"mmsi": m, "name": n} for m, n in SAAM_BGRA_NAMES.items()]
    return obsidian_notes.build_vault(
        overview,
        vessels,
        metocean,
        saam_fleet=saam_fleet,
        competitor_tugs=COMPETITOR_TUGS,
        now_iso=get_now_iso(),
    )


@app.post("/api/obsidian/export")
@app.post("/dashboard/api/obsidian/export")
async def obsidian_export(wait: bool = False):
    """Dispara o envio do vault de notas para o Supabase.

    Por padrão roda em **segundo plano** e responde na hora — o vault completo
    pode levar mais que o timeout do gateway (Cloudflare ~20 s). Use ``?wait=1``
    para o modo síncrono (diagnóstico), que aguarda e retorna o resumo.
    """
    if not obsidian_supabase.is_configured():
        return JSONResponse(
            {
                "ok": False,
                "error": (
                    "Supabase não configurado. Defina SUPABASE_URL, SUPABASE_KEY "
                    "e SUPABASE_BUCKET no .env."
                ),
                "status": obsidian_supabase.config_status(),
            },
            status_code=400,
        )
    if wait:
        result = await _run_obsidian_export_safe("manual")
        return JSONResponse(result, status_code=200 if result.get("ok") else 502)
    if _obsidian_export_running:
        return {
            "ok": True,
            "status": "running",
            "message": "Sincronização já em andamento.",
        }
    asyncio.create_task(_run_obsidian_export_safe("manual"))
    return {
        "ok": True,
        "status": "started",
        "message": "Sincronização iniciada em segundo plano. As notas aparecem no bucket em instantes.",
    }


@app.get("/api/obsidian/graph")
@app.get("/dashboard/api/obsidian/graph")
def obsidian_graph():
    """Grafo (nós/arestas) do modelo KRATOS para a visualização web."""
    overview = build_dashboard_overview_dict()
    vessels = list(latest_vessel_by_mmsi.values())
    saam_fleet = [{"mmsi": m, "name": n} for m, n in SAAM_BGRA_NAMES.items()]
    graph = obsidian_notes.build_graph(
        overview, vessels, saam_fleet=saam_fleet, competitor_tugs=COMPETITOR_TUGS
    )
    graph["ok"] = True
    return graph


@app.get("/graph")
@app.get("/graph/")
def graph_page():
    """Página de visualização do grafo KRATOS (Graph View nativo)."""
    return FileResponse(FRONTEND_DIR / "graph.html", media_type="text/html")


@app.get("/versao")
@app.get("/versao/")
def versao_page():
    """Painel de implementações da versão (responsivo)."""
    return FileResponse(FRONTEND_DIR / "versao.html", media_type="text/html")


async def _run_obsidian_export_safe(reason: str) -> dict:
    """Gera e sobe o vault em background, tolerante a falha (auto-sync)."""
    global _obsidian_export_running, _obsidian_last_export_ts
    if _obsidian_export_running:
        return {"ok": False, "skipped": "already_running"}
    _obsidian_export_running = True
    try:
        notes = await asyncio.to_thread(_build_obsidian_vault_notes)
        result = await obsidian_supabase.upload_notes_async(notes)
        result["generated"] = len(notes)
        result["reason"] = reason
        if result.get("uploaded"):
            _obsidian_last_export_ts = time.time()
        return result
    except Exception as e:
        return {"ok": False, "error": str(e), "reason": reason}
    finally:
        _obsidian_export_running = False


async def _obsidian_auto_export_if_due(reason: str):
    """Dispara o export automático respeitando debounce e configuração."""
    if not (_obsidian_auto_enabled() and obsidian_supabase.is_configured()):
        return
    interval = _obsidian_auto_interval()
    if interval > 0 and (time.time() - _obsidian_last_export_ts) < interval:
        return  # debounce: ainda dentro da janela mínima
    await _run_obsidian_export_safe(reason)


async def _obsidian_auto_sync_loop():
    """Loop periódico de sincronização do Obsidian (debounce natural pelo intervalo)."""
    interval = _obsidian_auto_interval()
    if interval <= 0:
        return
    await asyncio.sleep(20)
    while True:
        try:
            await _obsidian_auto_export_if_due("loop")
        except asyncio.CancelledError:
            break
        except Exception:
            pass
        try:
            await asyncio.sleep(interval)
        except asyncio.CancelledError:
            break


@app.get("/healthz")
def healthz():
    return {"ok": True}

@app.get("/api/areas")
def get_areas():
    ensure_live_worker_started()
    return {"areas": AREAS, "currentAreaKey": current_area_key}

@app.get("/api/vessels")
def get_vessels(since: int = 0, limit: int = 300, snapshot: bool = False):
    ensure_live_worker_started()
    if snapshot:
        items = []
        for v in latest_vessel_by_mmsi.values():
            if not isinstance(v, dict):
                continue
            items.append({k: val for k, val in v.items() if k != "raw"})
        return {
            "vessels": items,
            "lastSeq": last_vessel_seq,
            "count": len(items),
        }
    items = [v for v in recent_vessels if v.get("_seq", 0) > since]
    if limit > 0:
        items = items[-limit:]
    current_seq = recent_vessels[-1]["_seq"] if recent_vessels else since
    return {
        "vessels": items,
        "lastSeq": current_seq,
        "count": len(items)
    }


def _saam_operating_hours_rows():
    op = (tug_stats_state or {}).get("operating", {}) or {}
    rows = []
    for mmsi, ent in op.items():
        op_h = float(ent.get("operatingSec", 0)) / 3600.0
        rest_h = float(ent.get("restingSec", 0)) / 3600.0
        if op_h >= TUG_OP_LIMIT_HOURS:
            status = "limite"
        elif op_h >= TUG_OP_WARN_HOURS:
            status = "atencao"
        elif ent.get("recovered"):
            status = "recuperado"
        else:
            status = "ok"
        rows.append({
            "mmsi": mmsi,
            "name": ent.get("name") or mmsi,
            "operatingHours": round(op_h, 2),
            "restingHours": round(rest_h, 2),
            "moving": bool(ent.get("moving")),
            "recovered": bool(ent.get("recovered")),
            "status": status,
        })
    rows.sort(key=lambda r: r["operatingHours"], reverse=True)
    return rows


@app.get("/api/saam-operating-hours")
def get_saam_operating_hours():
    """Horas de operação (fadiga) por rebocador SAAM, para o painel e os alertas."""
    return {
        "ok": True,
        "limitHours": TUG_OP_LIMIT_HOURS,
        "warnHours": TUG_OP_WARN_HOURS,
        "recoveryHours": TUG_RECOVERY_HOURS,
        "tugs": _saam_operating_hours_rows(),
    }


@app.get("/api/saa-maneuver-names")
def get_saa_maneuver_names():
    """Nomes dos navios programados para manobra pela SAAM (EMP.RB = SAA).

    Usado no mapa para destacar o label desses navios quando o usuário ativa
    a visualização de nomes. Retorna os nomes crus e normalizados (sem
    acento/maiúsculas) para casar com o shipName do AIS no frontend.
    """
    ensure_saa_maneuvers_loaded()
    with saa_maneuvers_lock:
        snapshot = _dedupe_saa_maneuvers(list(saa_maneuvers_list))
    names = []
    seen = set()
    for item in snapshot:
        emp = str(item.get("empRb") or "").strip().upper()
        if emp != OWN_COMPANY_EMP_RB:
            continue
        raw_name = str(item.get("vesselName") or "").strip()
        norm = _normalize_text(raw_name)
        if not norm or norm in {"—", "-"} or norm in seen:
            continue
        seen.add(norm)
        names.append({"name": raw_name, "normalized": norm, "pob": item.get("pob") or ""})
    return {"ok": True, "company": OWN_COMPANY_EMP_RB, "count": len(names), "names": names}


@app.get("/api/programmed-names")
def get_programmed_vessel_names():
    """Nomes (normalizados) de TODOS os navios com manobra programada na
    Praticagem-RJ, de QUALQUER empresa (SAA/WIL/CAM/etc.).

    Usado pelo 'radar de oportunidade': um navio que entra na barra e NÃO
    consta aqui provavelmente nao tem contrato de reboque (alvo comercial).
    """
    ensure_saa_maneuvers_loaded()
    with saa_maneuvers_lock:
        snapshot = _dedupe_saa_maneuvers(list(saa_maneuvers_list))
    names = []
    seen = set()
    for item in snapshot:
        raw_name = str(item.get("vesselName") or "").strip()
        norm = _normalize_text(raw_name)
        if not norm or norm in {"—", "-"} or norm in seen:
            continue
        seen.add(norm)
        names.append({
            "normalized": norm,
            "empRb": str(item.get("empRb") or "").strip().upper(),
            "pob": item.get("pob") or "",
        })
    return {"ok": True, "count": len(names), "names": names}


def _dashboard_geofence_status_rows():
    ensure_geofences_loaded()
    occ = {o.get("geofenceId"): o for o in build_geofence_occupancy()}
    with geofence_lock:
        snapshot = list(geofences)
    vessels_full = list(latest_vessel_by_mmsi.values())
    # Linhas «base_rebocador»: SAA acende também se houver navio SAA em qualquer berço **ativo**
    # (geometria só; escopo «SAAM-BGRA» não esconde navio comercial do indicador SAA)
    saa_in_any_berco = False
    for g in snapshot:
        if g.get("type") != "berco" or not g.get("isActive", True):
            continue
        inside_geo_b = [v for v in vessels_full if is_inside_geofence(v, g)]
        if any(_vessel_matches_saa_maneuver_lamp(v) for v in inside_geo_b):
            saa_in_any_berco = True
            break

    rows = []
    for g in snapshot:
        gid = g.get("id")
        o = occ.get(gid, {})
        inside = o.get("insideVessels", [])
        gtype = g.get("type")
        active = bool(g.get("isActive", True))
        if active:
            inside_geo = [v for v in vessels_full if is_inside_geofence(v, g)]
            lamp_saam = any(bool(v.get("isSaamBgra")) for v in inside_geo)
            local_saa = any(_vessel_matches_saa_maneuver_lamp(v) for v in inside_geo)
        else:
            lamp_saam = False
            local_saa = False
        if gtype in ("berco", "polygon"):
            lamp_saa = local_saa
        elif gtype == "base_rebocador":
            lamp_saa = local_saa or saa_in_any_berco
        else:
            lamp_saa = False

        gname = (g.get("name") or "").strip().lower()
        if gname == "base brasco":
            em_manobra = "stand-by"
        elif gtype == "berco" and active:
            em_manobra = _berco_em_manobra_text(inside_geo)
        else:
            em_manobra = "—"

        lamp_saam_names: list[str] = []
        lamp_saa_names: list[str] = []
        if active:
            lamp_saam_names = [
                _vessel_display_name_for_tooltip(v)
                for v in inside_geo
                if bool(v.get("isSaamBgra"))
            ]
            lamp_saam_names = list(dict.fromkeys(lamp_saam_names))[:18]

            if gtype in ("berco", "polygon"):
                lamp_saa_names = [
                    _vessel_display_name_for_tooltip(v)
                    for v in inside_geo
                    if _vessel_matches_saa_maneuver_lamp(v)
                ]
                lamp_saa_names = list(dict.fromkeys(lamp_saa_names))[:18]
            elif gtype == "base_rebocador" and lamp_saa:
                seen: set[str] = set()
                for v in inside_geo:
                    if not _vessel_matches_saa_maneuver_lamp(v):
                        continue
                    label = _vessel_display_name_for_tooltip(v)
                    if label not in seen:
                        seen.add(label)
                        lamp_saa_names.append(label)
                if saa_in_any_berco:
                    for gb in snapshot:
                        if gb.get("type") != "berco" or not gb.get("isActive", True):
                            continue
                        bname = (gb.get("name") or "Berço").strip() or "Berço"
                        for v in vessels_full:
                            if not is_inside_geofence(v, gb):
                                continue
                            if not _vessel_matches_saa_maneuver_lamp(v):
                                continue
                            label = f"{bname} — {_vessel_display_name_for_tooltip(v)}"
                            if label not in seen:
                                seen.add(label)
                                lamp_saa_names.append(label)
                                if len(lamp_saa_names) >= 24:
                                    break
                        if len(lamp_saa_names) >= 24:
                            break

        rows.append(
            {
                "id": gid,
                "name": g.get("name"),
                "type": g.get("type"),
                "fleetScope": g.get("fleetScope", "all"),
                "isActive": bool(g.get("isActive", True)),
                "vesselCount": o.get("vesselCount", 0),
                "insideVessels": inside,
                "lampSaaManeuver": lamp_saa,
                "lampSaamInside": lamp_saam,
                "lampSaaNames": lamp_saa_names,
                "lampSaamNames": lamp_saam_names,
                "emManobra": em_manobra,
            }
        )
    return rows


def _in_maneuver_geofence(vessel, geofence_snapshot):
    for g in geofence_snapshot:
        if not g.get("isActive", True):
            continue
        if g.get("type") not in ("berco", "polygon"):
            continue
        try:
            if is_inside_geofence(vessel, g):
                return g
        except Exception:
            continue
    return None


def _competitor_status_rows(geofence_snapshot):
    rows = []
    for company, tugs in COMPETITOR_TUGS.items():
        for tug in tugs:
            mmsi = tug.get("mmsi")
            vessel = latest_vessel_by_mmsi.get(mmsi) if mmsi else None
            if vessel:
                geo = _in_maneuver_geofence(vessel, geofence_snapshot)
                geofence_name = geo.get("name", "—") if geo else "—"
                manobra = bool(geo)
                ship_name = (vessel.get("shipName") or "").strip() or tug.get("name", "")
                lat = vessel.get("latitude")
                lon = vessel.get("longitude")
            else:
                geofence_name = "—"
                manobra = False
                ship_name = tug.get("name", "")
                lat = None
                lon = None
            rows.append(
                {
                    "company": company,
                    "mmsi": mmsi,
                    "name": ship_name,
                    "insideManeuverGeofence": manobra,
                    "maneuverGeofenceName": geofence_name,
                    "latitude": lat,
                    "longitude": lon,
                }
            )
    return rows


def _market_share_rows(saa_snapshot):
    counts = {}
    for item in saa_snapshot:
        key = str(item.get("empRb") or "N/A").strip().upper() or "N/A"
        counts[key] = counts.get(key, 0) + 1
    total = sum(counts.values())
    rows = []
    for key, cnt in sorted(counts.items(), key=lambda kv: kv[1], reverse=True):
        share = (100.0 * cnt / total) if total else 0.0
        rows.append({
            "empRb": key,
            "count": cnt,
            "sharePct": round(share, 2),
            # EMP.RB "SAA" na Praticagem representa a SAAM (nossa empresa).
            "isOwnCompany": key == OWN_COMPANY_EMP_RB,
        })
    return {"rows": rows, "total": total}


def _parse_recorded_at_iso(value):
    if not value:
        return None
    try:
        txt = str(value).strip()
        if txt.endswith("Z"):
            txt = txt[:-1] + "+00:00"
        return datetime.fromisoformat(txt)
    except Exception:
        return None


def _market_share_windows(saa_snapshot):
    now = datetime.now()
    midnight = now.replace(hour=0, minute=0, second=0, microsecond=0)
    w7 = now.timestamp() - 7 * 86400
    w30 = now.timestamp() - 30 * 86400

    today_items = []
    w7_items = []
    w30_items = []
    for item in saa_snapshot:
        dt = _parse_recorded_at_iso(item.get("recordedAt"))
        if dt is None:
            continue
        ts = dt.timestamp()
        if dt >= midnight:
            today_items.append(item)
        if ts >= w7:
            w7_items.append(item)
        if ts >= w30:
            w30_items.append(item)
    return {
        "today": _market_share_rows(today_items),
        "last7d": _market_share_rows(w7_items),
        "last30d": _market_share_rows(w30_items),
    }


def _estimate_tugs_required(maneuver):
    def _f(v):
        if v is None:
            return 0.0
        txt = str(v).replace(",", ".").strip()
        try:
            return float(txt)
        except Exception:
            return 0.0

    loa = _f(maneuver.get("loa"))
    dwt = _f(maneuver.get("dwt"))
    if loa >= 300 or dwt >= 100000:
        return 4
    if loa >= 230 or dwt >= 60000:
        return 3
    if loa > 0:
        return 2
    return 2


def _estimate_maneuver_value(maneuver):
    """Valor comercial *relativo* da manobra (qualitativo, sem tabela de preços).

    Heurística por porte do navio + rebocadores estimados: quanto maior o navio
    e mais rebocadores necessários, maior o valor comercial da manobra.
    """
    def _f(v):
        try:
            return float(str(v).replace(",", ".").strip())
        except Exception:
            return 0.0

    loa = _f(maneuver.get("loa"))
    dwt = _f(maneuver.get("dwt"))
    tugs = _estimate_tugs_required(maneuver)
    if loa >= 300 or dwt >= 100000 or tugs >= 4:
        return "alto"
    if loa >= 230 or dwt >= 60000 or tugs >= 3:
        return "medio"
    return "padrao"


def _parse_pob_to_datetime(pob):
    s = str(pob or "").strip()
    if not s:
        return None
    m = None
    # dd/mm HH:MM
    try:
        import re

        m = re.match(r"^(\d{1,2})/(\d{1,2})[^\d]?(\d{1,2}):(\d{2})", s)
        if not m:
            return None
        day, mon, hh, mm = int(m.group(1)), int(m.group(2)), int(m.group(3)), int(m.group(4))
        y = datetime.now().year
        dt = datetime(y, mon, day, hh, mm)
        # ajuste simples de virada de ano
        if dt.timestamp() < datetime.now().timestamp() - 14 * 86400:
            dt = datetime(y + 1, mon, day, hh, mm)
        return dt
    except Exception:
        return None


def _simultaneous_maneuvers_summary(saa_snapshot):
    buckets = {}
    for item in saa_snapshot:
        dt = _parse_pob_to_datetime(item.get("pob"))
        if not dt:
            continue
        key = dt.strftime("%Y-%m-%d %H:00")
        buckets.setdefault(key, []).append(item)
    rows = []
    for key, items in buckets.items():
        if len(items) < 2:
            continue
        tug_demand = sum(_estimate_tugs_required(x) for x in items)
        rows.append(
            {
                "timeSlot": key,
                "maneuvers": len(items),
                "estimatedTugsNeeded": tug_demand,
                "vessels": [x.get("vesselName", "—") for x in items[:10]],
            }
        )
    rows.sort(key=lambda x: x["timeSlot"])
    return rows[:24]


def _schedule_row_key_without_pob(item):
    return "|".join(
        [
            str(item.get("vesselName") or "").strip().upper(),
            str(item.get("berthName") or "").strip().upper(),
            str(item.get("empRb") or "").strip().upper(),
            str(item.get("m") or "").strip().upper(),
            str(item.get("loa") or "").strip(),
            str(item.get("dwt") or "").strip(),
            str(item.get("status") or "").strip().upper(),
        ]
    )


def _schedule_signature(item):
    return _schedule_row_key_without_pob(item) + "|" + str(item.get("pob") or "").strip()


def _load_schedule_monitor_state():
    path = schedule_monitor_file_for_user(DASHBOARD_USER_ID)
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _save_schedule_monitor_state(state):
    path = schedule_monitor_file_for_user(DASHBOARD_USER_ID)
    path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")


def _build_schedule_changes(saa_snapshot):
    prev = _load_schedule_monitor_state()
    prev_rows = prev.get("rows") or []
    prev_ts = prev.get("updatedAt")

    now = get_now_iso()
    new_rows = []
    for item in saa_snapshot:
        new_rows.append(
            {
                "keyNoPob": _schedule_row_key_without_pob(item),
                "signature": _schedule_signature(item),
                "pob": str(item.get("pob") or "").strip(),
                "vesselName": item.get("vesselName") or "—",
                "berthName": item.get("berthName") or "—",
                "empRb": item.get("empRb") or "N/A",
            }
        )

    prev_sigs = {str(r.get("signature") or "") for r in prev_rows}
    new_sigs = {str(r.get("signature") or "") for r in new_rows}

    added = [r for r in new_rows if r["signature"] not in prev_sigs][:60]
    removed = [r for r in prev_rows if str(r.get("signature") or "") not in new_sigs][:60]

    prev_by_key = {}
    for r in prev_rows:
        k = str(r.get("keyNoPob") or "")
        if not k:
            continue
        prev_by_key.setdefault(k, []).append(r)
    new_by_key = {}
    for r in new_rows:
        k = r["keyNoPob"]
        new_by_key.setdefault(k, []).append(r)

    delayed = []
    advanced = []
    touched_keys = set(prev_by_key.keys()) & set(new_by_key.keys())
    for k in touched_keys:
        olds = sorted(prev_by_key[k], key=lambda x: str(x.get("pob") or ""))
        news = sorted(new_by_key[k], key=lambda x: str(x.get("pob") or ""))
        pairs = min(len(olds), len(news))
        for i in range(pairs):
            old_pob = str(olds[i].get("pob") or "")
            new_pob = str(news[i].get("pob") or "")
            if old_pob == new_pob:
                continue
            old_dt = _parse_pob_to_datetime(old_pob)
            new_dt = _parse_pob_to_datetime(new_pob)
            if not old_dt or not new_dt:
                continue
            delta_min = int(round((new_dt.timestamp() - old_dt.timestamp()) / 60.0))
            row = {
                "vesselName": news[i].get("vesselName") or "—",
                "berthName": news[i].get("berthName") or "—",
                "empRb": news[i].get("empRb") or "N/A",
                "oldPob": old_pob,
                "newPob": new_pob,
                "deltaMinutes": delta_min,
            }
            if delta_min > 0:
                delayed.append(row)
            elif delta_min < 0:
                advanced.append(row)

    changed = delayed[:40] + advanced[:40]
    summary = {
        "previousUpdatedAt": prev_ts,
        "currentUpdatedAt": now,
        "addedCount": len(added),
        "removedCount": len(removed),
        "delayedCount": len(delayed),
        "advancedCount": len(advanced),
        "anyChange": bool(added or removed or delayed or advanced),
        "added": added,
        "removed": removed,
        "delayed": delayed[:60],
        "advanced": advanced[:60],
        "changed": changed[:80],
    }

    _save_schedule_monitor_state({"updatedAt": now, "rows": new_rows})
    return summary


def _compass_from_deg(deg):
    """Direção cardeal (16 setores) a partir de um ângulo em graus."""
    try:
        d = float(deg)
    except (TypeError, ValueError):
        return None
    dirs = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
            "S", "SSO", "SO", "OSO", "O", "ONO", "NO", "NNO"]
    return dirs[int((d % 360 + 11.25) // 22.5) % 16]


def _fetch_tide_context(lat: float, lon: float) -> dict:
    """Maré e corrente reais via Open-Meteo Marine.

    Maré: nível atual (m), tendência e próxima virada (preia/baixa-mar).
    Corrente: velocidade (nós) e direção (de onde/para onde flui) na hora atual.
    """
    url = (
        "https://marine-api.open-meteo.com/v1/marine"
        f"?latitude={lat}&longitude={lon}"
        "&hourly=sea_level_height_msl,ocean_current_velocity,ocean_current_direction"
        "&timezone=America%2FSao_Paulo"
    )
    try:
        req = urllib.request.Request(url, headers={"Accept": "application/json"}, method="GET")
        with urllib.request.urlopen(req, timeout=10) as resp:
            body = json.loads(resp.read().decode("utf-8"))
        hourly = body.get("hourly") or {}
        times = hourly.get("time") or []
        levels = hourly.get("sea_level_height_msl") or []
        pairs = [(t, l) for t, l in zip(times, levels) if isinstance(l, (int, float))]
        if len(pairs) < 3:
            return {"tide": "Sem dados de maré online."}

        now = datetime.now()
        # Índice da hora mais próxima de agora.
        def _parse(t):
            try:
                return datetime.fromisoformat(t)
            except ValueError:
                return None
        idx = 0
        best = None
        for i, (t, _l) in enumerate(pairs):
            dt = _parse(t)
            if dt is None:
                continue
            diff = abs((dt - now).total_seconds())
            if best is None or diff < best:
                best = diff
                idx = i

        level = pairs[idx][1]
        prev_level = pairs[idx - 1][1] if idx > 0 else level
        delta = level - prev_level
        if delta > 0.02:
            trend = "subindo"
        elif delta < -0.02:
            trend = "descendo"
        else:
            trend = "estável"

        # Próxima virada: onde a derivada troca de sinal.
        next_turn = None
        for i in range(idx, len(pairs) - 1):
            d0 = pairs[i][1] - pairs[i - 1][1] if i > 0 else delta
            d1 = pairs[i + 1][1] - pairs[i][1]
            if d0 >= 0 > d1:
                next_turn = ("preia-mar", pairs[i][0], pairs[i][1])
                break
            if d0 <= 0 < d1:
                next_turn = ("baixa-mar", pairs[i][0], pairs[i][1])
                break

        text = f"{level:.2f} m, {trend}"
        if next_turn:
            kind, t_iso, t_lvl = next_turn
            hhmm = (_parse(t_iso) or now).strftime("%H:%M")
            text += f" · próxima {kind} ~{hhmm} ({t_lvl:.2f} m)"

        # Corrente marítima (mesma hora de referência da maré).
        cur_speed_kn = None
        cur_dir_deg = None
        cur_dir_label = None
        cur_text = "Sem dados de corrente online."
        try:
            vel = hourly.get("ocean_current_velocity") or []
            dirs = hourly.get("ocean_current_direction") or []
            if idx < len(vel) and isinstance(vel[idx], (int, float)):
                cur_speed_kn = round(float(vel[idx]) / 1.852, 2)  # km/h -> nós
            if idx < len(dirs) and isinstance(dirs[idx], (int, float)):
                cur_dir_deg = round(float(dirs[idx]))
                cur_dir_label = _compass_from_deg(cur_dir_deg)
            if cur_speed_kn is not None:
                cur_text = f"{cur_speed_kn:.2f} nós"
                if cur_dir_label:
                    cur_text += f" para {cur_dir_label}"
        except Exception:
            pass

        return {
            "tide": text,
            "tideLevelM": round(level, 2),
            "tideTrend": trend,
            "tideNextTurn": (
                {"type": next_turn[0], "time": next_turn[1], "levelM": round(next_turn[2], 2)}
                if next_turn else None
            ),
            "tideSource": "open-meteo-marine",
            "current": cur_text,
            "currentSpeedKn": cur_speed_kn,
            "currentDirectionDeg": cur_dir_deg,
            "currentDirection": cur_dir_label,
        }
    except Exception as exc:
        return {"tide": "Sem dados de maré online.", "tideError": str(exc)}


def _fetch_metocean_context():
    # Dados operacionais basicos para a baia de Guanabara (ponto medio)
    lat, lon = -22.90, -43.17
    url = (
        "https://api.open-meteo.com/v1/forecast"
        f"?latitude={lat}&longitude={lon}"
        "&current=temperature_2m,wind_speed_10m,wind_direction_10m"
        "&hourly=wind_speed_10m"
        "&timezone=America%2FSao_Paulo"
    )
    tide_ctx = _fetch_tide_context(lat, lon)
    try:
        req = urllib.request.Request(url, headers={"Accept": "application/json"}, method="GET")
        with urllib.request.urlopen(req, timeout=10) as resp:
            body = json.loads(resp.read().decode("utf-8"))
        cur = body.get("current") or {}
        wind_kmh = cur.get("wind_speed_10m")
        wind_kn = round(float(wind_kmh) / 1.852, 1) if isinstance(wind_kmh, (int, float)) else None
        result = {
            "source": "open-meteo",
            "temperatureC": cur.get("temperature_2m"),
            "windSpeedKmh": wind_kmh,
            "windSpeedKn": wind_kn,
            "windDirectionDeg": cur.get("wind_direction_10m"),
            "windDirection": _compass_from_deg(cur.get("wind_direction_10m")),
        }
        result.update(tide_ctx)
        return result
    except Exception as exc:
        result = {
            "source": "fallback",
            "temperatureC": None,
            "windSpeedKmh": None,
            "windSpeedKn": None,
            "windDirectionDeg": None,
            "windDirection": None,
            "error": str(exc),
        }
        result.update(tide_ctx)
        return result


def _question_normalized(question: str) -> str:
    q = (question or "").strip().lower()
    return "".join(
        c for c in unicodedata.normalize("NFD", q) if unicodedata.category(c) != "Mn"
    )


def _question_asks_about_competitors(question: str) -> bool:
    q = _question_normalized(question)
    if not q:
        return False
    if "concorrent" in q or "competidor" in q or "rival" in q:
        return True
    if "quem sao" in q or "quem e " in q or "quem e o" in q:
        if "nosso" in q or "nossa" in q or "nos " in q or "empresa" in q:
            return True
    return False


def _competitor_fleet_catalog_text() -> str:
    lines = []
    for company in sorted(COMPETITOR_TUGS.keys()):
        bits = []
        for tug in COMPETITOR_TUGS[company]:
            label = (tug.get("name") or "").strip() or "—"
            mmsi = str(tug.get("mmsi") or "").strip()
            bits.append(f"{label} (MMSI {mmsi})" if mmsi else label)
        lines.append(f"- {company}: " + ", ".join(bits))
    return "\n".join(lines)


def _strategy_fallback_local_followup(question: str, context: dict) -> str:
    if not _question_asks_about_competitors(question):
        return ""
    comps = context.get("competitors") or []
    manobrando = [c for c in comps if c.get("insideManeuverGeofence")]
    lines_manobra = []
    for c in manobrando[:12]:
        nm = (c.get("name") or "—").strip()
        co = c.get("company") or "—"
        gf = (c.get("maneuverGeofenceName") or "—").strip()
        lines_manobra.append(f"  • {co} — {nm} (geofence: {gf})")
    manobra_txt = (
        "\nAgora em geofence de manobra (AIS):\n" + "\n".join(lines_manobra)
        if lines_manobra
        else "\nNeste momento nenhum rebocador concorrente da lista aparece com AIS dentro de geofence de manobra."
    )
    return (
        "\n\n---\n\n"
        "Resposta local (sem Grok)\n\n"
        "Os concorrentes que este dashboard acompanha por AIS são os rebocadores das empresas WIL e CAM "
        "(além da frota SAAM-BGRA, a vossa). Catálogo configurado:\n\n"
        f"{_competitor_fleet_catalog_text()}\n"
        "Na programação da Praticagem, o campo EMP.RB associa cada manobra a uma empresa. "
        "IMPORTANTE: o código EMP.RB «SAA» representa a SAAM (a nossa empresa); WIL e CAM são concorrentes. "
        "O resumo «Market share» no quadro acima conta quantas linhas na base existem por código EMP.RB.\n"
        f"{manobra_txt}"
    )


def _is_simple_greeting(question: str) -> bool:
    q = _question_normalized(question)
    if not q:
        return False
    for ch in "!?.,;:":
        q = q.replace(ch, " ")
    parts = [p for p in q.split() if p]
    if not parts:
        return False
    if len(parts) <= 2 and parts[0] in ("ola", "oi", "hi", "hello", "hey", "eai"):
        return True
    if (
        len(parts) >= 2
        and len(parts) <= 4
        and parts[0] in ("bom", "boa")
        and parts[1] in ("dia", "tarde", "noite")
    ):
        return True
    return False


# ===== Expansões de conhecimento do KRATOS =====

def _bearing_compass(lat1, lon1, lat2, lon2):
    """Direção cardeal aproximada de (lat1,lon1) para (lat2,lon2)."""
    dlon = math.radians(lon2 - lon1)
    y = math.sin(dlon) * math.cos(math.radians(lat2))
    x = (
        math.cos(math.radians(lat1)) * math.sin(math.radians(lat2))
        - math.sin(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.cos(dlon)
    )
    brg = (math.degrees(math.atan2(y, x)) + 360) % 360
    dirs = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
            "S", "SSO", "SO", "OSO", "O", "ONO", "NO", "NNO"]
    return dirs[int((brg + 11.25) // 22.5) % 16]


def _compact_vessels_overview(limit: int = 150):
    """TODAS as embarcações com posição (mercantes incluídos), em formato compacto.

    Permite ao KRATOS responder "que navio está fundeado em tal ponto".
    """
    items = []
    for v in list(latest_vessel_by_mmsi.values()):
        lat, lon = v.get("latitude"), v.get("longitude")
        if lat is None or lon is None:
            continue
        mmsi = str(v.get("mmsi") or "")
        try:
            sog = float(v.get("sog") or 0)
        except (TypeError, ValueError):
            sog = 0.0
        entry = {
            "name": (v.get("shipName") or "").strip() or f"MMSI {mmsi}",
            "mmsi": mmsi,
            "category": v.get("shipCategory") or "outros",
            "lat": round(float(lat), 5),
            "lon": round(float(lon), 5),
            "sogKn": round(sog, 1),
            "moving": sog >= 0.5,
            "geofences": v.get("geofencesInside") or [],
        }
        if v.get("isSaamBgra"):
            entry["fleet"] = "SAAM"
        elif mmsi in COMPETITOR_COMPANY_BY_MMSI:
            entry["fleet"] = COMPETITOR_COMPANY_BY_MMSI[mmsi]
        items.append(entry)
    # Prioridade: frota própria/concorrentes, depois quem está em geofence, depois demais.
    items.sort(key=lambda x: (0 if x.get("fleet") else (1 if x["geofences"] else 2), x["name"]))
    return items[:limit]


def _geofences_summary():
    """Demarcação das geofences: vértices, centro e dimensão aproximada.

    Permite ao KRATOS orientar onde fica e como é demarcada cada área.
    """
    ensure_geofences_loaded()
    with geofence_lock:
        snap = list(geofences)
    out = []
    for g in snap:
        if not g.get("isActive", True):
            continue
        geom = g.get("geometry") or {}
        entry = {
            "name": (g.get("name") or "").strip(),
            "type": g.get("type"),
            "fleetScope": g.get("fleetScope", "all"),
        }
        coords = geom.get("coordinates") or []
        pts = [c for c in coords if isinstance(c, (list, tuple)) and len(c) == 2]
        if pts:
            lats = [p[0] for p in pts]
            lons = [p[1] for p in pts]
            entry["centerLat"] = round(sum(lats) / len(lats), 5)
            entry["centerLon"] = round(sum(lons) / len(lons), 5)
            entry["vertices"] = [[round(p[0], 5), round(p[1], 5)] for p in pts][:20]
            if g.get("type") == "corredor":
                # Corredor = rota: comprimento total ao longo da polilinha.
                length = sum(
                    _haversine_nm(pts[i][0], pts[i][1], pts[i + 1][0], pts[i + 1][1])
                    for i in range(len(pts) - 1)
                )
                entry["role"] = "corredor de trafego (rota navegavel)"
                entry["lengthNm"] = round(length, 2)
                entry["startLat"], entry["startLon"] = round(pts[0][0], 5), round(pts[0][1], 5)
                entry["endLat"], entry["endLon"] = round(pts[-1][0], 5), round(pts[-1][1], 5)
            else:
                entry["approxSpanNm"] = round(
                    _haversine_nm(min(lats), min(lons), max(lats), max(lons)), 2
                )
        center = geom.get("center") or []
        if len(center) == 2:
            entry["centerLat"] = round(center[0], 5)
            entry["centerLon"] = round(center[1], 5)
            entry["radiusMeters"] = geom.get("radiusMeters")
        out.append(entry)
    return out


def _distances_and_eta_summary(enriched_maneuvers):
    """Distância/ETA de cada rebocador (SAAM e concorrente) aos navios das próximas manobras."""
    vessels_by_norm = {}
    for v in latest_vessel_by_mmsi.values():
        n = _normalize_text(v.get("shipName"))
        if n and v.get("latitude") is not None:
            vessels_by_norm[n] = v
    tugs = []
    for mmsi in SAAM_BGRA_MMSI_SET:
        t = latest_vessel_by_mmsi.get(mmsi)
        if t and t.get("latitude") is not None:
            tugs.append((t, "SAAM"))
    for mmsi, company in COMPETITOR_COMPANY_BY_MMSI.items():
        t = latest_vessel_by_mmsi.get(mmsi)
        if t and t.get("latitude") is not None:
            tugs.append((t, company))
    rows = []
    for m in enriched_maneuvers[:10]:
        v = vessels_by_norm.get(_normalize_text(m.get("vesselName")))
        if not v:
            continue
        dists = []
        for t, fleet in tugs:
            d = _haversine_nm(t["latitude"], t["longitude"], v["latitude"], v["longitude"])
            try:
                sog = float(t.get("sog") or 0)
            except (TypeError, ValueError):
                sog = 0.0
            eta = round(d / sog * 60) if sog >= 0.5 else None
            dists.append({
                "tug": (t.get("shipName") or "").strip() or t.get("mmsi"),
                "fleet": fleet,
                "distanceNmStraight": round(d, 2),
                "etaMinStraightLine": eta,
                "note": "linha reta; trajeto real pelo corredor navegavel e maior",
            })
        dists.sort(key=lambda x: x["distanceNmStraight"])
        rows.append({
            "vessel": m.get("vesselName"),
            "pob": m.get("pob"),
            "vesselLat": round(v["latitude"], 5),
            "vesselLon": round(v["longitude"], 5),
            "tugDistances": dists[:8],
        })
    return rows


def _recent_tracks_summary():
    """Tendência de deslocamento recente dos rebocadores (nossos e concorrentes),
    derivada do buffer de posições (rastro/intenção de movimento)."""
    tracked = set(SAAM_BGRA_MMSI_SET) | set(COMPETITOR_COMPANY_BY_MMSI.keys())
    pts_by = {}
    for item in list(recent_vessels):
        m = str(item.get("mmsi") or "")
        if m in tracked and item.get("latitude") is not None:
            pts_by.setdefault(m, []).append(item)
    out = []
    for mmsi, pts in pts_by.items():
        if len(pts) < 2:
            continue
        first, last = pts[0], pts[-1]
        dist = _haversine_nm(first["latitude"], first["longitude"], last["latitude"], last["longitude"])
        t0 = _parse_vessel_timestamp_unix(first)
        t1 = _parse_vessel_timestamp_unix(last)
        minutes = round((t1 - t0) / 60.0, 1) if (t0 and t1 and t1 > t0) else None
        entry = {
            "name": (last.get("shipName") or "").strip() or f"MMSI {mmsi}",
            "fleet": "SAAM" if mmsi in SAAM_BGRA_MMSI_SET else COMPETITOR_COMPANY_BY_MMSI.get(mmsi),
            "points": len(pts),
            "distanceNm": round(dist, 2),
            "windowMinutes": minutes,
        }
        if dist >= 0.05:
            entry["trend"] = (
                f"deslocando-se para {_bearing_compass(first['latitude'], first['longitude'], last['latitude'], last['longitude'])}"
            )
        else:
            entry["trend"] = "praticamente parado"
        out.append(entry)
    return out


def _fleet_ais_status():
    """Status AIS de cada rebocador cadastrado (nosso e concorrente).

    Sem posição ou sem atualização há muito tempo => possivelmente saiu barra
    fora / fora da área de cobertura do Rio de Janeiro.
    """
    now = time.time()

    def status_row(mmsi: str, name: str, fleet: str):
        v = latest_vessel_by_mmsi.get(mmsi)
        row = {"mmsi": mmsi, "name": name, "fleet": fleet}
        if not v or v.get("latitude") is None:
            row["aisStatus"] = (
                "sem sinal AIS na cobertura (Rio de Janeiro) — possivelmente barra fora / fora da área"
            )
            return row
        ts = _parse_vessel_timestamp_unix(v)
        age_min = round((now - ts) / 60) if ts and now > ts else None
        row["lastLat"] = round(float(v["latitude"]), 5)
        row["lastLon"] = round(float(v["longitude"]), 5)
        if age_min is not None:
            row["lastSeenMinutesAgo"] = age_min
        if age_min is not None and age_min > 60:
            row["aisStatus"] = (
                f"sem atualização há {age_min} min — possivelmente saiu barra fora ou desligou o AIS"
            )
        else:
            row["aisStatus"] = "ativo"
        return row

    rows = [status_row(m, n, "SAAM") for m, n in SAAM_BGRA_NAMES.items()]
    for company, tugs in COMPETITOR_TUGS.items():
        rows.extend(status_row(t["mmsi"], t.get("name", t["mmsi"]), company) for t in tugs)
    return rows


# ===== Perfil do usuário (KRATOS conhece quem opera) =====
def kratos_profile_file() -> Path:
    return user_data_dir(DASHBOARD_USER_ID) / "kratos_user_profile.json"


def load_kratos_profile() -> dict:
    try:
        p = kratos_profile_file()
        if p.exists():
            data = json.loads(p.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        pass
    return {}


def save_kratos_profile(profile: dict):
    try:
        kratos_profile_file().write_text(
            json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except OSError:
        pass


def _extract_profile_tags(answer: str) -> str:
    """Extrai tags [PERFIL: ...] da resposta, persiste no perfil e remove do texto."""
    import re
    tags = re.findall(r"\[PERFIL:\s*([^\]]+)\]", answer)
    if not tags:
        return answer
    profile = load_kratos_profile()
    for tag in tags:
        for part in tag.split(";"):
            if "=" not in part:
                continue
            key, val = part.split("=", 1)
            key = _normalize_text(key)
            val = val.strip()
            if not val:
                continue
            if key in ("nome", "name"):
                profile["name"] = val
            elif key in ("funcao", "cargo", "role"):
                profile["role"] = val
            elif key in ("padrao", "padrão", "pattern", "preferencia"):
                patterns = profile.setdefault("patterns", [])
                if val not in patterns:
                    patterns.append(val)
                profile["patterns"] = patterns[-20:]
    save_kratos_profile(profile)
    return re.sub(r"\s*\[PERFIL:[^\]]+\]", "", answer).strip()


def _profile_instruction_block() -> str:
    """Bloco de instruções sobre o usuário (apresentação na 1ª conversa + memória)."""
    profile = load_kratos_profile()
    parts = []
    if profile.get("name"):
        linha = f"PERFIL DO USUARIO: nome {profile['name']}"
        if profile.get("role"):
            linha += f", funcao {profile['role']}"
        parts.append(linha + ". Chame-o pelo nome com naturalidade.")
        if profile.get("patterns"):
            parts.append("Padroes de decisao ja observados: " + "; ".join(profile["patterns"][:10]) + ".")
    else:
        parts.append(
            "PRIMEIRO CONTATO: voce ainda nao conhece este usuario. Apresente-se em UMA frase "
            "como KRATOS e pergunte o nome dele — nada mais. NAO traga status, dados nem "
            "recomendacoes neste primeiro momento; apenas cumprimente e fique a disposicao."
        )
    parts.append(
        "MEMORIA DE PERFIL: quando o usuario informar nome, funcao ou um padrao de decisao "
        "relevante (ex.: 'prefiro 4 rebocadores em navio acima de 300m'), registre adicionando "
        "ao FINAL da resposta a tag [PERFIL: nome=...; funcao=...; padrao=...] com apenas os "
        "campos novos. A tag e removida antes de exibir ao usuario — nunca a mencione."
    )
    return "\n".join(parts)


def _strategy_context_dict():
    ensure_geofences_loaded()
    ensure_saa_maneuvers_loaded()
    with geofence_lock:
        geofence_snapshot = list(geofences)
    with saa_maneuvers_lock:
        saa_snapshot = _dedupe_saa_maneuvers(list(saa_maneuvers_list))
    saam_positions = []
    for mmsi in SAAM_BGRA_MMSI_SET:
        v = latest_vessel_by_mmsi.get(mmsi)
        if not v:
            continue
        saam_positions.append(
            {
                "mmsi": mmsi,
                "name": (v.get("shipName") or "").strip() or f"MMSI {mmsi}",
                "latitude": v.get("latitude"),
                "longitude": v.get("longitude"),
                "speed": v.get("sog") if v.get("sog") is not None else v.get("speed"),
                "heading": v.get("heading"),
                "insideGeofences": get_vessel_geofences(v),
            }
        )
    upcoming = sorted(
        saa_snapshot,
        key=lambda x: str(x.get("pob") or ""),
    )[:40]
    # Enriquece cada manobra com leitura tática: rebocadores estimados,
    # valor comercial relativo e se é nossa (SAAM) ou de concorrente.
    enriched = []
    for m in upcoming:
        emp = str(m.get("empRb") or "").strip().upper()
        item = dict(m)
        item["estimatedTugs"] = _estimate_tugs_required(m)
        item["commercialValue"] = _estimate_maneuver_value(m)
        item["isOwn"] = emp == OWN_COMPANY_EMP_RB
        item["isCompetitor"] = emp in {"WIL", "CAM"}
        enriched.append(item)
    competitors = _competitor_status_rows(geofence_snapshot)
    schedule_changes = _build_schedule_changes(saa_snapshot)
    memory_items = load_strategy_memory()
    return {
        "timestamp": get_now_iso(),
        "saamTugs": saam_positions,
        "scheduledManeuvers": enriched,
        "scheduledManeuverTotal": len(saa_snapshot),
        "competitors": competitors,
        "marketShare": _market_share_rows(saa_snapshot),
        "simultaneousManeuvers": _simultaneous_maneuvers_summary(saa_snapshot),
        "scheduleChanges": schedule_changes,
        "metocean": _fetch_metocean_context(),
        "userLearnedNotes": memory_items[:30],
        # Expansões de conhecimento:
        "vesselsOverview": _compact_vessels_overview(),
        "tugOperatingHours": _saam_operating_hours_rows(),
        "fleetAisStatus": _fleet_ais_status(),
        "geofencesMap": _geofences_summary(),
        "maneuverDistances": _distances_and_eta_summary(enriched),
        "recentTracks": _recent_tracks_summary(),
        "userProfile": load_kratos_profile(),
    }


def _strategy_fallback_answer(question: str, context: dict) -> str:
    saam = context.get("saamTugs") or []
    comps = context.get("competitors") or []
    maneuvers = context.get("scheduledManeuvers") or []
    m_total = int(context.get("scheduledManeuverTotal", len(maneuvers)))
    manobrando = [c for c in comps if c.get("insideManeuverGeofence")]
    top = (context.get("marketShare") or {}).get("rows") or []
    top_txt = ", ".join(f"{x['empRb']}: {x['count']}" for x in top[:4]) if top else "sem dados"
    intro = (
        "Olá. Bora olhar o cenário operacional agora (modo local, sem chamada ao Grok):\n\n"
        if _is_simple_greeting(question)
        else "Leitura rápida do cenário operacional (modo local, sem Grok):\n\n"
    )
    hint = (
        "Opcional: defina a variável de ambiente XAI_API_KEY (API xAI) para respostas com Grok "
        "usando este contexto — por exemplo num ficheiro `.env` na raiz do projeto ou nas variáveis do servidor."
    )
    follow = _strategy_fallback_local_followup(question, context)
    return (
        intro
        + f"Neste momento temos {len(saam)} rebocadores SAAM com posição AIS e {m_total} manobras na base da Praticagem.\n"
        + f"Concorrentes em geofence de manobra agora: {len(manobrando)}. Market share por EMP.RB (contagem): {top_txt}.\n"
        + f"Sobre a sua pergunta ('{question}'), com este recorte eu começaria por estes pontos práticos:\n"
        + "- priorizar janelas com sobreposição de manobras e necessidade de mais rebocadores;\n"
        + "- monitorar entradas simultâneas de concorrentes nas geofences críticas;\n"
        + "- revisar alocação com base no mix do turno (SAA = nossa SAAM; WIL e CAM concorrentes).\n"
        + follow
        + "\n\n"
        + hint
    )


def _assistant_profile_instruction() -> str:
    profile = ASSISTANT_PROFILE
    if profile == "executivo":
        return (
            "Perfil executivo: comece pelo impacto e recomendacao, use poucas frases e priorize decisao."
        )
    if profile == "operacional":
        return (
            "Perfil operacional de patio: detalhe janelas, geofences, alocacao e risco de conflito."
        )
    if profile == "despacho":
        return (
            "Perfil despacho: resposta curta, direta e orientada a acao imediata."
        )
    return (
        "Perfil hibrido: seja objetivo, mas inclua contexto essencial quando ajudar na decisao."
    )


KRATOS_SYSTEM_PROMPT = (
    "Voce e o KRATOS — assistente operacional estrategista de navegacao e apoio "
    "portuario do Porto do Rio de Janeiro e da Baia de Guanabara. Sua missao e manter "
    "o usuario sempre antecipado: ler o cenario, prever movimentos e municiar a equipe "
    "com informacao acionavel. Ao se apresentar, use o nome KRATOS.\n\n"
    "LINGUAGEM: use SEMPRE a terminologia da area de navegacao e apoio portuario "
    "(manobra, atracacao, desatracacao, fundeio, praticagem, POB, berco, calado, "
    "rebocador, janela de manobra, alocacao de frota). NUNCA use termos ou metaforas "
    "de jogos/xadrez (nada de 'tabuleiro', 'peca', 'jogada', 'adversario', 'lance'). "
    "Fale como um profissional experiente do setor falando com outro.\n\n"
    "QUEM E QUEM (campo EMP.RB da Praticagem): 'SAA' = SAAM, a NOSSA empresa (nossa frota); "
    "'WIL' e 'CAM' sao empresas concorrentes. Trate SAA como nos.\n\n"
    "COMO VOCE ATUA:\n"
    "- Leia o cenario operacional: posicao e deslocamento de cada rebocador (nosso e "
    "concorrente), a programacao de manobras (POB), as caracteristicas dos navios "
    "(LOA/boca/DWT) e o valor comercial relativo de cada manobra (campo commercialValue: "
    "alto/medio/padrao).\n"
    "- Antecipe a concorrencia: a partir da posicao/rumo dos rebocadores WIL/CAM e da "
    "programacao, avalie para qual manobra eles podem estar se posicionando e o impacto.\n"
    "- Reconheca padroes: use as notas em userLearnedNotes e o historico da conversa para "
    "entender o modo de operar de cada concorrente.\n"
    "- Considere as condicoes meteoceanicas (campo metocean): vento (windSpeedKn/windDirection), "
    "mare (tide/tideTrend), CORRENTE (currentSpeedKn/currentDirection) e temperatura (temperatureC). "
    "Vento, mare e CORRENTE afetam a janela de manobra e o numero de rebocadores — varias regras da "
    "NPCP-RJ limitam manobra por corrente (ex.: <=0,5 no, <=0,8 no). Cruze a corrente atual com o "
    "limite do terminal e sinalize risco/oportunidade.\n"
    "- Proteja a nossa programacao: recomende acoes que ganhem manobra ou melhorem o "
    "posicionamento SEM comprometer os compromissos ja assumidos pela SAAM.\n"
    "- Enxergue TODO o espelho d'agua: vesselsOverview lista todas as embarcacoes visiveis "
    "(mercantes incluidos) com posicao lat/lon, categoria, velocidade e geofences — use para "
    "responder 'que navio esta fundeado em tal ponto'.\n"
    "- Conheca a demarcacao das areas: geofencesMap traz vertices, centro e dimensao de cada "
    "geofence (bercos, base de rebocadores, poligonos) — oriente sobre limites quando perguntado.\n"
    "- Use distancias e ETA com ressalva: maneuverDistances traz distancia EM LINHA RETA "
    "(distanceNmStraight) e o ETA correspondente. IMPORTANTE: os rebocadores navegam por "
    "CORREDORES DE TRAFEGO (canais navegaveis, limitados pela profundidade) — como estradas no "
    "mar. Por isso o trajeto real e o tempo sao MAIORES que a linha reta. Trate a linha reta como "
    "estimativa minima/referencia; deixe claro que a distancia efetiva segue o corredor "
    "navegavel. Em geofencesMap, os itens com role 'corredor de trafego' sao essas rotas "
    "(com lengthNm, inicio e fim) cadastradas pelo usuario e nomeadas; use-as para raciocinar "
    "sobre o caminho real, citando o corredor pelo nome quando fizer sentido. Quando houver rastro "
    "recente (recentTracks), use-o para refinar o trajeto.\n"
    "- Leia a tendencia de deslocamento: recentTracks mostra o rastro recente dos rebocadores "
    "(direcao, distancia percorrida, janela de tempo) — antecipe para onde cada um esta indo.\n"
    "- Controle de fadiga: tugOperatingHours traz as HORAS DE OPERACAO do dia de cada rebocador "
    "SAAM (operatingHours), o estado (operando/parado), se esta recuperado e o status (ok/atencao "
    ">=7h/limite >=8h). Por seguranca, o limite e 8h de operacao por rebocador; ao recomendar "
    "alocacao, EVITE rebocadores em 'atencao' ou 'limite' e prefira os de menor carga (revezamento). "
    "Operacao = qualquer movimento (manobra ou deslocamento); repouso = atracado/parado.\n"
    "- Verificacao de frota: fleetAisStatus traz o status AIS de cada rebocador cadastrado "
    "(nosso e concorrente). Se um rebocador esta 'sem sinal' ou sem atualizacao ha muito tempo, "
    "informe que ele possivelmente saiu BARRA FORA / esta fora da area de cobertura do Rio de "
    "Janeiro (ou com AIS desligado). Se perguntarem por rebocador ou MMSI que nao consta na "
    "frota cadastrada, procure em vesselsOverview; nao constando, diga que nao ha sinal na "
    "cobertura do Rio e lembre que e possivel cadastra-lo no painel 'Frota' do mapa.\n\n"
    "FOCO: mantenha a conversa exclusivamente em operacoes portuarias, navegacao, apoio "
    "maritimo, meteorologia operacional, seguranca da navegacao e normas maritimas "
    "(NORMAM/Marinha do Brasil, SOLAS, MARPOL, COLREG). Pode orientar sobre essas normas com "
    "base no seu conhecimento, recomendando confirmar na publicacao oficial em temas criticos. "
    "Se o usuario desviar para assunto fora desse escopo, redirecione com cortesia para o "
    "contexto operacional.\n\n"
    "COMO VOCE FALA: portugues claro, direto e natural, como um parceiro operacional ao "
    "lado da equipe de manobra.\n"
    "UNIDADES E SIGLAS POR EXTENSO: ao falar (e tambem ao escrever), pronuncie as unidades e "
    "siglas por extenso, nunca soletre as letras. Use: 'milhas nauticas' para nm/NM; "
    "'milha nautica' no singular; 'nos' para kn/knots (ex.: '8 nos'); 'metros' para m; "
    "'graus' para o simbolo de grau; 'SAAM' fala-se 'saam' (uma palavra, nao soletrar); "
    "'POB' fala-se 'P-O-B'; 'MMSI' fala-se 'M-M-S-I'; 'AIS' fala-se 'A-I-S'; 'BG' diga "
    "'Baia de Guanabara'; 'EMP.RB' diga 'empresa de reboque'. Ex.: '0,8 milha nautica', "
    "'reboque a 6 nos'.\n"
    "REGRA DE OURO — SEJA COMEDIDO E REATIVO:\n"
    "- Responda APENAS o que o usuario perguntou. NUNCA despeje dados, listas ou relatorios "
    "nao solicitados. O contexto operacional que voce recebe e so para consulta — use o "
    "trecho relevante a pergunta, nada alem.\n"
    "- Respostas CURTAS, em partes: 1 a 3 frases. Se houver muito a dizer, de o essencial e "
    "PERGUNTE se o usuario quer mais detalhes, em vez de mandar tudo de uma vez.\n"
    "- Nao se antecipe com alertas ou recomendacoes a cada turno; so recomende quando "
    "perguntado ou quando o usuario claramente pedir uma decisao.\n"
    "- Se faltar dado, diga com transparencia. Faca perguntas de volta quando ajudar a decidir.\n"
    "SAUDACAO: se o usuario apenas cumprimentar (ex.: 'ola', 'bom dia'), responda a saudacao "
    "de forma breve e cordial, pergunte o nome dele (se ainda nao souber) e coloque-se a "
    "disposicao — SEM disparar dados operacionais ou status. Espere ele perguntar."
)

# Guia de telas do app: o KRATOS conhece a propria interface e conduz tours.
KRATOS_APP_GUIDE = (
    "VOCE CONHECE A INTERFACE DO APP (KRATOS — Inteligencia Naval Estrategica). "
    "Use este guia para explicar telas e funcionalidades quando perguntarem:\n"
    "PAGINA MAPA (inicial): "
    "Topo esquerda: logo KRATOS. "
    "Topo centro: lampadas da frota SAAM por sigla (ex.: PX, PA, CH, HL, LT, AT) — azul = na base, "
    "verde = fora da base (operando), cinza = offline; legenda logo abaixo. "
    "Topo direita: chips de mare (nivel do mar com seta de tendencia), vento (kn e direcao), "
    "corrente maritima (kn e direcao) e temperatura; botoes DB (abre o Painel Estrategico), "
    "GR (abre o Grafo Estrategico) e N "
    "(mostra/oculta os nomes das embarcacoes — navios programados para manobra da SAAM aparecem "
    "com etiqueta dourada). "
    "Barra lateral esquerda (dock), de cima para baixo: Status (conexao AIS e mensagens), "
    "Area (area de monitoramento), Filtros (por tipo de embarcacao e frota), Geofences (criar e "
    "editar bercos, bases, poligonos e CORREDORES DE TRAFEGO — rotas navegaveis nomeadas — direto "
    "no mapa), SAAM-BGRA (painel da frota com status), "
    "Entrada/Saida BG (fluxo de embarcacoes na baia e grafico 24h), Tempo (previsao, mare e "
    "vento) e Frota (incluir, substituir ou remover MMSIs de rebocadores SAAM/WIL/CAM). "
    "No mapa: embarcacoes sao setas coloridas por tipo (verde carga/conteiner, vermelho "
    "petroleiro, azul passageiros, teal rebocador); a seta aponta o rumo; circulo = parada/"
    "fundeada; frota SAAM em dourado brilhante; concorrentes WIL/CAM com anel vermelho pulsante; "
    "clique numa embarcacao abre popup com dados (LOA, boca, velocidade, rumo, geofences) e botao "
    "de criar rastro (trilha). "
    "Canto inferior esquerdo: caixa 'KRATOS Insights' datilografada com leituras do cenario ao "
    "vivo e botao de fone (abre conversa por VOZ AO VIVO comigo). "
    "Canto inferior direito: contador BG (dentro) / MAR (fora). Ha um splash de abertura com o simbolo K.\n"
    "PAGINA PAINEL ESTRATEGICO (botao DB): "
    "Topo: botao Manual (manual do usuario) e link de volta ao mapa. Secoes na ordem: "
    "1) KRATOS — Assistente Estrategico: conversa comigo por texto, botoes Voz ao vivo (fone), "
    "Falar (ditado por microfone), Voz (leitura das respostas), Gerar relatorio, Gerar insights, "
    "Limpar conversa e campo de Aprendizado para minha memoria; "
    "2) Monitor de alteracoes da programacao (atrasos, adiantamentos, entradas e saidas); "
    "3) Status das geofences (lampadas SAA e SAAM indicando manobra por area); "
    "4) Manobras SAA — tabela da programacao da Praticagem-RJ (POB, navio, caracteristicas), com "
    "destaque azul antes do horario, amarelo faltando 30 min e verde piscando apos a POB; "
    "sincroniza automaticamente a cada 5 min; "
    "5) Frota SAAM — grafico de manobras, horas em geofence e milhas nauticas por rebocador; "
    "6) Market share (Praticagem-RJ) — grafico de rosca por empresa, com a fatia SAAM em dourado "
    "e janelas hoje/7 dias/30 dias; "
    "7) Integracao Obsidian (exportacao de dados para vault).\n"
    "PAGINA GRAFO ESTRATEGICO (botao GR): visualizacao em grafo (2D/3D) das conexoes entre as "
    "entidades da operacao.\n"
    "TOUR GUIADO: se o usuario pedir um tour ou apresentacao das telas, conduza POR PARTES: "
    "pergunte em qual pagina ele esta (mapa ou painel), explique UMA area por turno (2 a 4 "
    "frases) e pergunte se quer seguir para a proxima area. NUNCA despeje o tour inteiro de uma "
    "vez. Encerre o tour quando ele pedir."
)


# Conhecimento estrategico das normas locais (NPCP-RJ) — o KRATOS usa como
# referencia para antecipar DEMANDA DE REBOCADOR, janelas de manobra e restricoes.
# Fonte: NPCP-RJ (3a Revisao), atualizada pelas Portarias CPRJ/COMOPNAV/MB
# no 11 (06/03/2026) e no 110 (14/05/2026), Capitania dos Portos do Rio de Janeiro.
KRATOS_NPCP_KNOWLEDGE = (
    # Foco operacional atual: Baia de Guanabara. Fonte: NPCP-RJ (3a Rev., Mod.2)
    # + Portarias CPRJ no 11 e no 110/2026. Detalhe humano em
    # docs/conhecimento/npcp-rj-bg-terminais.md (fichas por terminal).
    """NORMAS LOCAIS DO PORTO (NPCP-RJ) — CONHECIMENTO ESTRATEGICO — FOCO: BAIA DE GUANABARA.
Fontes: NPCP-RJ (3a Revisao, Mod.2) e Portarias CPRJ/COMOPNAV/MB no 11 (06/03/2026) e no 110 (14/05/2026), em vigor. Use para antecipar demanda de rebocadores (numero, tipo e TTE — tonelada de tracao estatica), janelas e restricoes; cite o terminal/regra aplicada e, em decisao critica, recomende confirmar na publicacao oficial. As demais subzonas da ZP-15 (Itaguai/Sepetiba, Angra/TPAR, Acu, Paraty, Forno) constam na norma, mas o foco operacional atual e a Baia de Guanabara.

POR QUE IMPORTA PARA A SAAM: cada terminal define QUANTOS rebocadores e qual TTE minimo a manobra exige — demanda direta da nossa frota. Onde ha dispensa (navios/EAM com DP e sistemas plenos), a demanda cai. Janelas de mare/vento e restricoes diurno/noturno CONCENTRAM manobras no tempo — antecipe picos e sobreposicao.

CAIS COMERCIAL (Gamboa / Sao Cristovao / Caju):
- Gamboa: LOA <=165 m e calado <=8 m = 2 rebocadores TKM >=43 TTE; LOA 165-200 m com giro = 2 (1 azimutal + 1 TKM) ou 3 TKM, todos >=43 TTE (sem giro: 2 TKM >=43 TTE); LOA >200 m ou calado >8 m = 2 azimutais >=43 TTE. Giro leve recomendado (atracar BE para carregar; BB para descarregar).
- Sao Cristovao: LOA <=120 m (diurno/noturno) = 2 rebocadores >=40 TTE; 120-150 m = 3 >=40 TTE (2 se tiver bow thruster); 150-185 m (SO DIURNO) = 3 >=40 TTE sendo 2 azimutais (2 azimutais com thruster; ate 185 m exige cabecos 150-178 livres de navio com boca >30 m).
- Calados por trecho: cabecos 36-110 = 10,30 m (max 11,00 com mare); 110-129 = 9,00 (max 9,70); 129-205 = 8,50 (max 9,00; LOA <=185 m); 205-216 = 8,20 (max 9,00).

TECON-RJ (ate 349 m de LOA; calado 14,50 m, ate 15,30 com mare; prioridade do canal: porta-conteineres do TECON e Ro-Ro; programacao previa + aviso a Praticagem VHF 12 com 30 min). Escada de reboque:
- LOA <=155 m: 2 rebocadores, somatorio 80 TTE (>=1 azimutal, min 40/un.);
- 155-200 m: 2 azimutais, somatorio 80 TTE;
- 200-250 m (ou DWT 40-60 mil t): 3 rebocadores, somatorio 90 TTE (2 azimutais); com bow thruster: 2 azimutais, somatorio 80;
- 250-290 m (ou DWT 60-80 mil t): 3 rebocadores, somatorio 140 TTE; com thruster: somatorio 120;
- LOA >290 m sem thruster: 4 rebocadores, somatorio 160 TTE (2 azimutais + 2 TKM, min 40/un.); com thruster: 3, somatorio 140;
- Gigantes com calado 14,5-14,6 m: LOA 295-335 m (boca 42-48,5 m) = 2x60 TTE + 2x55 TTE, todos azimutais; LOA 335-349 m (boca 48,5-52 m) = 2x70 TTE + 2x60 TTE, todos azimutais. Dois praticos quando LOA >295 m ou boca >42 m; vento <=15 nos; manobras diurnas para os muito grandes.
- Restricao de porte no trecho 197-216 (via TECON): LOA 120-150 m diurno/noturno; 150-185 m SO DIURNO; acima de 185 m nao passa.

TERMINAL DE OLEO (cabecos 197-205): navios-tanque LOA <=185 m = 3 rebocadores >=45 TTE (2 azimutais + 1 TKM multi-eixo). Canal por calado/bordo: entrada com calado 6,10-7,80 m SO via canal do TECON; saida atracado BE via Canal Comercial (calado <=6,10 m); atracado BB sai via TECON.

TERMINAL ALMIRANTE TAMANDARE (Ilha d'Agua, Petrobras — pieres PP-I, PP-II, PS-I, PS-II; canal de 10 milhas dragado a 17 m). TODOS os rebocadores azimutais >=50 TTE:
- PP-I (LOA 279,5 m): DWT <=60 mil t e calado <=12 m = 2; calado >12 m = 3; DWT >60 mil = 4.
- PP-II (LOA 259 m): 3 (DWT <=60 mil, calado <=12 m) a 4 (DWT 60-135 mil ou calado >12 m; 90-135 mil = diurno).
- PS-I (LOA 186,4 m): 2 a 3 conforme giro. PS-II (LOA 175 m): DWT <=7 mil = 2; 7-15 mil = 3; >15 mil = 4.
- Calado >11,50 m vindo de fora da BG: entrada e saida DIURNAS. Alta demanda recorrente de reboque pesado.

GLP TAIC/TAIR (Ilhas Comprida e Redonda, Petrobras): TAIC = 3 azimutais >=45 TTE (calado leve ate 6,40 m) ou 4 azimutais >=45 TTE (carregado ate 10,60 m). TAIR = minimo 3 azimutais >=45 TTE; atracacao SO DIURNA e contra a corrente. Os rebocadores ESCOLTAM o navio desde antes da Ponte Rio-Niteroi ate o terminal (entrada) e ate a Ponte (saida).

GNL (terminal flexivel de regaseificacao, pieres PG-1/PG-2, LNGC ate LOA 315 m): entrada = 2 rebocadores (azimutais ou TKM) somatorio >=120 TTE com CABOS PASSADOS no vao central da Ponte Rio-Niteroi; atracacao/desatracacao/saida = 3 azimutais >=40 TTE; apos atracado, 1 rebocador >=45 TTE de PRONTIDAO 24 h. Entrada so diurna; zona de seguranca 500-600 m; transito interno <=8 nos; VHF 13.

MANGUINHOS (quadro de boias, navio ate LOA ~190 m): amarracao = 4 rebocadores >=45 TTE (3 azimutais + 1 TKM); desamarracao = 4 (2 azimutais + 2 TKM); STANDBY de 2 rebocadores >=45 TTE prontos em ate 15 min enquanto o navio estiver amarrado. Diurno; mare vazante; vento <=16 nos.

NEOLUBES (Ponte do Thun, Shell; LOA ate 206 m): 2 azimutais >=50 TTE em TODAS as manobras; SO DIURNO; 1 navio por vez no canal; vento <=15 nos.

NITEROI / CAJU — dispensas e casos leves (menos demanda):
- Brasco Base Niteroi: navio-tipo com DP dispensa rebocador; manobras SO nos estofos de mare; 1 navio por vez no canal.
- MacLaren Ilha da Conceicao e BHGE Caximbau: navio tipo 1 (LOA 146 m) = 2 rebocadores >=40 TTE (azimutal/TKM/ASD); tipo 2 (LOA 97 m) com sistemas plenos = dispensado.
- Braskem (LOA 130 m): 2 TKM >=35 TTE; diurno. Subsea7 (LOA 85,5 m): dispensado com sistemas plenos.
- RBT/CLIP (Caju): tipo 1 EAM (LOA 160 m) dispensado; tipo 2 carga geral (LOA 140 m) = 2 TKM/ASD >=45 TTE.

ESTALEIROS: EISA (Ilha do Governador) = entrada/saida de cascos SOB REBOQUE obrigatorio (DWT <=45 mil t = 2 azimutais; 45-55 mil = 3 azimutais; apoio maritimo com DP = dispensado); assessoria por AMD, transferencia pratico/AMD na isobata de 10 m. Maua (Niteroi): bercos ate LOA 230 m; sem exigencia explicita de rebocador na norma.

FERRO-GUSA (Cais Comercial, cabecos ~95-99 <-> Area de Fundeio no 3, alternativa no 9): comboio ~150 m = 1 rebocador >=45 TTE (conduz a balsa) + 1 >=25 TTE (apoio a atracacao); 3 nos; sem praticagem obrigatoria; visibilidade >=2 milhas nauticas, vento <=20 nos; diurno/noturno; VHF 12. Oportunidade recorrente.

CARGA PERIGOSA (regra geral): DWT >=40.000 t = 2 rebocadores (TKM/ASD) >=60 TTE com cabos passados; DWT 5.000-40.000 t com carga perigosa = 2 rebocadores >=45 TTE.

REGRAS DE OURO: quando o rebocador (ou bow thruster) e obrigatorio, o Comandante NAO pode dispensa-lo. Praticagem obrigatoria para AB >=500 (Lei 9.537/97); os praticos aguardam nos Pontos de Espera de Pratico (PEP). ERU da ZP-15: minimo 58 fainas/quadrimestre por pratico; periodo de escala das 11h01 as 11h do dia seguinte.

FUNDEADOUROS E CONTATO (BG): Areas de Fundeio numeradas (ex.: no 2, no 3, no 9) para espera/staging. Praticagem no VHF canal 12 ('Praticagem Rio', via Atalaia); VHF 16 chamada/seguranca; GNL opera no VHF 13.

USO: cruze o porte do navio (LOA/boca/calado/DWT) e o terminal da manobra programada com as exigencias acima para estimar QUANTOS rebocadores e qual TTE a manobra demanda; aponte as janelas (diurno/mare) que concentram demanda; cite a regra aplicada."""
)


# ===== Conhecimento carregado pelo usuario (upload no chat do KRATOS) =====
KNOWLEDGE_MAX_BYTES = 15 * 1024 * 1024          # 15 MB por arquivo
KNOWLEDGE_INJECT_BUDGET = 12000                 # teto de chars de trechos injetados por resposta
KNOWLEDGE_WHOLE_DOC_MAX = 8000                  # docs menores entram inteiros no contexto
_knowledge_lock = threading.Lock()


def knowledge_dir() -> Path:
    p = user_data_dir(DASHBOARD_USER_ID) / "knowledge"
    p.mkdir(parents=True, exist_ok=True)
    return p


def _knowledge_index_file() -> Path:
    return knowledge_dir() / "index.json"


def _knowledge_text_path(doc_id: str) -> Path:
    return knowledge_dir() / f"{doc_id}.txt"


def _load_knowledge_index() -> list:
    try:
        p = _knowledge_index_file()
        if p.exists():
            data = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return list(data.get("docs") or [])
            if isinstance(data, list):
                return data
    except (OSError, json.JSONDecodeError):
        pass
    return []


def _save_knowledge_index(docs: list):
    _knowledge_index_file().write_text(
        json.dumps({"docs": docs}, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _extract_pdf_text(raw: bytes) -> str:
    import fitz  # PyMuPDF — import tardio; PDF é opcional
    parts = []
    with fitz.open(stream=raw, filetype="pdf") as d:
        for page in d:
            parts.append(page.get_text())
    return "\n".join(parts)


def _summarize_text(text: str, max_len: int = 240) -> str:
    s = " ".join(text.split())
    return s[:max_len] + ("…" if len(s) > max_len else "")


def add_knowledge_document(name: str, ext: str, raw: bytes) -> dict:
    ext = (ext or "").lower().lstrip(".")
    if ext in ("txt", "md", "markdown", "text"):
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", "ignore")
    elif ext == "pdf":
        text = _extract_pdf_text(raw)
    else:
        raise ValueError("formato nao suportado (use PDF, TXT ou MD)")
    text = (text or "").strip()
    if not text:
        raise ValueError("nao foi possivel extrair texto do arquivo")
    doc_id = uuid.uuid4().hex[:12]
    doc = {
        "id": doc_id,
        "name": (str(name or "documento").strip() or "documento")[:120],
        "ext": ext,
        "chars": len(text),
        "summary": _summarize_text(text),
        "addedAt": get_now_iso(),
    }
    with _knowledge_lock:
        _knowledge_text_path(doc_id).write_text(text, encoding="utf-8")
        docs = _load_knowledge_index()
        docs.append(doc)
        _save_knowledge_index(docs)
    return doc


def delete_knowledge_document(doc_id: str) -> bool:
    with _knowledge_lock:
        docs = _load_knowledge_index()
        keep = [d for d in docs if d.get("id") != doc_id]
        if len(keep) == len(docs):
            return False
        _save_knowledge_index(keep)
    try:
        _knowledge_text_path(doc_id).unlink()
    except OSError:
        pass
    return True


def _read_knowledge_text(doc_id: str) -> str:
    try:
        return _knowledge_text_path(doc_id).read_text(encoding="utf-8")
    except OSError:
        return ""


def _extract_relevant_excerpt(text: str, norm: str, q_tokens: list, max_len: int = 3000, window: int = 400) -> str:
    positions = []
    for t in q_tokens:
        i = norm.find(t)
        if i >= 0:
            positions.append(i)
    if not positions:
        return text[:max_len]
    positions.sort()
    parts = []
    used = 0
    last_end = 0
    for p in positions:
        start = max(last_end, p - window)
        end = min(len(text), p + window)
        if start >= end:
            continue
        seg = text[start:end]
        if used + len(seg) > max_len:
            seg = seg[: max(0, max_len - used)]
        if seg:
            parts.append(seg)
            used += len(seg)
            last_end = end
        if used >= max_len:
            break
    return " […] ".join(parts)


def _build_user_knowledge_block(question: str, with_excerpts: bool = True) -> str:
    """Bloco de conhecimento dos documentos carregados pelo usuario: indice (sempre)
    + trechos relevantes a pergunta (somente no chat de texto)."""
    docs = _load_knowledge_index()
    if not docs:
        return ""
    lines = ["CONHECIMENTO CARREGADO PELO USUARIO (documentos que voce deve considerar e citar quando usar):"]
    for d in docs:
        lines.append(f"- {d.get('name')} ({str(d.get('ext','')).upper()}, {d.get('chars',0)} caracteres): {d.get('summary','')}")
    if not with_excerpts:
        return "\n".join(lines)
    q_tokens = [t for t in _normalize_text(question).split() if len(t) > 3]
    budget = KNOWLEDGE_INJECT_BUDGET
    scored = []
    for d in docs:
        text = _read_knowledge_text(d.get("id", ""))
        if not text:
            continue
        norm = _normalize_text(text)
        score = sum(norm.count(t) for t in q_tokens) if q_tokens else 0
        scored.append((score, d, text, norm))
    # docs pequenos entram inteiros; demais por relevancia (score) decrescente
    scored.sort(key=lambda x: (x[1].get("chars", 0) <= KNOWLEDGE_WHOLE_DOC_MAX, x[0]), reverse=True)
    excerpts = []
    for score, d, text, norm in scored:
        if budget <= 0:
            break
        if d.get("chars", 0) <= KNOWLEDGE_WHOLE_DOC_MAX:
            chunk = text[:budget]
            excerpts.append(f"\n=== {d.get('name')} (integra) ===\n{chunk}")
            budget -= len(chunk)
        elif score > 0 and q_tokens:
            chunk = _extract_relevant_excerpt(text, norm, q_tokens, max_len=min(budget, 3000))
            if chunk:
                excerpts.append(f"\n=== {d.get('name')} (trechos) ===\n{chunk}")
                budget -= len(chunk)
    if excerpts:
        lines.append("\nTRECHOS RELEVANTES PARA A PERGUNTA:")
        lines.extend(excerpts)
    return "\n".join(lines)


def _ask_grok_with_context(question: str, context: dict, history: list | None = None) -> str:
    if not GROK_API_KEY:
        return _strategy_fallback_answer(question, context)
    user_knowledge = _build_user_knowledge_block(question, with_excerpts=True)
    messages = [
        {
            "role": "system",
            "content": (
                KRATOS_SYSTEM_PROMPT
                + "\n\n" + KRATOS_APP_GUIDE
                + "\n\n" + KRATOS_NPCP_KNOWLEDGE
                + (("\n\n" + user_knowledge) if user_knowledge else "")
                + "\n\n" + _profile_instruction_block()
                + " " + _assistant_profile_instruction()
            ),
        }
    ]
    # Memoria de conversa: ate as ultimas 8 trocas (papel user/assistant).
    for turn in (history or [])[-8:]:
        role = "assistant" if str(turn.get("role")) == "assistant" else "user"
        text = str(turn.get("content") or "").strip()
        if text:
            messages.append({"role": role, "content": text[:2000]})
    payload = {
        "model": GROK_MODEL,
        "temperature": 0.5,
        "messages": messages + [
            {
                "role": "user",
                "content": (
                    "Contexto operacional JSON:\n"
                    + json.dumps(context, ensure_ascii=False)
                    + "\n\nPergunta:\n"
                    + question
                ),
            },
        ],
    }
    req = urllib.request.Request(
        url="https://api.x.ai/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {GROK_API_KEY}",
            "Content-Type": "application/json",
            "Accept": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=25) as resp:
            body = json.loads(resp.read().decode("utf-8"))
        return (
            (((body.get("choices") or [{}])[0]).get("message") or {}).get("content")
            or "Sem resposta do Grok."
        )
    except (urllib.error.URLError, TimeoutError, json.JSONDecodeError) as exc:
        return _strategy_fallback_answer(question, context) + f"\n\nErro Grok: {exc}"


def build_dashboard_overview_dict():
    ensure_live_worker_started()
    ensure_geofences_loaded()
    ensure_saa_maneuvers_loaded()
    with geofence_lock:
        gfs = list(geofences)
    with saam_geofence_stats_lock:
        tug_snapshot = json.loads(json.dumps(tug_stats_state))
    with saa_maneuvers_lock:
        saa_snapshot = _dedupe_saa_maneuvers(list(saa_maneuvers_list))
    chart = []
    by_m = tug_snapshot.get("byMmsi", {})
    for mmsi in SAAM_MMSI_ABBR:
        row = by_m.get(mmsi, {})
        chart.append(
            {
                "mmsi": mmsi,
                "abbr": row.get("abbr") or SAAM_MMSI_ABBR.get(mmsi, mmsi),
                "name": row.get("name", ""),
                "totalManeuvers": int(row.get("totalManeuvers", 0)),
                "totalHours": round(float(row.get("totalSeconds", 0)) / 3600.0, 2),
                "totalNauticalMiles": round(float(row.get("totalNauticalMiles", 0)), 2),
            }
        )
    market_share = _market_share_rows(saa_snapshot)
    market_share_windows = _market_share_windows(saa_snapshot)
    competitors = _competitor_status_rows(gfs)
    schedule_changes = _build_schedule_changes(saa_snapshot)
    return {
        "ok": True,
        "userId": DASHBOARD_USER_ID,
        "geofences": gfs,
        "geofenceStatus": _dashboard_geofence_status_rows(),
        "occupancy": build_geofence_occupancy(),
        "tugStats": tug_snapshot,
        "tugChart": chart,
        "saaManeuvers": saa_snapshot,
        "marketShare": market_share,
        "marketShareWindows": market_share_windows,
        "competitors": competitors,
        "scheduleChanges": schedule_changes,
    }


@app.get("/api/dashboard/overview")
def dashboard_overview():
    return build_dashboard_overview_dict()


@app.get("/dashboard/api/overview")
def dashboard_overview_under_dashboard_path():
    """Mesmo JSON que /api/dashboard/overview (registado antes das rotas /api/geofences/*)."""
    return build_dashboard_overview_dict()


@app.get("/api/dashboard/saa-maneuvers")
def get_saa_maneuvers():
    with saa_maneuvers_lock:
        return {"ok": True, "items": _dedupe_saa_maneuvers(list(saa_maneuvers_list))}


@app.post("/api/dashboard/saa-maneuvers")
async def append_saa_maneuver(request: Request):
    payload = await request.json()
    item = {
        "id": str(uuid.uuid4()),
        "vesselName": payload.get("vesselName", "").strip() or "—",
        "berthName": payload.get("berthName", "").strip() or "—",
        "empRb": (payload.get("empRb") or "SAA").strip(),
        "status": payload.get("status", "").strip() or "—",
        "note": (payload.get("note") or "").strip(),
        "recordedAt": get_now_iso(),
        **_saa_dims_from_payload(payload),
    }
    with saa_maneuvers_lock:
        saa_maneuvers_list.insert(0, item)
        save_saa_maneuvers()
    return {"ok": True, "item": item}


@app.post("/api/dashboard/strategy-assistant")
async def strategy_assistant(request: Request):
    payload = await request.json()
    question = str(payload.get("question") or "").strip()
    action = str(payload.get("action") or "ask").strip().lower()
    learn_note = str(payload.get("learnNote") or "").strip()
    raw_history = payload.get("history")
    history = []
    if isinstance(raw_history, list):
        for turn in raw_history:
            if isinstance(turn, dict) and turn.get("content"):
                history.append({"role": turn.get("role"), "content": turn.get("content")})
    if learn_note:
        append_strategy_memory(learn_note, author="user")
    if not question:
        return JSONResponse({"ok": False, "error": "question obrigatoria"}, status_code=400)
    context = _strategy_context_dict()
    map_view = payload.get("mapView")
    if isinstance(map_view, dict):
        context["userMapView"] = map_view
    if action == "report":
        question = (
            "Gere um relatorio executivo com: panorama operacional, market share (hoje/7d/30d), "
            "concorrentes manobrando, riscos meteoceanicos (vento/mare) e recomendacoes de alocacao. "
            + question
        )
    elif action == "insights":
        question = (
            "Gere insights acionaveis e hipoteses de ganho competitivo com base no contexto. "
            + question
        )
    answer = await asyncio.to_thread(_ask_grok_with_context, question, context, history)
    answer = _extract_profile_tags(answer)
    append_strategy_memory(f"Pergunta: {question}\nResposta: {answer[:1200]}", author="assistant")
    # Telemetria/auditoria: registra a interação (sem PII além do que o usuário digita).
    low = answer.lower()
    no_data = any(s in low for s in ["sem dado", "não há dado", "nao ha dado", "não tenho", "nao tenho", "sem grok", "erro grok"])
    log_kratos_event("chat", {
        "action": action,
        "mode": "grok" if GROK_API_KEY else "local",
        "question": question[:500],
        "answerPreview": answer[:500],
        "answerChars": len(answer),
        "historyTurns": len(history),
        "noData": no_data,
        "viaVoice": bool(payload.get("viaVoice")),
        "userName": (load_kratos_profile() or {}).get("name"),
    })
    return {"ok": True, "answer": answer, "context": context}


@app.post("/api/kratos/feedback")
async def kratos_feedback(request: Request):
    """Feedback leve do usuário sobre uma resposta do KRATOS ('isso foi útil?')."""
    payload = await request.json()
    useful = bool(payload.get("useful"))
    log_kratos_event("feedback", {
        "useful": useful,
        "acted": bool(payload.get("acted")),
        "question": str(payload.get("question") or "")[:300],
        "answerPreview": str(payload.get("answerPreview") or "")[:300],
        "userName": (load_kratos_profile() or {}).get("name"),
    })
    return {"ok": True}


@app.post("/dashboard/api/kratos/feedback")
async def kratos_feedback_dash(request: Request):
    return await kratos_feedback(request)


# ----- Conhecimento carregado no chat (upload via JSON+base64, sem multipart) -----
@app.post("/api/kratos/knowledge")
async def kratos_knowledge_upload(request: Request):
    import base64
    data = await _read_json_body(request)
    name = str(data.get("name") or "documento")
    ext = str(data.get("ext") or "").lower().lstrip(".")
    b64 = data.get("contentBase64")
    if not b64:
        return JSONResponse({"ok": False, "error": "arquivo vazio"}, status_code=400)
    try:
        raw = base64.b64decode(str(b64).split(",")[-1])
    except (ValueError, TypeError):
        return JSONResponse({"ok": False, "error": "conteudo invalido"}, status_code=400)
    if len(raw) > KNOWLEDGE_MAX_BYTES:
        return JSONResponse({"ok": False, "error": "arquivo acima de 15 MB"}, status_code=400)
    try:
        doc = await asyncio.to_thread(add_knowledge_document, name, ext, raw)
    except ImportError:
        return JSONResponse(
            {"ok": False, "error": "Leitura de PDF indisponivel no servidor (PyMuPDF nao instalado). Envie TXT ou MD."},
            status_code=400,
        )
    except ValueError as exc:
        return JSONResponse({"ok": False, "error": str(exc)}, status_code=400)
    log_kratos_event("knowledge", {"name": doc["name"], "ext": doc["ext"], "chars": doc["chars"]})
    return {"ok": True, "doc": doc}


@app.get("/api/kratos/knowledge")
async def kratos_knowledge_list():
    return {"ok": True, "docs": _load_knowledge_index()}


@app.post("/api/kratos/knowledge/delete")
async def kratos_knowledge_delete(request: Request):
    data = await _read_json_body(request)
    ok = delete_knowledge_document(str(data.get("id") or "").strip())
    return {"ok": bool(ok)}


@app.post("/dashboard/api/kratos/knowledge")
async def kratos_knowledge_upload_dash(request: Request):
    return await kratos_knowledge_upload(request)


@app.get("/dashboard/api/kratos/knowledge")
async def kratos_knowledge_list_dash():
    return await kratos_knowledge_list()


@app.post("/dashboard/api/kratos/knowledge/delete")
async def kratos_knowledge_delete_dash(request: Request):
    return await kratos_knowledge_delete(request)


# ===== Relatório executivo formatado (PDF / DOCX) + compartilhamento =====
def _build_report_payload() -> dict:
    """Reúne os dados reais do cenário para o relatório executivo."""
    ov = build_dashboard_overview_dict()
    ctx = _strategy_context_dict()
    metocean = ctx.get("metocean") or {}
    msw = ov.get("marketShareWindows") or {}
    tug_chart = ov.get("tugChart") or []
    competitors = ov.get("competitors") or []
    saa = ov.get("saaManeuvers") or []
    ophours = _saam_operating_hours_rows()
    try:
        insights = _build_kratos_insights()
    except Exception:
        insights = []
    saam_active = [t for t in (ctx.get("saamTugs") or []) if t.get("latitude") is not None]
    comp_inside = [c for c in competitors if c.get("insideManeuverGeofence")]
    return {
        "generatedAt": get_now_iso(),
        "panorama": {
            "saamTugs": len(saam_active),
            "scheduledManeuvers": len(saa),
            "competitorsManeuvering": len(comp_inside),
        },
        "marketShare": msw,
        "tugChart": tug_chart,
        "operatingHours": ophours,
        "competitors": competitors,
        "metocean": metocean,
        "insights": insights,
    }


def _fmt_metocean_line(m: dict) -> str:
    parts = []
    if isinstance(m.get("temperatureC"), (int, float)):
        parts.append(f"Temperatura {m['temperatureC']:.0f} °C")
    if isinstance(m.get("windSpeedKn"), (int, float)):
        wd = m.get("windDirection") or ""
        parts.append(f"Vento {m['windSpeedKn']:.0f} kn {wd}".strip())
    if isinstance(m.get("tideLevelM"), (int, float)):
        parts.append(f"Maré {m['tideLevelM']:.2f} m ({m.get('tideTrend','')})".strip())
    if isinstance(m.get("currentSpeedKn"), (int, float)):
        cd = m.get("currentDirection") or ""
        parts.append(f"Corrente {m['currentSpeedKn']:.2f} kn {cd}".strip())
    return " · ".join(parts) or "Sem dados meteoceânicos."


def _report_filename(ext: str) -> str:
    return "KRATOS_Relatorio_" + datetime.now().strftime("%Y%m%d_%H%M") + "." + ext


def _render_report_pdf(p: dict) -> bytes:
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable)
    from reportlab.lib.styles import ParagraphStyle

    ACC = HexColor("#0e7bb8"); DARK = HexColor("#0b2237"); MUT = HexColor("#4a6076")
    body = ParagraphStyle("b", fontName="Helvetica", fontSize=9.8, textColor=DARK, leading=13.5)
    h1 = ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=17, textColor=DARK, spaceAfter=2)
    sub = ParagraphStyle("sub", fontName="Helvetica", fontSize=9.5, textColor=MUT, spaceAfter=10)
    h2 = ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=12, textColor=ACC, spaceBefore=11, spaceAfter=5)

    def P(t, s=body): return Paragraph(t, s)

    def tbl(data, widths):
        t = Table(data, colWidths=widths)
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"), ("FONTSIZE", (0, 0), (-1, -1), 8.6),
            ("TEXTCOLOR", (0, 0), (-1, -1), DARK), ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#c7d4e0")),
            ("TOPPADDING", (0, 0), (-1, -1), 3.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
            ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#e8f1f8")), ("TEXTCOLOR", (0, 0), (-1, 0), ACC),
        ]))
        return t

    def hdr(c, d):
        c.saveState(); c.setFillColor(DARK); c.rect(0, A4[1] - 16 * mm, A4[0], 16 * mm, fill=1, stroke=0)
        c.setFillColor(HexColor("#35c8ff")); c.setFont("Helvetica-Bold", 11); c.drawString(16 * mm, A4[1] - 10.5 * mm, "KRATOS")
        c.setFillColor(HexColor("#9ab6d4")); c.setFont("Helvetica", 8)
        c.drawString(38 * mm, A4[1] - 10.5 * mm, "Inteligência Naval Estratégica · Relatório operacional")
        c.setFillColor(MUT); c.setFont("Helvetica", 7.5)
        c.drawCentredString(A4[0] / 2, 8 * mm, f"Autor: Jossian Brito · Baía de Guanabara · Página {c.getPageNumber()}")
        c.restoreState()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm,
                            topMargin=24 * mm, bottomMargin=16 * mm,
                            title="KRATOS — Relatório operacional", author="Jossian Brito")
    when = datetime.now().strftime("%d/%m/%Y %H:%M")
    E = [P("Relatório operacional — KRATOS", h1), P(f"Gerado em {when} · Baía de Guanabara", sub),
         HRFlowable(width="100%", thickness=1, color=ACC), Spacer(1, 6)]

    pan = p.get("panorama", {})
    E.append(P("Panorama", h2))
    E.append(P(f"Rebocadores SAAM com sinal: <b>{pan.get('saamTugs',0)}</b> · "
               f"Manobras programadas: <b>{pan.get('scheduledManeuvers',0)}</b> · "
               f"Concorrentes manobrando agora: <b>{pan.get('competitorsManeuvering',0)}</b>."))

    msw = p.get("marketShare") or {}
    def share_rows(win):
        rows = (win or {}).get("rows", []) if isinstance(win, dict) else (win or [])
        return [[r.get("empRb", "-"), str(r.get("count", 0)), f"{r.get('sharePct', 0)}%"] for r in rows[:6]] or [["—", "—", "—"]]
    E.append(P("Market share (Praticagem-RJ)", h2))
    for label, key in [("Hoje", "today"), ("7 dias", "last7d"), ("30 dias", "last30d")]:
        E.append(P(f"<b>{label}</b>", body))
        E.append(tbl([["Empresa", "Manobras", "Participação"]] + share_rows(msw.get(key)), [60 * mm, 50 * mm, 58 * mm]))
        E.append(Spacer(1, 4))

    tc = p.get("tugChart") or []
    oph = {r["mmsi"]: r for r in (p.get("operatingHours") or [])}
    if tc:
        E.append(P("Frota SAAM (atividade e fadiga)", h2))
        rows = [["Rebocador", "Manobras", "Horas em geofence", "Milhas náut.", "Operação hoje"]]
        for t in tc:
            op = oph.get(t.get("mmsi"), {})
            oph_txt = f"{op.get('operatingHours', 0)} h ({op.get('status','-')})" if op else "—"
            rows.append([t.get("name") or t.get("abbr") or t.get("mmsi"),
                         str(t.get("totalManeuvers", 0)), str(t.get("totalHours", 0)),
                         str(t.get("totalNauticalMiles", 0)), oph_txt])
        E.append(tbl(rows, [44 * mm, 24 * mm, 34 * mm, 26 * mm, 40 * mm]))

    E.append(P("Condições meteoceânicas", h2))
    E.append(P(_fmt_metocean_line(p.get("metocean") or {})))

    ins = p.get("insights") or []
    if ins:
        E.append(P("Leitura do KRATOS", h2))
        for s in ins[:10]:
            E.append(P("• " + str(s)))

    doc.build(E, onFirstPage=hdr, onLaterPages=hdr)
    return buf.getvalue()


def _render_report_docx(p: dict) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ACC = RGBColor(0x0E, 0x7B, 0xB8); DARK = RGBColor(0x0B, 0x22, 0x37); MUT = RGBColor(0x4A, 0x60, 0x76)
    doc = Document()
    for s in doc.sections:
        s.top_margin = Mm(18); s.bottom_margin = Mm(16); s.left_margin = Mm(16); s.right_margin = Mm(16)
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"; normal.size = Pt(10)

    def heading(txt, size=16, color=DARK):
        pp = doc.add_paragraph(); r = pp.add_run(txt); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
        return pp

    def h2(txt):
        pp = doc.add_paragraph(); pp.space_before = Pt(8); r = pp.add_run(txt); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = ACC
        return pp

    when = datetime.now().strftime("%d/%m/%Y %H:%M")
    heading("Relatório operacional — KRATOS")
    sp = doc.add_paragraph(); r = sp.add_run(f"Gerado em {when} · Baía de Guanabara · Autor: Jossian Brito"); r.font.size = Pt(9.5); r.font.color.rgb = MUT

    pan = p.get("panorama", {})
    h2("Panorama")
    doc.add_paragraph(f"Rebocadores SAAM com sinal: {pan.get('saamTugs',0)} · "
                      f"Manobras programadas: {pan.get('scheduledManeuvers',0)} · "
                      f"Concorrentes manobrando agora: {pan.get('competitorsManeuvering',0)}.")

    msw = p.get("marketShare") or {}
    h2("Market share (Praticagem-RJ)")
    for label, key in [("Hoje", "today"), ("7 dias", "last7d"), ("30 dias", "last30d")]:
        doc.add_paragraph().add_run(label).bold = True
        win = msw.get(key) or {}
        rows = ((win.get("rows", []) if isinstance(win, dict) else win) or [])[:6]
        tb = doc.add_table(rows=1, cols=3); tb.style = "Light Grid Accent 1"
        hc = tb.rows[0].cells; hc[0].text = "Empresa"; hc[1].text = "Manobras"; hc[2].text = "Participação"
        for r in rows:
            c = tb.add_row().cells
            c[0].text = str(r.get("empRb", "-")); c[1].text = str(r.get("count", 0)); c[2].text = f"{r.get('sharePct',0)}%"

    tc = p.get("tugChart") or []
    oph = {r["mmsi"]: r for r in (p.get("operatingHours") or [])}
    if tc:
        h2("Frota SAAM (atividade e fadiga)")
        tb = doc.add_table(rows=1, cols=5); tb.style = "Light Grid Accent 1"
        hc = tb.rows[0].cells
        for i, t in enumerate(["Rebocador", "Manobras", "Horas geofence", "Milhas náut.", "Operação hoje"]):
            hc[i].text = t
        for t in tc:
            op = oph.get(t.get("mmsi"), {})
            c = tb.add_row().cells
            c[0].text = str(t.get("name") or t.get("abbr") or t.get("mmsi"))
            c[1].text = str(t.get("totalManeuvers", 0)); c[2].text = str(t.get("totalHours", 0))
            c[3].text = str(t.get("totalNauticalMiles", 0))
            c[4].text = (f"{op.get('operatingHours',0)} h ({op.get('status','-')})" if op else "—")

    h2("Condições meteoceânicas")
    doc.add_paragraph(_fmt_metocean_line(p.get("metocean") or {}))

    ins = p.get("insights") or []
    if ins:
        h2("Leitura do KRATOS")
        for s in ins[:10]:
            doc.add_paragraph(str(s), style="List Bullet")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


@app.post("/api/kratos/report-file")
async def kratos_report_file(request: Request):
    data = await _read_json_body(request)
    fmt = str(data.get("format") or "pdf").lower().strip()
    if fmt not in ("pdf", "docx"):
        return JSONResponse({"ok": False, "error": "formato invalido (pdf|docx)"}, status_code=400)
    try:
        payload = await asyncio.to_thread(_build_report_payload)
        if fmt == "pdf":
            blob = await asyncio.to_thread(_render_report_pdf, payload)
            media = "application/pdf"
        else:
            blob = await asyncio.to_thread(_render_report_docx, payload)
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except ImportError as exc:
        lib = "python-docx" if fmt == "docx" else "reportlab"
        return JSONResponse({"ok": False, "error": f"biblioteca {lib} nao instalada no servidor ({exc})."}, status_code=400)
    except Exception as exc:
        return JSONResponse({"ok": False, "error": f"falha ao gerar relatorio: {exc}"}, status_code=500)
    log_kratos_event("report", {"format": fmt, "bytes": len(blob)})
    fname = _report_filename(fmt)
    return Response(content=blob, media_type=media,
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@app.post("/dashboard/api/kratos/report-file")
async def kratos_report_file_dash(request: Request):
    return await kratos_report_file(request)


# ===== Página do administrador (métricas, eficácia, auditoria) =====

def _admin_authorized(request: Request) -> bool:
    if not ADMIN_TOKEN:
        return False  # sem token configurado, admin fica indisponível (seguro por padrão)
    given = (
        request.headers.get("X-Admin-Token")
        or request.query_params.get("token")
        or ""
    ).strip()
    return given == ADMIN_TOKEN


def _build_admin_metrics(events: list) -> dict:
    from collections import Counter
    now = time.time()
    day = 86400
    chats = [e for e in events if e.get("type") == "chat"]
    fb = [e for e in events if e.get("type") == "feedback"]
    voice = [e for e in events if e.get("type") == "voice_session"]
    errors = [e for e in events if e.get("type") == "error"]

    def since(items, secs):
        return [e for e in items if (now - (e.get("ts") or 0)) <= secs]

    chats_24h = since(chats, day)
    chats_7d = since(chats, 7 * day)
    no_data = [e for e in chats if e.get("noData")]
    useful_yes = [e for e in fb if e.get("useful")]
    useful_no = [e for e in fb if not e.get("useful")]
    acted = [e for e in fb if e.get("acted")]

    # Uso por dia (últimos 14 dias) para o gráfico de evolução.
    by_day = Counter()
    for e in chats:
        ts = e.get("ts")
        if ts:
            d = datetime.fromtimestamp(ts).strftime("%Y-%m-%d")
            by_day[d] += 1
    daily = [{"day": d, "chats": by_day[d]} for d in sorted(by_day.keys())[-14:]]

    # Temas mais perguntados (heurística por palavra-chave).
    topics = Counter()
    kw = {
        "concorrente": ["concorrent", "wil", "cam"],
        "manobra/programação": ["manobra", "pob", "programa"],
        "frota SAAM": ["saam", "rebocador", "frota"],
        "maré/vento": ["maré", "mare", "vento", "tempo"],
        "geofence/corredor": ["geofence", "berço", "berco", "corredor"],
        "market share": ["market", "share", "participa"],
        "distância/ETA": ["distância", "distancia", "eta", "milha"],
    }
    for e in chats:
        q = _normalize_text(e.get("question") or "")
        for topic, terms in kw.items():
            if any(_normalize_text(t) in q for t in terms):
                topics[topic] += 1

    total_fb = len(useful_yes) + len(useful_no)
    return {
        "totals": {
            "chats": len(chats),
            "chats24h": len(chats_24h),
            "chats7d": len(chats_7d),
            "voiceSessions": len(voice),
            "feedbacks": len(fb),
            "errors": len(errors),
        },
        "efficacy": {
            "usefulYes": len(useful_yes),
            "usefulNo": len(useful_no),
            "usefulPct": round(100.0 * len(useful_yes) / total_fb, 1) if total_fb else None,
            "actedCount": len(acted),
            "feedbackRatePct": round(100.0 * len(fb) / len(chats), 1) if chats else None,
        },
        "quality": {
            "noDataCount": len(no_data),
            "noDataPct": round(100.0 * len(no_data) / len(chats), 1) if chats else None,
            "errorCount": len(errors),
            "lastErrors": errors[-10:][::-1],
        },
        "dailyChats": daily,
        "topics": [{"topic": t, "count": c} for t, c in topics.most_common(8)],
    }


@app.get("/api/admin/overview")
async def admin_overview(request: Request):
    if not _admin_authorized(request):
        return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    events = load_kratos_events(limit=KRATOS_EVENTS_MAX)
    metrics = _build_admin_metrics(events)
    # Correlação com resultado operacional (snapshot atual).
    try:
        ov = build_dashboard_overview_dict()
        correlation = {
            "marketShare": ov.get("marketShare"),
            "tugChart": ov.get("tugChart"),
            "scheduledManeuverTotal": len(ov.get("saaManeuvers") or []),
        }
    except Exception:
        correlation = {}
    profile = load_kratos_profile()
    return {
        "ok": True,
        "generatedAt": get_now_iso(),
        "metrics": metrics,
        "correlation": correlation,
        "userProfile": profile,
        "voiceConfigured": bool(GROK_API_KEY),
    }


@app.get("/api/admin/conversations")
async def admin_conversations(request: Request, limit: int = 200):
    """Log de conversas (auditoria): perguntas/respostas e feedback recentes."""
    if not _admin_authorized(request):
        return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    events = load_kratos_events(limit=KRATOS_EVENTS_MAX)
    convo = [e for e in events if e.get("type") in ("chat", "feedback", "voice_session", "error")]
    return {"ok": True, "events": convo[-limit:][::-1]}


async def _read_json_body(request: Request) -> dict:
    try:
        data = await request.json()
        return data if isinstance(data, dict) else {}
    except (json.JSONDecodeError, ValueError):
        return {}


# ----- Controle de acesso por convite: APIs públicas (porta de entrada) -----
@app.get("/api/access/status")
async def access_status(request: Request):
    token = (
        request.cookies.get(ACCESS_COOKIE_NAME)
        or request.query_params.get("access")
        or ""
    ).strip()
    has = (
        (not ACCESS_CONTROL_ON)
        or (bool(ADMIN_TOKEN) and token == ADMIN_TOKEN)
        or access_token_is_valid(token)
    )
    return {"ok": True, "accessControl": ACCESS_CONTROL_ON, "hasValidAccess": bool(has)}


@app.post("/api/access/validate")
async def access_validate(request: Request):
    data = await _read_json_body(request)
    token = str(data.get("token") or "").strip()
    valid = (bool(ADMIN_TOKEN) and token == ADMIN_TOKEN) or access_token_is_valid(token)
    if not valid:
        return JSONResponse({"ok": True, "valid": False})
    resp = JSONResponse({"ok": True, "valid": True})
    secure = request.headers.get("x-forwarded-proto", request.url.scheme) == "https"
    resp.set_cookie(
        ACCESS_COOKIE_NAME, token, max_age=ACCESS_COOKIE_MAX_AGE,
        httponly=True, samesite="lax", secure=secure,
    )
    return resp


@app.post("/api/access/logout")
async def access_logout():
    resp = JSONResponse({"ok": True})
    resp.delete_cookie(ACCESS_COOKIE_NAME)
    return resp


@app.get("/entrar")
async def access_gate_page():
    return FileResponse(FRONTEND_DIR / "entrar.html", media_type="text/html")


# ----- Gestão de convites: APIs do administrador (exigem ADMIN_TOKEN) -----
@app.get("/api/admin/invites")
async def admin_invites_list(request: Request):
    if not _admin_authorized(request):
        return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    now = time.time()
    invites = [_public_invite_view(inv, now) for inv in _load_access_invites()]
    invites.sort(key=lambda x: x.get("createdAt") or 0, reverse=True)
    return {"ok": True, "accessControl": ACCESS_CONTROL_ON, "invites": invites}


@app.post("/api/admin/invites")
async def admin_invites_create(request: Request):
    if not _admin_authorized(request):
        return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    data = await _read_json_body(request)
    inv = create_access_invite(data.get("label"), data.get("expiresInDays"))
    return {"ok": True, "invite": _public_invite_view(inv)}


@app.post("/api/admin/invites/revoke")
async def admin_invites_revoke(request: Request):
    if not _admin_authorized(request):
        return JSONResponse({"ok": False, "error": "não autorizado"}, status_code=401)
    data = await _read_json_body(request)
    ok = revoke_access_invite(str(data.get("token") or "").strip())
    return {"ok": bool(ok)}


@app.get("/admin")
async def admin_page(request: Request):
    return FileResponse(FRONTEND_DIR / "admin.html", media_type="text/html")


@app.get("/admin/")
async def admin_page_slash(request: Request):
    return FileResponse(FRONTEND_DIR / "admin.html", media_type="text/html")


@app.post("/dashboard/api/strategy-assistant")
async def strategy_assistant_under_dashboard_path(request: Request):
    """Mesmo comportamento de /api/dashboard/strategy-assistant para subpath /dashboard."""
    return await strategy_assistant(request)


# ===== Gestão de frota (inclusão/substituição/remoção de rebocadores) =====
VALID_FLEETS = {"SAAM", "WIL", "CAM"}


@app.get("/api/fleet")
def get_fleet():
    ensure_fleet_config_loaded()
    return {"ok": True, "fleet": _current_fleet_config()}


@app.post("/api/fleet/tug")
async def upsert_fleet_tug(request: Request):
    """Inclui ou substitui um rebocador na frota.

    Body: {fleet: "SAAM"|"WIL"|"CAM", mmsi, name, abbr?, replaceMmsi?}
    Se replaceMmsi vier, o rebocador antigo é removido (substituição).
    """
    ensure_fleet_config_loaded()
    payload = await request.json()
    fleet = str(payload.get("fleet") or "").strip().upper()
    mmsi = "".join(ch for ch in str(payload.get("mmsi") or "") if ch.isdigit())
    name = str(payload.get("name") or "").strip().upper()
    abbr = str(payload.get("abbr") or "").strip().upper()
    replace_mmsi = "".join(ch for ch in str(payload.get("replaceMmsi") or "") if ch.isdigit())
    if fleet not in VALID_FLEETS:
        return JSONResponse({"ok": False, "error": f"fleet deve ser uma de {sorted(VALID_FLEETS)}"}, status_code=400)
    if not (7 <= len(mmsi) <= 9):
        return JSONResponse({"ok": False, "error": "MMSI inválido (use 7 a 9 dígitos)."}, status_code=400)
    if not name:
        return JSONResponse({"ok": False, "error": "Informe o nome do rebocador."}, status_code=400)

    cfg = _current_fleet_config()
    saam = [t for t in cfg["saam"] if t["mmsi"] not in (mmsi, replace_mmsi)]
    comps = {
        c: [t for t in tugs if t["mmsi"] not in (mmsi, replace_mmsi)]
        for c, tugs in cfg["competitors"].items()
    }
    if fleet == "SAAM":
        saam.append({"mmsi": mmsi, "name": name, "abbr": abbr or mmsi[-2:]})
    else:
        comps.setdefault(fleet, []).append({"mmsi": mmsi, "name": name})
    _rebuild_fleet_structures(saam, comps)
    save_fleet_config()
    return {"ok": True, "fleet": _current_fleet_config()}


@app.post("/api/fleet/tug/remove")
async def remove_fleet_tug(request: Request):
    """Remove um rebocador da frota. Body: {mmsi}."""
    ensure_fleet_config_loaded()
    payload = await request.json()
    mmsi = "".join(ch for ch in str(payload.get("mmsi") or "") if ch.isdigit())
    if not mmsi:
        return JSONResponse({"ok": False, "error": "MMSI obrigatório."}, status_code=400)
    cfg = _current_fleet_config()
    saam = [t for t in cfg["saam"] if t["mmsi"] != mmsi]
    comps = {c: [t for t in tugs if t["mmsi"] != mmsi] for c, tugs in cfg["competitors"].items()}
    _rebuild_fleet_structures(saam, comps)
    save_fleet_config()
    return {"ok": True, "fleet": _current_fleet_config()}


# ===== KRATOS Voz ao vivo (xAI Realtime Voice API) =====
KRATOS_VOICE_ID = (os.getenv("XAI_VOICE") or "leo").strip() or "leo"
KRATOS_REALTIME_MODEL = (os.getenv("XAI_REALTIME_MODEL") or "grok-voice-latest").strip()


def _kratos_voice_instructions(map_view: dict | None = None) -> str:
    """Instruções do agente de voz: persona + perfil do usuário + contexto operacional
    completo (insights, embarcações, geofences, distâncias) e visão do mapa."""
    try:
        insights = _build_kratos_insights()
    except Exception:
        insights = []
    contexto = "\n".join(f"- {s}" for s in insights[:10])

    # Embarcações em formato compacto (1 linha cada) para perguntas tipo
    # "que navio está fundeado em tal ponto".
    vessels_lines = []
    try:
        for v in _compact_vessels_overview(limit=120):
            fleet = f"/{v['fleet']}" if v.get("fleet") else ""
            geo = f" em {','.join(v['geofences'][:2])}" if v.get("geofences") else ""
            estado = "movendo" if v.get("moving") else "parado"
            vessels_lines.append(
                f"{v['name']}{fleet} ({v['category']}, {v['lat']},{v['lon']}, {v['sogKn']}kn {estado}){geo}"
            )
    except Exception:
        pass
    vessels_block = "; ".join(vessels_lines)

    geo_lines = []
    try:
        for g in _geofences_summary():
            dims = f", raio {g['radiusMeters']}m" if g.get("radiusMeters") else (
                f", ~{g.get('approxSpanNm')}nm" if g.get("approxSpanNm") is not None else ""
            )
            geo_lines.append(
                f"{g['name']} ({g['type']}, centro {g.get('centerLat')},{g.get('centerLon')}{dims})"
            )
    except Exception:
        pass

    dist_lines = []
    try:
        ctx = _strategy_context_dict()
        for r in (ctx.get("maneuverDistances") or [])[:6]:
            tops = ", ".join(
                f"{d['tug']}/{d['fleet']} {d['distanceNmStraight']} milhas nauticas (linha reta)"
                for d in r.get("tugDistances", [])[:4]
            )
            dist_lines.append(f"{r['vessel']} (POB {r['pob']}): {tops}")
    except Exception:
        pass

    mapa = ""
    if isinstance(map_view, dict):
        c = map_view.get("center") or {}
        mapa = (
            "\n\nVISAO ATUAL DO MAPA DO USUARIO: centro "
            f"{c.get('lat')},{c.get('lng')}, zoom {map_view.get('zoom')}"
        )
        b = map_view.get("bounds") or {}
        if b:
            mapa += f", area visivel de {b.get('south')},{b.get('west')} a {b.get('north')},{b.get('east')}"

    voice_knowledge = _build_user_knowledge_block("", with_excerpts=False)
    return (
        KRATOS_SYSTEM_PROMPT
        + "\n\n" + KRATOS_APP_GUIDE
        + "\n\n" + KRATOS_NPCP_KNOWLEDGE
        + (("\n\n" + voice_knowledge) if voice_knowledge else "")
        + "\n\n" + _profile_instruction_block()
        + "\n\nVOZ: fale em portugues do Brasil, tom calmo, tecnico e com autoridade, "
        "como um parceiro operacional ao lado da equipe. Respostas CURTAS (1 a 2 frases por "
        "turno), so o que foi perguntado — NUNCA leia listas nem despeje o contexto abaixo; "
        "ele e apenas para consulta. Se houver muito a dizer, de o essencial e pergunte se quer "
        "mais. Se o usuario so cumprimentar, responda a saudacao, pergunte o nome (se nao souber) "
        "e fique a disposicao, sem disparar dados. Se interromper, pare e ouca. Nunca invente "
        "dados.\n\n"
        "CONTEXTO OPERACIONAL (apenas para consulta, NAO leia em voz alta):\n" + contexto
        + ("\n\nGEOFENCES (demarcacao): " + "; ".join(geo_lines) if geo_lines else "")
        + ("\n\nDISTANCIAS REBOCADOR->NAVIO (proximas manobras): " + " | ".join(dist_lines) if dist_lines else "")
        + ("\n\nEMBARCACOES NO RADAR: " + vessels_block if vessels_block else "")
        + mapa
    )


def _mint_realtime_client_secret() -> dict:
    """Cunha um token efemero no xAI para o navegador conectar ao Realtime.

    A XAI_API_KEY fica apenas no servidor; o browser usa o token via
    subprotocolo `xai-client-secret.<token>`.
    """
    req = urllib.request.Request(
        url="https://api.x.ai/v1/realtime/client_secrets",
        data=json.dumps({"expires_after": {"seconds": 300}}).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {GROK_API_KEY}",
            "Content-Type": "application/json",
            "Accept": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=15) as resp:
        return json.loads(resp.read().decode("utf-8"))


@app.post("/api/kratos/voice-session")
async def kratos_voice_session(request: Request = None):
    """Inicia uma sessão de voz ao vivo: token efêmero + configuração da sessão.

    Aceita corpo opcional {"mapView": {center:{lat,lng}, zoom, bounds:{...}}}
    para o KRATOS saber o que o usuário está olhando no mapa.
    """
    map_view = None
    if request is not None:
        try:
            body = await request.json()
            if isinstance(body, dict) and isinstance(body.get("mapView"), dict):
                map_view = body["mapView"]
        except Exception:
            map_view = None
    if not GROK_API_KEY:
        return JSONResponse(
            {"ok": False, "error": "XAI_API_KEY não configurada no servidor."},
            status_code=503,
        )
    try:
        secret = await asyncio.to_thread(_mint_realtime_client_secret)
    except urllib.error.HTTPError as exc:
        detail = ""
        try:
            detail = exc.read().decode("utf-8")[:300]
        except Exception:
            pass
        log_kratos_event("error", {"where": "voice_session", "code": exc.code, "detail": detail[:200]})
        return JSONResponse(
            {
                "ok": False,
                "error": f"xAI recusou a criação do token ({exc.code}). "
                "Verifique se a chave tem o endpoint Voice habilitado no console.x.ai.",
                "detail": detail,
            },
            status_code=502,
        )
    except Exception as exc:
        log_kratos_event("error", {"where": "voice_session", "detail": str(exc)[:200]})
        return JSONResponse({"ok": False, "error": f"Falha ao cunhar token: {exc}"}, status_code=502)
    token = secret.get("value") or ""
    if not token:
        return JSONResponse({"ok": False, "error": "Resposta do xAI sem token."}, status_code=502)
    log_kratos_event("voice_session", {"voice": KRATOS_VOICE_ID, "hasMapView": bool(map_view)})
    return {
        "ok": True,
        "token": token,
        "expiresAt": secret.get("expires_at"),
        "model": KRATOS_REALTIME_MODEL,
        "voice": KRATOS_VOICE_ID,
        "instructions": _kratos_voice_instructions(map_view),
    }


@app.post("/dashboard/api/kratos/voice-session")
async def kratos_voice_session_under_dashboard_path(request: Request = None):
    """Mesmo comportamento para o subpath /dashboard."""
    return await kratos_voice_session(request)


def _wind_label(metocean: dict) -> str:
    spd = metocean.get("windSpeedKmh")
    if not isinstance(spd, (int, float)):
        return ""
    kn = spd / 1.852
    if kn >= 22:
        nivel = "forte"
    elif kn >= 12:
        nivel = "moderado"
    else:
        nivel = "fraco"
    return f"{kn:.0f} kn ({nivel})"


def _build_kratos_insights() -> list[str]:
    """Insights curtos gerados por regras a partir do contexto operacional real.

    Rápido e sempre disponível (sem chamada externa); alimenta a caixa
    datilografada do mapa.
    """
    ctx = _strategy_context_dict()
    insights: list[str] = []

    saam = ctx.get("saamTugs") or []
    comps = ctx.get("competitors") or []
    maneuvers = ctx.get("scheduledManeuvers") or []
    m_total = int(ctx.get("scheduledManeuverTotal", len(maneuvers)))
    simult = ctx.get("simultaneousManeuvers") or []
    changes = ctx.get("scheduleChanges") or {}
    metocean = ctx.get("metocean") or {}
    market = (ctx.get("marketShare") or {}).get("rows") or []

    # Rebocadores SAAM ativos / em geofence
    saam_em_geo = [t for t in saam if t.get("insideGeofences")]
    if saam:
        if saam_em_geo:
            nomes = ", ".join((t.get("name") or "").split()[-1] for t in saam_em_geo[:3])
            insights.append(f"Frota SAAM: {len(saam)} no radar; {len(saam_em_geo)} em geofence agora ({nomes}).")
        else:
            insights.append(f"Frota SAAM: {len(saam)} rebocadores no radar, nenhum em geofence de manobra no momento.")
    else:
        insights.append("Frota SAAM sem posição AIS recente — monitorando reconexão.")

    # Rebocadores SAAM na base de rebocador (matching pelo TIPO da geofence,
    # estável mesmo que o nome mude — ex.: "BASE BRASCO" / "base rebocador").
    with geofence_lock:
        base_geo_names = {
            _normalize_text(g.get("name"))
            for g in geofences
            if g.get("type") == "base_rebocador" and g.get("isActive", True)
        }
    base_label = "BASE BRASCO"
    with geofence_lock:
        for g in geofences:
            if g.get("type") == "base_rebocador" and g.get("isActive", True):
                base_label = (g.get("name") or base_label).strip() or base_label
                break
    na_base = []
    for t in saam:
        geos = {_normalize_text(g) for g in (t.get("insideGeofences") or [])}
        if geos & base_geo_names:
            short = (t.get("name") or "").split()[-1] or (t.get("name") or "")
            na_base.append(short)
    if base_geo_names:
        if na_base:
            insights.append(f"{base_label}: {len(na_base)} rebocador(es) SAAM na base — {', '.join(na_base)}.")
        else:
            insights.append(f"{base_label}: nenhum rebocador SAAM na base agora — frota em operação/posicionamento.")

    # Concorrentes manobrando
    manobrando = [c for c in comps if c.get("insideManeuverGeofence")]
    if manobrando:
        alvo = manobrando[0]
        insights.append(
            f"Atenção: {alvo.get('company')} ({alvo.get('name')}) manobrando em {alvo.get('maneuverGeofenceName')}. "
            f"{len(manobrando)} concorrente(s) ativo(s)."
        )
    else:
        insights.append("Nenhum rebocador concorrente (WIL/CAM) em geofence de manobra agora — janela favorável.")

    # Próxima manobra programada
    if maneuvers:
        nxt = maneuvers[0]
        emp = str(nxt.get("empRb") or "").upper()
        dono = "SAAM" if emp == OWN_COMPANY_EMP_RB else (emp or "—")
        insights.append(
            f"Próxima manobra na programação: {nxt.get('vesselName','—')} às {nxt.get('pob','—')} — EMP.RB {emp or '—'} [{dono}]."
        )
    insights.append(f"Programação da praticagem: {m_total} manobra(s) na base de dados.")

    # Simultaneidade
    if simult:
        s0 = simult[0]
        insights.append(
            f"Simultaneidade às {s0.get('timeSlot')}: {s0.get('maneuvers')} manobras, "
            f"~{s0.get('estimatedTugsNeeded')} rebocadores necessários. Antecipar alocação."
        )

    # Mudanças de programação
    if changes.get("anyChange"):
        insights.append(
            f"Programação alterada: {changes.get('delayedCount',0)} atraso(s), "
            f"{changes.get('advancedCount',0)} adiantamento(s), {changes.get('addedCount',0)} entrada(s)."
        )

    # Meteocean
    wind = _wind_label(metocean)
    if wind:
        insights.append(f"Vento na Baía: {wind}. Avaliar reforço de rebocador em manobra sensível.")
    tide = metocean.get("tide")
    if isinstance(tide, str) and tide.strip():
        insights.append(f"Maré: {tide}")
    cur_kn = metocean.get("currentSpeedKn")
    if isinstance(cur_kn, (int, float)):
        cdir = metocean.get("currentDirection")
        ctxt = f"Corrente: {cur_kn:.2f} nós" + (f" para {cdir}" if cdir else "")
        if cur_kn >= 0.8:
            ctxt += " — atenção: acima de limites de manobra de vários terminais."
        insights.append(ctxt)

    # Market share (SAA = nós)
    if market:
        nosso = next((r for r in market if str(r.get("empRb")).upper() == OWN_COMPANY_EMP_RB), None)
        if nosso:
            insights.append(f"Market share SAAM: {nosso.get('sharePct',0)}% das manobras na base.")

    insights.append("KRATOS vigiando o porto em tempo real — um passo à frente.")
    return [s for s in insights if s]


@app.get("/api/kratos/insights")
def kratos_insights():
    ensure_live_worker_started()
    return {"ok": True, "generatedAt": get_now_iso(), "insights": _build_kratos_insights()}


@app.get("/api/geofences")
def get_geofences():
    ensure_geofences_loaded()
    with geofence_lock:
        return {"geofences": geofences}


@app.post("/api/geofences")
async def create_geofence(request: Request):
    payload = await request.json()
    geofence = {
        "id": str(uuid.uuid4()),
        "name": payload.get("name", "Novo geofence"),
        "type": payload.get("type", "berco"),
        "geometry": payload.get("geometry", {}),
        "fleetScope": payload.get("fleetScope", "all"),
        "isActive": bool(payload.get("isActive", True)),
        "color": payload.get("color", "#35c8ff"),
        "createdAt": get_now_iso(),
        "updatedAt": get_now_iso(),
    }
    with geofence_lock:
        geofences.append(geofence)
        save_geofences()
    return {"ok": True, "geofence": geofence}


@app.put("/api/geofences/{geofence_id}")
async def update_geofence(geofence_id: str, request: Request):
    payload = await request.json()
    with geofence_lock:
        for geofence in geofences:
            if geofence.get("id") == geofence_id:
                geofence["name"] = payload.get("name", geofence.get("name"))
                geofence["type"] = payload.get("type", geofence.get("type"))
                geofence["geometry"] = payload.get("geometry", geofence.get("geometry", {}))
                geofence["fleetScope"] = payload.get("fleetScope", geofence.get("fleetScope", "all"))
                geofence["isActive"] = bool(payload.get("isActive", geofence.get("isActive", True)))
                geofence["color"] = payload.get("color", geofence.get("color", "#35c8ff"))
                geofence["updatedAt"] = get_now_iso()
                save_geofences()
                return {"ok": True, "geofence": geofence}
    return {"ok": False, "error": "Geofence não encontrado"}


@app.delete("/api/geofences/{geofence_id}")
def delete_geofence(geofence_id: str):
    with geofence_lock:
        for g in geofences:
            if g.get("id") == geofence_id and _is_bg_interno_geofence(g):
                return {"ok": False, "error": "Geofence persistente nao pode ser apagado"}
        before = len(geofences)
        geofences[:] = [g for g in geofences if g.get("id") != geofence_id]
        if len(geofences) == before:
            return {"ok": False, "error": "Geofence não encontrado"}
        save_geofences()
    return {"ok": True}


@app.get("/api/geofences/occupancy")
def geofence_occupancy():
    return {"occupancy": build_geofence_occupancy()}


@app.get("/api/geofences/{geofence_id}/vessels")
def geofence_vessels(geofence_id: str):
    occupancy = build_geofence_occupancy()
    for item in occupancy:
        if item.get("geofenceId") == geofence_id:
            return {"ok": True, "vessels": item.get("insideVessels", []), "geofence": item}
    return {"ok": False, "error": "Geofence não encontrado"}


def _read_dashboard_html() -> str:
    path = FRONTEND_DIR / "dashboard.html"
    return path.read_text(encoding="utf-8")


def _json_for_inline_script(obj) -> str:
    """JSON seguro dentro de <script> (evita quebra por </script> ou < no payload)."""
    s = json.dumps(obj, ensure_ascii=False, default=str)
    return s.replace("<", "\\u003c")


def _dashboard_html_with_bootstrap() -> str:
    """Injeta window.__DASHBOARD_OVERVIEW__ para a página funcionar sem XHR a /api (ex.: proxy na 8080)."""
    html = _read_dashboard_html()
    try:
        payload = _json_for_inline_script(build_dashboard_overview_dict())
    except Exception:
        payload = _json_for_inline_script({"ok": False, "error": "Falha ao montar overview"})
    inject = f"<script>window.__DASHBOARD_OVERVIEW__={payload};</script>"
    if "</head>" in html:
        return html.replace("</head>", inject + "\n</head>", 1)
    return inject + html


@app.get("/dashboard")
def dashboard_page():
    """HTML com dados embutidos (bootstrap) + fallback por API no cliente."""
    return HTMLResponse(_dashboard_html_with_bootstrap())


@app.get("/dashboard/")
def dashboard_page_trailing_slash():
    return HTMLResponse(_dashboard_html_with_bootstrap())


@app.get("/dashboard.html")
def dashboard_page_html_file():
    """Alias com extensão (útil se o browser ou proxy esperar .html)."""
    return HTMLResponse(_dashboard_html_with_bootstrap())


@app.post("/api/mode")
async def set_mode(request: Request):
    global current_mode
    data = await request.json()
    if data.get("syncPraticagemSaa") in (True, 1, "1", "true"):
        return await _sync_saa_from_praticagem_impl()
    mode = data.get("mode", "live").lower()
    if mode != "live":
        return {"ok": False, "error": "Modo mock removido. Use apenas live."}
    if not AISSTREAM_API_KEY:
        return {"ok": False, "error": "AISSTREAM_API_KEY ausente no backend."}
    current_mode = "live"
    ensure_live_worker_started()
    return {"ok": True, "status": get_status()}

@app.post("/api/area")
async def set_area(request: Request):
    global current_area_key, last_subscription_update_monotonic
    data = await request.json()
    area = data.get("areaKey", "rio").lower()
    if area not in AREAS:
        return {"ok": False, "error": f"Área inválida: {area}"}
    if area == current_area_key:
        return {"ok": True, "status": get_status(), "updatedSubscription": False}

    if current_mode == "live":
        ensure_live_worker_started()
        now = time.monotonic()
        elapsed = now - last_subscription_update_monotonic
        if elapsed < 1.0:
            await asyncio.sleep(1.0 - elapsed)
        last_subscription_update_monotonic = time.monotonic()

    current_area_key = area
    if current_mode == "live":
        live_subscription_update_event.set()
    return {"ok": True, "status": get_status(), "updatedSubscription": current_mode == "live"}

# --- WebSocket relay ---
@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    try:
        await relay_cached_stream(websocket)
    except WebSocketDisconnect:
        pass
    except Exception as e:
        await websocket.send_json({"type": "ais_error", "payload": {"error": str(e)}})

async def relay_cached_stream(websocket: WebSocket):
    """
    Relay para o frontend usando apenas o buffer local (`recent_vessels`).
    A conexão com AISStream fica centralizada no worker de background.
    """
    ensure_live_worker_started()
    last_sent_seq = 0
    if recent_vessels:
        last_sent_seq = recent_vessels[-1]["_seq"]
    await websocket.send_json({"type": "status", "payload": get_status()})
    next_status_at = time.monotonic() + 1.0

    while True:
        await asyncio.sleep(0.2)
        new_items = [v for v in recent_vessels if v.get("_seq", 0) > last_sent_seq]
        for vessel_payload in new_items:
            await websocket.send_json({"type": "ais", "payload": vessel_payload})
            last_sent_seq = vessel_payload.get("_seq", last_sent_seq)
        now = time.monotonic()
        if now >= next_status_at:
            await websocket.send_json({"type": "status", "payload": get_status()})
            next_status_at = now + 1.0

async def relay_live(websocket: WebSocket):
    global live_connected, last_error, last_ais_message_at, total_messages
    total_messages = 0
    backoff_seconds = 2

    while True:
        live_connected = False
        last_error = None
        try:
            async with websockets.connect(
                AISSTREAM_URL,
                ping_interval=30,
                ping_timeout=30,
            ) as ais_ws:
                live_subscription_update_event.clear()
                await ais_ws.send(json.dumps(build_live_subscription()))
                live_connected = True
                backoff_seconds = 2
                await websocket.send_json({"type": "status", "payload": get_status()})

                while True:
                    if live_subscription_update_event.is_set():
                        live_subscription_update_event.clear()
                        await ais_ws.send(json.dumps(build_live_subscription()))
                        await websocket.send_json({"type": "status", "payload": get_status()})
                        continue

                    try:
                        msg = await asyncio.wait_for(ais_ws.recv(), timeout=1.0)
                    except asyncio.TimeoutError:
                        continue
                    except ConnectionClosed as e:
                        last_error = f"Conexao AISStream encerrada: {e}"
                        live_connected = False
                        await websocket.send_json({"type": "status", "payload": get_status()})
                        break

                    try:
                        data = json.loads(msg)
                        if "error" in data:
                            last_error = data["error"]
                            await websocket.send_json({"type": "ais_error", "payload": data})
                            await websocket.send_json({"type": "status", "payload": get_status()})
                            continue
                        vessel = extract_normalized_vessel(data)
                        if vessel:
                            total_messages += 1
                            last_ais_message_at = vessel["payload"]["timestamp"]
                            push_recent_vessel(vessel["payload"])
                            await websocket.send_json(vessel)
                            await websocket.send_json({"type": "status", "payload": get_status()})
                    except Exception as e:
                        last_error = f"Falha ao ler mensagem do AISStream: {e}"
                        await websocket.send_json({"type": "status", "payload": get_status()})
        except WebSocketDisconnect:
            break
        except Exception as e:
            last_error = f"Erro no socket AISStream: {e}"
            live_connected = False
            try:
                await websocket.send_json({"type": "ais_error", "payload": {"error": last_error}})
                await websocket.send_json({"type": "status", "payload": get_status()})
            except Exception:
                break

            await asyncio.sleep(backoff_seconds)
            backoff_seconds = min(backoff_seconds * 2, 30)

    live_connected = False

async def live_background_worker():
    global live_connected, last_error, last_ais_message_at, total_messages
    backoff_seconds = 2
    while True:
        live_connected = False
        last_error = None
        try:
            async with websockets.connect(
                AISSTREAM_URL,
                ping_interval=30,
                ping_timeout=30,
            ) as ais_ws:
                live_subscription_update_event.clear()
                await ais_ws.send(json.dumps(build_live_subscription()))
                live_connected = True
                backoff_seconds = 2
                while True:
                    if live_subscription_update_event.is_set():
                        live_subscription_update_event.clear()
                        await ais_ws.send(json.dumps(build_live_subscription()))
                    try:
                        msg = await asyncio.wait_for(ais_ws.recv(), timeout=1.0)
                    except asyncio.TimeoutError:
                        continue
                    except ConnectionClosed as e:
                        last_error = f"Conexao AISStream encerrada: {e}"
                        live_connected = False
                        break
                    data = json.loads(msg)
                    if "error" in data:
                        last_error = data["error"]
                        continue
                    vessel = extract_normalized_vessel(data)
                    if vessel:
                        total_messages += 1
                        last_ais_message_at = vessel["payload"]["timestamp"]
                        push_recent_vessel(vessel["payload"])
        except Exception as e:
            last_error = f"Erro no socket AISStream: {e}"
            live_connected = False
        await asyncio.sleep(backoff_seconds)
        backoff_seconds = min(backoff_seconds * 2, 30)

async def relay_mock(websocket: WebSocket):
    global total_messages, last_ais_message_at, live_connected
    live_connected = False
    while True:
        vessels = generate_mock_vessels(current_area_key)
        for vessel in vessels:
            print(f"[MOCK] Enviando alvo simulado: {vessel}")
            await websocket.send_json({"type": "ais", "payload": vessel})
            total_messages += 1
            last_ais_message_at = vessel["timestamp"]
        print("[MOCK] Status enviado para frontend.")
        await websocket.send_json({"type": "status", "payload": get_status()})
        await asyncio.sleep(2)

def extract_normalized_vessel(data):
    try:
        metadata = data.get("MetaData") or data.get("Metadata") or {}
        message_type = data.get("MessageType", "Unknown")
        message_wrapper = data.get("Message", {})
        message_body = message_wrapper.get(message_type) or next(iter(message_wrapper.values()), {})
        mmsi = str(metadata.get("MMSI") or message_body.get("UserID") or "desconhecido")
        cached = vessel_state_by_mmsi.get(mmsi, {})
        latitude = (
            metadata.get("latitude") or metadata.get("Latitude") or message_body.get("Latitude") or message_body.get("latitude")
        )
        longitude = (
            metadata.get("longitude") or metadata.get("Longitude") or message_body.get("Longitude") or message_body.get("longitude")
        )
        if latitude is None or longitude is None:
            return None
        incoming_name = metadata.get("ShipName") or message_body.get("Name")
        ship_name = incoming_name or cached.get("shipName") or "Sem nome"
        incoming_ship_type_code = normalize_ship_type_code(metadata, message_body)
        ship_type_code = incoming_ship_type_code if incoming_ship_type_code is not None else cached.get("shipTypeCode")
        is_saam_bgra = mmsi in SAAM_BGRA_MMSI_SET
        ship_category = "rebocador_servico" if is_saam_bgra else infer_ship_category(ship_type_code, ship_name)
        latitude_f = float(latitude)
        longitude_f = float(longitude)
        fallback_geofence = classify_geofence(latitude_f, longitude_f)

        ais_length_m, ais_beam_m = extract_ship_dimensions_meters(message_type, message_body)
        cached_ais_length = cached.get("lengthMetersAis")
        cached_ais_beam = cached.get("beamMetersAis")

        ais_ref = extract_ship_ref_offsets(message_type, message_body)
        ref_offsets = ais_ref if ais_ref is not None else cached.get("refOffsets")

        if ais_length_m is not None:
            length_source = "ais_dimension"
            length_m = ais_length_m
            beam_m = ais_beam_m if ais_beam_m is not None else cached_ais_beam
        elif cached_ais_length is not None:
            length_source = "ais_dimension_cached"
            length_m = float(cached_ais_length)
            beam_m = cached_ais_beam
        else:
            length_source = "estimated_category"
            length_m = estimate_length_from_category(ship_category)
            beam_m = None

        vessel_state_by_mmsi[mmsi] = {
            "shipName": ship_name,
            "shipTypeCode": ship_type_code,
            "shipCategory": ship_category,
            "fleet": SAAM_BGRA_FLEET_NAME if is_saam_bgra else None,
            "lengthMetersAis": ais_length_m if ais_length_m is not None else cached_ais_length,
            "beamMetersAis": ais_beam_m if ais_beam_m is not None else cached_ais_beam,
            "refOffsets": ref_offsets,
        }
        vessel_data = {
            "source": current_mode,
            "messageType": message_type,
            "mmsi": mmsi,
            "shipName": ship_name,
            "shipTypeCode": ship_type_code,
            "shipCategory": ship_category,
            "fleet": SAAM_BGRA_FLEET_NAME if is_saam_bgra else None,
            "isSaamBgra": is_saam_bgra,
            "geofence": fallback_geofence,
            "latitude": latitude_f,
            "longitude": longitude_f,
            "sog": float(message_body.get("Sog") or message_body.get("SpeedOverGround") or 0),
            "cog": float(message_body.get("Cog") or message_body.get("CourseOverGround") or 0),
            "heading": float(message_body.get("TrueHeading") or message_body.get("Heading") or 0),
            "navStatus": message_body.get("NavigationalStatus"),
            "lengthMeters": length_m,
            "beamMeters": beam_m,
            "refToBow": ref_offsets[0] if ref_offsets else None,
            "refToStern": ref_offsets[1] if ref_offsets else None,
            "refToPort": ref_offsets[2] if ref_offsets else None,
            "refToStarboard": ref_offsets[3] if ref_offsets else None,
            "lengthSource": length_source,
            "timestamp": metadata.get("time_utc") or metadata.get("timeUTC") or get_now_iso(),
            "raw": data
        }
        vessel_data["geofencesInside"] = get_vessel_geofences(vessel_data)
        vessel_data["inRebocadorBase"] = vessel_in_rebocador_base(vessel_data)
        if vessel_data["geofencesInside"] and not vessel_data.get("geofence"):
            vessel_data["geofence"] = vessel_data["geofencesInside"][0]
        latest_vessel_by_mmsi[mmsi] = vessel_data
        update_saam_nautical_miles(vessel_data)
        update_saam_fleet_geofence_stats(vessel_data)
        update_saam_operating_hours(vessel_data)
        save_vessels_snapshot()

        return {
            "type": "ais",
            "payload": vessel_data
        }
    except Exception:
        return None

def get_now_iso():
    return datetime.utcnow().isoformat()

def generate_mock_vessels(area_key):
    presets = {
        "suape": [
            {"mmsi": "710000101", "shipName": "TUG SUAPE ALFA", "shipTypeCode": 52, "latitude": -8.403, "longitude": -34.969, "sog": 8.2, "cog": 112, "lengthMeters": 32, "beamMeters": 11, "lengthSource": "mock"},
            {"mmsi": "710000102", "shipName": "NAVIO RECIFE STAR", "shipTypeCode": 70, "latitude": -8.377, "longitude": -34.912, "sog": 11.4, "cog": 221, "lengthMeters": 190, "beamMeters": 28, "lengthSource": "mock"}
        ],
        "santos": [
            {"mmsi": "710000201", "shipName": "TUG SANTOS BRAVO", "shipTypeCode": 52, "latitude": -23.992, "longitude": -46.307, "sog": 7.1, "cog": 41, "lengthMeters": 30, "beamMeters": 10, "lengthSource": "mock"},
            {"mmsi": "710000202", "shipName": "CARGUEIRO ATLANTICO SUL", "shipTypeCode": 70, "latitude": -24.038, "longitude": -46.283, "sog": 9.8, "cog": 78, "lengthMeters": 210, "beamMeters": 32, "lengthSource": "mock"}
        ],
        "rio": [
            {"mmsi": "710000301", "shipName": "TUG GUANABARA", "shipTypeCode": 52, "latitude": -22.882, "longitude": -43.146, "sog": 6.8, "cog": 132, "lengthMeters": 28, "beamMeters": 9, "lengthSource": "mock"},
            {"mmsi": "710000302", "shipName": "RIO BAY TRADER", "shipTypeCode": 70, "latitude": -22.930, "longitude": -43.180, "sog": 10.6, "cog": 212, "lengthMeters": 200, "beamMeters": 30, "lengthSource": "mock"}
        ],
        "paranagua": [
            {"mmsi": "710000401", "shipName": "TUG PARANAGUA DELTA", "shipTypeCode": 52, "latitude": -25.522, "longitude": -48.501, "sog": 5.2, "cog": 94, "lengthMeters": 29, "beamMeters": 10, "lengthSource": "mock"},
            {"mmsi": "710000402", "shipName": "PR PORT CONTAINER", "shipTypeCode": 70, "latitude": -25.565, "longitude": -48.472, "sog": 8.7, "cog": 176, "lengthMeters": 205, "beamMeters": 31, "lengthSource": "mock"}
        ],
        "bahia": [
            {"mmsi": "710000501", "shipName": "TUG TODOS OS SANTOS", "shipTypeCode": 52, "latitude": -12.915, "longitude": -38.661, "sog": 4.7, "cog": 63, "lengthMeters": 27, "beamMeters": 9, "lengthSource": "mock"},
            {"mmsi": "710000502", "shipName": "BAHIA MINERAL", "shipTypeCode": 80, "latitude": -12.806, "longitude": -38.485, "sog": 12.0, "cog": 149, "lengthMeters": 240, "beamMeters": 40, "lengthSource": "mock"}
        ],
        "brasil_sudeste": [
            {"mmsi": "710000601", "shipName": "COSTA SUDESTE 01", "shipTypeCode": 70, "latitude": -23.300, "longitude": -42.500, "sog": 13.6, "cog": 205, "lengthMeters": 215, "beamMeters": 33, "lengthSource": "mock"},
            {"mmsi": "710000602", "shipName": "COSTA SUDESTE 02", "shipTypeCode": 60, "latitude": -24.200, "longitude": -44.100, "sog": 12.9, "cog": 35, "lengthMeters": 160, "beamMeters": 26, "lengthSource": "mock"}
        ]
    }
    from datetime import datetime
    drift = ((datetime.utcnow().timestamp() % 10) / 10000)
    vessels = presets.get(area_key, presets["suape"])
    for i, vessel in enumerate(vessels):
        vessel = vessel.copy()
        vessel["latitude"] += drift * (1 if i == 0 else -1)
        vessel["longitude"] += drift * (-1 if i == 0 else 1)
        vessel["heading"] = vessel["cog"]
        vessel["navStatus"] = 0
        vessel["shipCategory"] = infer_ship_category(vessel.get("shipTypeCode"), vessel.get("shipName"))
        vessel["timestamp"] = get_now_iso()
        yield vessel


async def _praticagem_auto_sync_loop():
    sec = int(os.getenv("PRATICAGEM_AUTO_SYNC_SECONDS", "60") or "60")
    if sec <= 0:
        return
    await asyncio.sleep(10)
    while True:
        try:
            await _sync_saa_from_praticagem_impl()
            # Após atualizar as manobras, reflete no Obsidian (se devido).
            await _obsidian_auto_export_if_due("praticagem")
        except asyncio.CancelledError:
            break
        except Exception:
            pass
        try:
            await asyncio.sleep(sec)
        except asyncio.CancelledError:
            break


@app.on_event("startup")
async def startup_event():
    global live_worker_task, _praticagem_auto_sync_task, _obsidian_auto_sync_task
    load_geofences()
    load_tug_stats()
    load_saa_maneuvers()
    if current_mode == "live" and AISSTREAM_API_KEY and live_worker_task is None:
        live_worker_task = asyncio.create_task(live_background_worker())
    if int(os.getenv("PRATICAGEM_AUTO_SYNC_SECONDS", "60") or "60") > 0:
        _praticagem_auto_sync_task = asyncio.create_task(_praticagem_auto_sync_loop())
    if _obsidian_auto_enabled() and _obsidian_auto_interval() > 0:
        _obsidian_auto_sync_task = asyncio.create_task(_obsidian_auto_sync_loop())


@app.on_event("shutdown")
async def shutdown_event():
    global live_worker_task, _praticagem_auto_sync_task, _obsidian_auto_sync_task
    try:
        save_tug_stats()
    except Exception:
        pass
    if _praticagem_auto_sync_task:
        _praticagem_auto_sync_task.cancel()
        try:
            await _praticagem_auto_sync_task
        except (asyncio.CancelledError, Exception):
            pass
        _praticagem_auto_sync_task = None
    if _obsidian_auto_sync_task:
        _obsidian_auto_sync_task.cancel()
        try:
            await _obsidian_auto_sync_task
        except (asyncio.CancelledError, Exception):
            pass
        _obsidian_auto_sync_task = None
    if live_worker_task:
        live_worker_task.cancel()
        live_worker_task = None


@app.get("/frontend/dashboard.html")
def redirect_frontend_dashboard_to_canonical():
    """Ficheiro estático não inclui dados embutidos; força a rota /dashboard."""
    return RedirectResponse(url="/dashboard", status_code=307)


app.mount("/frontend", StaticFiles(directory=str(FRONTEND_DIR)), name="frontend")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=int(os.getenv("PORT", "8080")),
        reload=False,
    )
